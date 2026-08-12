from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent / "deliverables"
TODAY = date.today().strftime("%B %d, %Y")

NAVY = "0F2747"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
PALE_BLUE = "F7FAFF"
GREEN = "047857"
PALE_GREEN = "ECFDF5"
AMBER = "92400E"
PALE_AMBER = "FFFBEB"
RED = "B91C1C"
PALE_RED = "FEF2F2"
INK = "1E293B"
MUTED = "64748B"
LINE = "D8E2EE"
WHITE = "FFFFFF"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=None, bold=None, color=INK, italic=None, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_document(doc: Document, guide: bool, running_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.7 if guide else 11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18 if guide else 1.12

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 16, 8),
        ("Heading 2", 13.5, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.7 if guide else 11)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4 if guide else 5)
        style.paragraph_format.line_spacing = 1.18 if guide else 1.12

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(running_title)
    set_run(r, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)


def add_cover(doc: Document, kicker: str, title: str, subtitle: str, audience: str, version: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.upper())
    set_run(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_run(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_run(r, size=14, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [("Audience", audience), ("Version", version), ("Prepared", TODAY)]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], NAVY)
        set_cell_shading(row.cells[1], PALE_BLUE)
        row.cells[0].text = label
        row.cells[1].text = value
        for run in row.cells[0].paragraphs[0].runs:
            set_run(run, size=9.5, bold=True, color=WHITE)
        for run in row.cells[1].paragraphs[0].runs:
            set_run(run, size=9.5, color=INK)
    set_table_geometry(table, [1800, 7560])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("Fraud Transaction Detector | AML Anomaly Intelligence Platform")
    set_run(r, size=10, bold=True, color=NAVY)
    doc.add_page_break()


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_heading(doc: Document, text: str, level=2) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r)


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc: Document, label: str, text: str, kind="info") -> None:
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (PALE_GREEN, GREEN),
        "warning": (PALE_AMBER, AMBER),
        "risk": (PALE_RED, RED),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_run(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run(r, color=INK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9.2, bold=True, color=WHITE)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            for run in cell.paragraphs[0].runs:
                set_run(run, size=9.1, color=INK)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def add_flow(doc: Document, steps: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(steps))
    table.style = "Table Grid"
    widths = [9360 // len(steps)] * len(steps)
    widths[-1] += 9360 - sum(widths)
    for idx, (label, detail) in enumerate(steps):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE if idx % 2 == 0 else PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run(r, size=9, bold=True, color=BLUE)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(detail)
        set_run(r, size=8.4, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_run(r, size=8.7, color=NAVY, font="Consolas")
    set_table_geometry(table, [9360])


def build_user_guide() -> Path:
    doc = Document()
    setup_document(doc, guide=True, running_title="Fraud Transaction Detector | User Guide")
    add_cover(
        doc,
        "Operational Handbook",
        "Fraud Transaction Detector - Full User Guide",
        "Step-by-step operation of transaction intake, model comparison, governed training, and case review",
        "Administrators, AML analysts, reviewers, operators, and project evaluators",
        "2.0",
    )

    add_title(doc, "How to Use This Guide")
    add_para(doc, "This guide explains how to start the application, navigate every available Angular page, run the offline model-comparison workflow, operate the bank-scale Half-Space Trees lifecycle, configure production safeguards, and investigate generated cases. It describes the current implementation as of the document date.")
    add_callout(doc, "Important scope", "The system identifies suspicious or anomalous transactions. It does not declare that fraud has been proven. A human reviewer remains responsible for the final case decision.", "warning")
    add_heading(doc, "Contents", 2)
    add_bullets(doc, [
        "1. System overview and roles",
        "2. Installation prerequisites and startup",
        "3. Authentication and navigation",
        "4. Uploading transaction data",
        "5. Comparing anomaly models as data grows",
        "6. Operating governed incremental training",
        "7. Validating models and controlling layered rollout",
        "8. Model tuning, anomaly configuration, and cold start",
        "9. Case management and STR XML generation",
        "10. Daily operating procedure, troubleshooting, and glossary",
    ])

    add_title(doc, "1. System Overview and User Roles")
    add_para(doc, "The application consists of an Angular user interface, one Spring Boot modular-monolith backend, a Python FastAPI machine-learning service, and Microsoft SQL Server. Users work only through Angular or the Spring API. Spring owns database access and business decisions; Python owns model training and scoring.")
    add_flow(doc, [("Angular", "Port 4200\nUser interface"), ("Spring Boot", "Port 8080\nBusiness and database layer"), ("FastAPI", "Port 8000\nML training and scoring"), ("SQL Server", "Transactions, cases, configuration, audit")])
    add_heading(doc, "User roles", 2)
    add_table(doc, ["Role", "Primary responsibilities", "Restricted functions"], [
        ["REVIEWER", "Review cases, inspect model evidence, add notes, mark false positive, generate STR XML, create manual cases.", "Cannot access administrator configuration and training-operation pages."],
        ["ANALYST", "Operational review and investigation where assigned by deployment policy.", "Administrator-only model controls remain unavailable."],
        ["ADMIN / AML_ADMIN", "All reviewer functions plus training operations, tuning, anomaly configuration, cold-start configuration, validation, promotion, and rollback.", "Every governed deployment action requires an audit reason."],
    ], [1600, 4700, 3060])
    add_callout(doc, "Registration security", "Self-registration always creates a REVIEWER account. Administrative roles must be assigned through a controlled administrative/database process; they cannot be selected on the public registration page.", "info")

    add_title(doc, "2. Prerequisites and Startup")
    add_heading(doc, "Required software and services", 2)
    add_bullets(doc, [
        "Microsoft SQL Server with the fraud-transaction-detector database and all applicable migration scripts already executed.",
        "Java 17 and Maven for the Spring Boot application.",
        "Python virtual environment with dependencies from fraud-ml-service/requirements.txt.",
        "Node.js, npm, and Angular dependencies for fraud-transaction-ui.",
        "Ports 1433, 8000, 8080, and 4200 available locally, unless configuration is changed.",
    ])
    add_heading(doc, "Recommended startup order", 2)
    add_steps(doc, [
        "Start SQL Server and verify that the fraud-transaction-detector database is reachable.",
        "Start the Python ML service on port 8000.",
        "Start the Spring Boot application on port 8080.",
        "Start the Angular development server on port 4200.",
        "Open http://localhost:4200 and sign in.",
    ])
    add_heading(doc, "Commands", 2)
    add_code(doc, [
        "# Terminal 1 - Python ML service",
        "cd fraud-ml-service",
        ".\\.venv\\Scripts\\python.exe -m uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload",
        "",
        "# Terminal 2 - Spring Boot",
        "cd fraud-transaction-detector",
        "mvn spring-boot:run",
        "",
        "# Terminal 3 - Angular UI",
        "cd fraud-transaction-ui",
        "npm install   # first run or after dependency changes",
        "npm start",
    ])
    add_heading(doc, "Startup checks", 2)
    add_table(doc, ["Check", "Expected result", "If it fails"], [
        ["http://localhost:8080/health", "status = UP and modular-monolith architecture", "Check Java logs, database connectivity, and port 8080."],
        ["http://localhost:8000/docs", "FastAPI OpenAPI page loads", "Check the Python virtual environment and requirements."],
        ["http://localhost:4200", "Login page loads on a white interface", "Check npm output and browser console."],
    ], [2500, 3300, 3560])

    add_title(doc, "3. Authentication and Navigation")
    add_heading(doc, "Create a reviewer account", 2)
    add_steps(doc, [
        "Open the Register page from the login screen.",
        "Enter full name, username, and a password of at least six characters.",
        "Select Create account. The account is created with the REVIEWER role.",
        "Return to the Login page and sign in with the new credentials.",
    ])
    add_heading(doc, "Sign in and sign out", 2)
    add_para(doc, "On successful login, the browser stores the JWT and current-user information locally for the active session. Use the sign-out icon at the bottom of the side menu to clear the session. If the token expires, sign in again.")
    add_heading(doc, "Side-menu map", 2)
    add_table(doc, ["Menu item", "Purpose", "Access"], [
        ["Upload Data", "Import Excel or CSV transaction files.", "Authenticated"],
        ["Research Comparison", "Create a database snapshot and compare five offline anomaly models across 10%, 25%, 50%, and 100% chronological partitions.", "Authenticated"],
        ["Model Governance", "Review model evidence and control exact-version layered canary promotion or Isolation Forest rollback.", "Authenticated; actions are admin-only"],
        ["Training Operations", "Close business days, create immutable windows, export Parquet data, and train HST or Online OCSVM candidates.", "ADMIN / AML_ADMIN"],
        ["Case Management", "Review generated cases, evidence, notes, and final actions.", "Authenticated"],
        ["Manual Case", "Search transactions and create a case on demand.", "Authenticated"],
        ["Model Tuning", "Edit model-search and optimization configuration.", "ADMIN / AML_ADMIN"],
        ["Anomaly Config", "Configure the production batch fallback and its case threshold.", "ADMIN / AML_ADMIN"],
        ["Cold Start", "Configure behavior for accounts with insufficient transaction history.", "ADMIN / AML_ADMIN"],
    ], [1800, 5480, 2080])
    add_para(doc, "The side menu is collapsible on desktop and becomes a toggleable mobile menu on smaller screens.")

    add_title(doc, "4. Uploading Transaction Data")
    add_heading(doc, "Upload workflow", 2)
    add_steps(doc, [
        "Open Upload Data.",
        "Choose an Excel or CSV file containing transaction records.",
        "Confirm the selected filename and select Upload.",
        "Wait for the success popup. Review the returned success and failure counts if present.",
        "Correct rejected rows in the source file and upload a corrected batch if necessary.",
    ])
    add_heading(doc, "Expected transaction fields", 2)
    add_table(doc, ["Field group", "Representative fields", "Purpose"], [
        ["Identity", "TransactionID, AccountID", "Links transactions, accounts, predictions, alerts, and cases."],
        ["Transaction", "TransactionAmount, TransactionDate, TransactionType", "Core behavioral and chronological features."],
        ["Context", "Location, Channel, LoginAttempts", "Location-change, channel, and access-risk signals."],
        ["Customer/account", "CustomerAge, CustomerOccupation, AccountBalance", "Peer, demographic, and amount-to-balance context."],
        ["History", "PreviousTransactionDate", "Time-gap and behavioral history context when supplied."],
    ], [1800, 3200, 4360])
    add_callout(doc, "Operational behavior", "Uploading data and training are intentionally separate. Data may arrive through file upload or API ingestion; training can run later after business hours against transactions already stored in SQL Server.", "success")

    add_title(doc, "5. Comparing Anomaly Models as Data Grows")
    add_para(doc, "This is the primary academic and analytical feature. The Research Comparison page snapshots currently available database transactions, orders them from oldest to newest by transaction_date, and evaluates the same five offline models as the training population grows. These results never activate the layered production architecture.")
    add_heading(doc, "Run a comparison", 2)
    add_steps(doc, [
        "Open Research Comparison. The page initially shows only the comparison action and explanatory text.",
        "Select Train. No new upload is required; the system creates a database snapshot from stored transactions.",
        "The snapshot is divided into oldest 10%, 25%, 50%, and 100% partitions. A minimum of 200 rows is required.",
        "The five offline models are trained sequentially for every available partition.",
        "Monitor the progress label and wait for the Training complete popup.",
        "Review the best-fit recommendation, model ranking, and every model's growth detail on the same page.",
    ])
    add_heading(doc, "Models in the research comparison", 2)
    add_table(doc, ["Model", "What it detects", "Scale position"], [
        ["Isolation Forest", "Globally isolated observations through randomized tree partitions.", "Production batch fallback and offline comparison."],
        ["Local Outlier Factor (LOF)", "Points whose local density is substantially lower than their neighbors.", "Offline/capped research only at bank scale."],
        ["One-Class SVM", "A learned boundary around normal observations.", "Offline/capped research only; kernel training is not suitable for millions of rows."],
        ["Elliptic Envelope", "Outliers relative to robust covariance assumptions.", "Offline/capped research only; unsuitable for high-dimensional bank-scale production."],
        ["PCA Reconstruction", "Observations poorly reconstructed by a lower-dimensional representation.", "Offline research and diagnostic interpretation."],
    ], [2100, 4700, 2560])
    add_heading(doc, "How to interpret the report", 2)
    add_table(doc, ["Metric", "Meaning", "Preferred direction"], [
        ["Proxy quality", "Unsupervised separation quality on a common future holdout; it is not fraud accuracy.", "Higher"],
        ["Stability", "Agreement under bootstrap/resampling and small data perturbations.", "Higher"],
        ["Rate control", "How reasonably and consistently the model controls anomaly volume.", "Higher"],
        ["Fit score", "0.50 x Proxy Quality + 0.35 x Stability + 0.15 x Rate Control, reported on a 0-100 scale.", "Higher"],
        ["Anomaly rate", "Percentage of evaluated transactions classified as anomalous.", "Context-dependent; extreme rates require review"],
        ["Quality change", "Difference between the earliest and full-data partition quality.", "Stable or improving"],
        ["Latency and artifact size", "Operational cost per row and model storage footprint.", "Lower, subject to quality"],
    ], [1800, 5660, 1900])
    add_callout(doc, "Do not report false accuracy", "The source data has no complete confirmed-fraud label. Therefore, the comparison must not be described as accuracy, precision, recall, or F1 for the whole population. Those supervised measures become valid only after reliable reviewer labels exist.", "risk")
    add_heading(doc, "Recommended interpretation sequence", 2)
    add_steps(doc, [
        "Check whether quality improves, remains stable, or declines from 10% to 100%.",
        "Reject a high-quality model if its resampling stability is poor.",
        "Inspect anomaly rate and rate control so the model does not overwhelm investigators.",
        "Compare latency and artifact size before recommending a production candidate.",
        "Treat the top score as decision support, not an automatic production deployment.",
    ])

    add_title(doc, "6. Governed Incremental Training Operations")
    add_para(doc, "Training Operations is the bank-scale workflow for Half-Space Trees (HST) and Online One-Class SVM. Unlike the research comparison, it uses immutable Parquet datasets, bounded chunks, asynchronous jobs, and a formal model registry.")
    add_heading(doc, "Step 1 - Close the business day", 2)
    add_steps(doc, [
        "Open Training Operations as ADMIN or AML_ADMIN.",
        "Confirm Application health is UP.",
        "Select the business date and choose Close day.",
        "Wait for the Business day closed popup. Closed dates cannot silently change underneath a reproducible training window.",
    ])
    add_heading(doc, "Step 2 - Create the training window", 2)
    add_table(doc, ["Field", "Guidance"], [
        ["Training type", "DAILY_INCREMENTAL for normal end-of-day learning; WEEKLY_BATCH or FULL_REBUILD for controlled larger jobs; BACKTEST for historical evaluation."],
        ["Feature version", "Use the persisted feature contract expected by the trainer, for example AML_FEATURES_V2 or the configured current version."],
        ["Model type", "Use HALF_SPACE_TREES or ONLINE_ONE_CLASS_SVM according to the approved training objective."],
        ["Segment", "Leave blank for GLOBAL or enter an approved customer/business segment."],
        ["From / To", "Every date in the range must be closed."],
        ["Cutoff timestamp", "The latest event time permitted in the immutable training snapshot."],
    ], [2200, 7160])
    add_steps(doc, [
        "Complete all training-window fields.",
        "Select Create run.",
        "Select the new run in Recent pipeline activity.",
        "Choose Generate dataset. The export runs asynchronously.",
        "Use Refresh status until the run reaches DATASET_READY.",
        "Optionally enter a compatible base model version for incremental continuation.",
        "Choose Train candidate and refresh until CANDIDATE_READY or TRAINING_FAILED.",
    ])
    add_heading(doc, "Training-run statuses", 2)
    add_table(doc, ["Status", "Meaning", "Operator action"], [
        ["CREATED", "Window recorded but export not queued.", "Generate dataset."],
        ["QUEUED / EXPORTING", "Feature rows are being exported in bounded chunks.", "Wait and refresh."],
        ["DATASET_READY", "Parquet files and checksums are complete.", "Start candidate training."],
        ["TRAINING", "Python is learning from chronological record batches.", "Wait and refresh."],
        ["CANDIDATE_READY", "Immutable candidate artifact has been registered.", "Continue in Model Governance."],
        ["FAILED / TRAINING_FAILED", "Export or training stopped safely.", "Read the failure reason, correct the cause, and create/retry a controlled run."],
    ], [2200, 4300, 2860])

    add_title(doc, "7. Model Governance")
    add_para(doc, "The Model Governance page prevents a newly trained model from affecting production automatically. It exposes registry status, active production pointers, silent-validation evidence, and an immutable deployment history.")
    add_heading(doc, "Validate a candidate", 2)
    add_steps(doc, [
        "Allow the candidate to collect silent scores on real transactions. It must not affect votes or cases during this observation period.",
        "Open Model Governance and select the candidate version.",
        "Review the observation sample, anomaly rate, production alert rate, agreement, Jaccard overlap, score percentiles, daily stability, and reviewed outcomes.",
        "Choose Validate. A PASSED result changes the model to VALIDATED; it does not deploy it.",
        "If the result is INSUFFICIENT_DATA, continue collecting observations. If FAILED, investigate the displayed gate failure.",
    ])
    add_heading(doc, "Promote a validated model", 2)
    add_steps(doc, [
        "Select a model whose registry status is VALIDATED.",
        "Enter an audit reason of at least ten characters describing the governance decision.",
        "Choose Promote.",
        "Verify that the active pointer now identifies the promoted version and the audit timeline contains the promotion event.",
    ])
    add_heading(doc, "Roll back a champion", 2)
    add_steps(doc, [
        "Select the active CHAMPION and confirm that a previous compatible model is available.",
        "Enter a clear operational reason, such as abnormal alert volume or post-deployment drift.",
        "Choose Roll back.",
        "Verify that the previous model is restored atomically and the rollback appears in deployment history.",
    ])
    add_callout(doc, "Governance guarantee", "Validation, promotion, and rollback are separate actions. Artifact checksums are reverified before deployment, and every production change is authenticated, atomic, reasoned, auditable, and reversible while a previous champion is available.", "success")
    add_heading(doc, "Validate and activate the layered architecture", 2)
    add_steps(doc, [
        "Select the customer peer group in Layered architecture rollout and run shadow validation.",
        "Review alert volume, agreement, overlap, daily stability, synthetic recall, reviewed outcomes, latency, score availability, and the exact HST and Online OCSVM versions.",
        "Confirm the report is PASSED and recent enough for the configured deployment-age limit.",
        "Enter a small canary percentage and an audit reason, then activate the layered canary.",
        "Verify that the pointer locks the exact risk policy, HST, Online OCSVM, and validation versions.",
        "Expand exposure only after operational review. Account routing remains stable because the canary uses deterministic account hashing.",
        "Choose Rollback to Isolation Forest if volume, latency, availability, or investigation outcomes become unacceptable.",
    ])
    add_callout(doc, "Production decision method", "Layered production uses normalized component scores, configurable versioned weights, confidence handling, and hard-rule overrides rather than model vote counting.", "warning")

    add_title(doc, "8. Configuration Pages")
    add_heading(doc, "Model Tuning", 2)
    add_para(doc, "Use Model Tuning to edit configured optimization controls and model-specific search ranges. Reload saved values before changing them, modify only settings you understand, and save. Parameter search may use a capped recent sample while final fitting uses the complete selected partition.")
    add_bullets(doc, [
        "Change one related group at a time and record why it changed.",
        "Use the comparison page to observe whether tuning improves quality without harming stability or anomaly-rate control.",
        "Do not use large kernel-model searches on millions of rows.",
    ])
    add_heading(doc, "Anomaly Configuration", 2)
    add_para(doc, "Anomaly Configuration controls the compatibility decision path used while segments are still unpromoted. An explicit layered rollback pointer forces Isolation Forest with zero layered traffic. Do not interpret this page as the layered weighted-risk policy editor.")
    add_callout(doc, "Current production boundary", "LOF, kernel One-Class SVM, Elliptic Envelope, and PCA Reconstruction remain available for offline research reports but cannot influence production alerts or cases.", "warning")
    add_heading(doc, "Cold Start", 2)
    add_para(doc, "Cold-start controls define behavior when an account has insufficient trustworthy history. Boolean values use dropdowns. Review the minimum-history threshold and fallback behavior, then save. Cold start prevents unstable profile-derived decisions from being treated like mature-account evidence.")

    add_title(doc, "9. Case Management")
    add_heading(doc, "Automatic case creation", 2)
    add_para(doc, "A case is created only when the active production decision meets the configured suspicious threshold. The case preserves the transaction, account, priority, source alert, and latest model prediction evidence.")
    add_heading(doc, "Review a generated case", 2)
    add_steps(doc, [
        "Open Case Management and filter the queue if needed.",
        "Select a case to load its details.",
        "Review transaction ID, account, priority, assignment, origin, and status.",
        "Inspect Decision evidence: risk level, final risk score, policy/model versions, component scores, legacy anomaly votes where applicable, reason codes, and learning-control decision.",
        "Add investigation notes with supporting context.",
        "Choose Generate STR XML if activity remains suspicious, or Mark False Positive if the alert is explainable.",
    ])
    add_table(doc, ["Action", "System result", "Restriction"], [
        ["Generate STR XML", "Downloads a draft XML, updates case status to STR_GENERATED, and records a system note.", "A false-positive case cannot generate an STR."],
        ["Mark False Positive", "Updates the case and related alert outcome and records a system note.", "A case with generated STR cannot be marked false positive."],
        ["Add Note", "Appends an attributed, timestamped analyst note.", "Use factual language; notes form part of the audit trail."],
    ], [2100, 4660, 2600])
    add_callout(doc, "STR boundary", "The generated XML is a draft technical artifact. Institutional approval, regulatory validation, and submission remain outside this prototype unless separately integrated.", "warning")
    add_heading(doc, "Create a manual case", 2)
    add_steps(doc, [
        "Open Manual Case.",
        "Search by transaction ID, account ID, or available search text.",
        "Select the relevant transaction and review its details.",
        "Enter a clear case title, priority, and optional assignee.",
        "Choose Create Case. If a recent case already exists for the transaction, the system returns the existing case instead of duplicating it.",
    ])

    add_title(doc, "10. Daily Operating Procedure")
    add_heading(doc, "During business hours", 2)
    add_bullets(doc, [
        "Monitor transaction ingestion and application health.",
        "Review generated cases and prioritize HIGH-risk evidence.",
        "Record analyst notes and resolve false positives promptly so learning eligibility can be audited.",
        "Escalate unavailable ML service, abnormal case volume, or persistent API errors.",
    ])
    add_heading(doc, "After business hours", 2)
    add_steps(doc, [
        "Confirm ingestion for the day is complete.",
        "Close the business date in Training Operations.",
        "Create the approved DAILY_INCREMENTAL HST or Online OCSVM training run.",
        "Generate the immutable dataset and wait for DATASET_READY.",
        "Train the candidate and verify CANDIDATE_READY.",
        "Do not promote immediately; allow a silent observation period and run governance validation.",
    ])
    add_heading(doc, "Periodic research review", 2)
    add_bullets(doc, [
        "Run Research Comparison after meaningful data growth, not after every transaction.",
        "Compare 10%, 25%, 50%, and 100% learning curves using the same chronology.",
        "Document quality, stability, anomaly rate, latency, artifact size, and reviewer outcomes.",
        "Revisit model and threshold choices when drift, false-positive rate, or operational load changes.",
    ])

    add_title(doc, "11. Troubleshooting")
    add_table(doc, ["Symptom", "Likely cause", "Recommended action"], [
        ["Angular cannot reach the backend", "Spring is stopped, wrong port, expired token, or CORS mismatch.", "Check http://localhost:8080/health, sign in again, and verify allowed origins."],
        ["Training page reports too few rows", "Database snapshot contains fewer than 200 eligible transactions.", "Load more valid data and rerun later."],
        ["Dataset export fails", "Business date not closed, no eligible feature rows, output path collision, or schema/config issue.", "Read the run failure reason; close all dates and verify export configuration."],
        ["Candidate training fails", "Feature-schema mismatch, missing dependency, incompatible base version, or artifact error.", "Check Spring and Python logs; correct the issue without replacing the current champion."],
        ["Validate button disabled", "Wrong role or registry status is not CANDIDATE/VALIDATED.", "Use an admin account and select an eligible model."],
        ["Promote button disabled", "Model has not passed validation or user is not admin.", "Complete validation and verify role."],
        ["Rollback disabled", "Selected model is not active champion or has no previous model pointer.", "Select the active champion and confirm rollback availability."],
        ["No model evidence on an old case", "The transaction predates enriched prediction logging.", "Use transaction and case notes; evidence appears for newer predictions."],
        ["Unexpected database column error", "Required migration was not executed against the active database.", "Run the matching phase migration and verification script in SSMS, then restart Spring."],
    ], [2600, 3600, 3160])

    add_title(doc, "12. Glossary")
    add_table(doc, ["Term", "Definition"], [
        ["Anomaly", "A transaction whose behavior differs from learned or configured normal patterns; not proof of fraud."],
        ["Batch fallback", "Isolation Forest decision used when layered scoring is explicitly rolled back or cannot be accepted safely."],
        ["Candidate", "A newly trained immutable model version not approved for production."],
        ["Champion", "The active approved model for a model type and segment."],
        ["Challenger", "A compatible non-active model retained for comparison or rollback."],
        ["Cold start", "Account state with insufficient trustworthy history for mature profile-based scoring."],
        ["Feature version", "Versioned contract describing the exact persisted input vector expected by a model."],
        ["HST", "Half-Space Trees, a bounded-state incremental anomaly model used for bank-scale production learning."],
        ["Layered scoring", "Weighted production decision using rules, customer and peer behaviour, HST, and Online OCSVM evidence."],
        ["Model segment", "A routing scope such as RETAIL_GENERAL; GLOBAL is the fallback scope."],
        ["Online OCSVM", "Incremental One-Class SVM production model; different from the offline kernel One-Class SVM research model."],
        ["Parquet", "Columnar file format used for efficient, checksummed, chunked training datasets."],
        ["Silent scoring", "Challenger scoring that is logged for validation but cannot affect alerts or cases."],
        ["STR", "Suspicious Transaction Report; here generated as a draft XML for review."],
    ], [2200, 7160])

    path = OUTPUT_DIR / "Fraud_Transaction_Detector_Full_User_Guide_v2.docx"
    doc.save(path)
    return path


def build_system_report() -> Path:
    doc = Document()
    setup_document(doc, guide=False, running_title="Fraud Transaction Detector | System Design and Scalability Report")
    add_cover(
        doc,
        "Technical and Academic Report",
        "Fraud Transaction Detector - System Design, Implementation, and Scalability",
        "How the platform compares models as data grows, governs production learning, and supports explainable AML case review",
        "Project supervisors, professors, technical reviewers, architects, and AML stakeholders",
        "2.0",
    )

    add_title(doc, "Executive Summary")
    add_para(doc, "The Fraud Transaction Detector is an anomaly-intelligence and AML case-management platform designed around a central research question: how does model behavior change as more transaction data becomes available? The system answers that question through chronological learning-curve comparison while also providing a safer production path for bank-scale transaction volumes.")
    add_para(doc, "The implementation combines an Angular interface, a Spring Boot modular monolith, a Python FastAPI ML service, and Microsoft SQL Server. Spring owns transactional business logic, configuration, governance, and audit persistence. Python owns model fitting, artifact handling, and prediction. The browser never accesses the ML service or database directly.")
    add_callout(doc, "Core contribution", "The project separates research comparison from production decisioning. Five classical unsupervised models remain offline for learning-curve research, while production combines deterministic rules, customer and peer behaviour, Half-Space Trees, and Online One-Class SVM through a versioned weighted risk policy. Isolation Forest remains the rollback fallback.", "success")
    add_heading(doc, "Contents", 2)
    add_bullets(doc, [
        "1. Problem definition and design principles",
        "2. Architecture and technology choices",
        "3. Data ingestion, persistence, and feature engineering",
        "4. Model-comparison methodology",
        "5. Production model strategy and governance lifecycle",
        "6. Scalability for millions of daily transactions",
        "7. Case management, explainability, and feedback",
        "8. Security, reliability, and auditability",
        "9. Kafka and clustering roadmap",
        "10. Limitations, validation, and future work",
    ])

    add_title(doc, "1. Problem Definition")
    add_heading(doc, "Why transaction anomaly detection is difficult", 2)
    add_bullets(doc, [
        "Banks process very large and continuously growing transaction streams.",
        "Confirmed fraud labels are sparse, delayed, and affected by investigation selection bias.",
        "Normal behavior differs by customer, account maturity, location, channel, time, and business segment.",
        "A model that appears useful on 2,000 rows may become unstable, slow, or operationally noisy on millions.",
        "False positives consume investigator capacity and can harm customer experience.",
        "Production model changes require reproducibility, approval, rollback, and evidence rather than an automatic retrain-and-replace process.",
    ])
    add_heading(doc, "System objective", 2)
    add_para(doc, "The system detects suspicious deviations, logs model evidence, creates review cases only when the configured production decision is met, and supports a reviewer outcome. It deliberately avoids stating that an anomaly is confirmed fraud.")
    add_heading(doc, "Design principles", 2)
    add_table(doc, ["Principle", "How it is applied"], [
        ["Separation of concerns", "Angular presents workflows; Spring owns business/database rules; Python owns ML."],
        ["Chronological integrity", "Training partitions and holdouts respect transaction time and use oldest observations first."],
        ["Research-production separation", "Expensive classical models remain offline; scalable governed models control production."],
        ["Human accountability", "Cases end in reviewer decisions; no model claims fraud confirmation."],
        ["Reproducibility", "Training windows, cutoffs, feature versions, checksums, parameters, and artifacts are persisted."],
        ["Fail-safe change", "Candidates score silently before validation; promotion is explicit and rollback is atomic."],
        ["Bounded resource use", "Training data is exported/read in chunks; HST state does not grow with lifetime transaction count."],
    ], [2600, 6760])

    add_title(doc, "2. Current Architecture")
    add_flow(doc, [("Angular UI", "Authentication, model reports, governance, cases"), ("Spring Boot", "Business rules, APIs, SQL ownership, audit"), ("FastAPI", "Training, artifact verification, scoring"), ("SQL Server + files", "Operational data, registry, Parquet, model bundles")])
    add_heading(doc, "Why a modular monolith for Java", 2)
    add_para(doc, "The project originally considered separate Java authentication, gateway, and case services. They were consolidated into one Spring Boot modular monolith because the system is currently developed and operated as one product, uses one SQL Server database, and benefits from local transactions across ingestion, prediction logging, alert creation, and case creation.")
    add_bullets(doc, [
        "Lower deployment and debugging complexity for the current team and academic environment.",
        "No network hop between authentication, transaction, fraud, and case modules.",
        "Database transactions remain coherent across tightly coupled business actions.",
        "Internal package/module boundaries preserve a future path to extraction if independent scale or ownership justifies it.",
    ])
    add_heading(doc, "Component responsibilities", 2)
    add_table(doc, ["Component", "Responsibilities", "Explicit non-responsibilities"], [
        ["Angular", "Login, upload, comparison reports, training operations, governance, configuration, case review.", "No SQL access and no direct Python calls."],
        ["Spring Boot", "JWT, transaction ingestion, account context, persisted features, training orchestration, model registry, deployment pointers, alerts, cases, STR XML.", "Does not fit ML algorithms."],
        ["FastAPI", "Offline research training, HST and Online OCSVM incremental training, artifact bundles, and persisted-feature scoring.", "Does not access SQL Server directly, aggregate final business risk, or decide case workflow."],
        ["SQL Server", "System of record for transactions, configuration, features, learning decisions, predictions, registry, validation, deployments, alerts, and cases.", "Does not perform ML training."],
        ["Parquet/artifact storage", "Immutable chunked datasets and versioned model bundles with checksums.", "Does not replace operational SQL records."],
    ], [1700, 4900, 2760])
    add_heading(doc, "Application ports and call direction", 2)
    add_code(doc, [
        "Browser :4200 -> Spring Boot :8080 -> FastAPI :8000",
        "                              |",
        "                              +-> SQL Server :1433",
        "                              +-> Parquet/model artifact paths",
    ])

    add_title(doc, "3. Data Ingestion and Persistence")
    add_heading(doc, "Two ingestion modes", 2)
    add_para(doc, "Bulk upload accepts Excel and CSV for historical or batch loading. The transaction API supports ongoing system integration. Both routes persist transactions in SQL Server, allowing later training to run independently of file upload.")
    add_heading(doc, "Why training is decoupled from upload", 2)
    add_bullets(doc, [
        "A bank may ingest transactions continuously through APIs throughout the day.",
        "Training can be scheduled after business hours or triggered by drift instead of after each upload.",
        "One database snapshot can represent all eligible data regardless of its ingestion channel.",
        "Operational ingestion latency is not blocked by expensive model fitting.",
    ])
    add_heading(doc, "Important persistence groups", 2)
    add_table(doc, ["Data group", "Examples", "Purpose"], [
        ["Operational", "transactions, bulk_upload_batches", "Immutable transaction history and ingestion audit."],
        ["Prediction", "fraud_prediction_logs, fraud_alerts", "Model inputs/outputs, scores, risk, reasons, and alert review."],
        ["Cases", "case records, notes, str_reports", "Human investigation and reporting trail."],
        ["Configuration", "app_config, anomaly configuration", "Runtime thresholds, feature/training settings, and safety controls."],
        ["AML feature/profile", "persisted feature vectors, observed/trusted profile state, learning eligibility", "Stable model input and contamination-safe behavioral learning."],
        ["Training/governance", "aml_business_days, aml_training_runs, aml_model_registry, aml_model_validations, aml_active_models, aml_model_deployments", "Reproducible training and controlled lifecycle."],
    ], [1900, 4200, 3260])

    add_title(doc, "4. Feature Engineering and Learning Safety")
    add_heading(doc, "Behavioral feature strategy", 2)
    add_para(doc, "A transaction is evaluated relative to both its absolute characteristics and the account's historical behavior. Representative persisted features include time since previous transaction, location change, amount relative to account averages and maxima, amount z-score, rolling-window amount ratios, amount-to-balance ratio, transaction hour/day, night/weekend flags, login risk, peer context, velocity, novelty, and profile confidence.")
    add_heading(doc, "Observed profile versus trusted profile", 2)
    add_para(doc, "A naive online learner can contaminate its own normal profile by learning from suspicious transactions. The implementation separates observed behavior from trusted learning eligibility. Every transaction may be observed and logged, but only transactions satisfying policy can update trusted profiles or enter incremental/batch learning.")
    add_flow(doc, [("Transaction", "Persist immutable event"), ("Feature context", "Load historical profile and peer state"), ("Prediction", "Score persisted versioned vector"), ("Learning decision", "Eligible, quarantined, or excluded"), ("Profile update", "Trusted state changes only when allowed")])
    add_heading(doc, "Cold-start handling", 2)
    add_para(doc, "New or low-history accounts lack reliable personal baselines. Cold-start configuration prevents weak profile evidence from being interpreted like mature-account behavior. Profile confidence, peer context, and safe fallback logic reduce unstable early decisions.")

    add_title(doc, "5. Model-Comparison Methodology")
    add_heading(doc, "Central experiment", 2)
    add_para(doc, "For the same database snapshot, transactions are ordered by transaction_date and divided into nested oldest-first partitions: 10%, 25%, 50%, and 100%. Each of the five research models is trained and evaluated for every available partition. The approach reveals whether more data improves, stabilizes, or degrades model behavior.")
    add_flow(doc, [("10% oldest", "Early-data baseline"), ("25% oldest", "Initial growth"), ("50% oldest", "Mid-scale behavior"), ("100%", "Full snapshot behavior")])
    add_heading(doc, "Fairness controls", 2)
    add_bullets(doc, [
        "The same chronological snapshot and ordering policy are used for every model.",
        "A common future holdout is used rather than evaluating only on each model's training rows.",
        "Random seeds and tuning limits are configurable and recorded.",
        "Metrics include computational cost, not only anomaly separation.",
        "The browser receives summaries and sampled/aggregated evidence rather than millions of points.",
    ])
    add_heading(doc, "Current research models", 2)
    add_table(doc, ["Model", "Strength", "Reason it is not the universal bank-scale answer"], [
        ["Isolation Forest", "Efficient global anomaly separation and useful batch baseline.", "Still non-incremental in common implementations; should train on a controlled representative sample at very large scale."],
        ["LOF", "Strong local-density interpretation.", "Full-data neighbor search and memory cost do not scale safely to millions; novelty mode still requires retained fit structure."],
        ["Kernel One-Class SVM", "Flexible nonlinear normal boundary.", "Kernel time/memory growth becomes impractical on very large datasets."],
        ["Elliptic Envelope", "Interpretable robust covariance outliers when assumptions hold.", "High-dimensional categorical/behavioral features and non-Gaussian distributions violate assumptions and increase cost."],
        ["PCA Reconstruction", "Useful reconstruction-error view and dimensional diagnostics.", "Linear components may miss nonlinear behavior; batch decomposition needs careful incremental redesign."],
    ], [2200, 3500, 3660])
    add_heading(doc, "Unsupervised scoring", 2)
    add_table(doc, ["Measure", "Definition and role"], [
        ["Proxy quality", "Measures unsupervised separation/structure on a common validation set. It is evidence of anomaly usefulness, not confirmed-fraud accuracy."],
        ["Resampling stability", "Measures consistency of anomaly decisions when the training sample is bootstrapped or perturbed."],
        ["Rate control", "Rewards operationally plausible and controlled anomaly volumes rather than indiscriminate flagging."],
        ["Fit score", "0.50 x Proxy Quality + 0.35 x Stability + 0.15 x Rate Control, each normalized to 0-100."],
        ["Growth trend", "Tracks the change in quality and behavior from early partitions to the full snapshot."],
        ["Operational cost", "Training duration, prediction latency per row, and artifact size."],
    ], [2400, 6960])
    add_callout(doc, "Scientific limitation", "Without representative confirmed labels, the project cannot honestly claim population accuracy, precision, recall, F1, or a confusion matrix. Reviewer outcomes later provide labeled evidence, but that sample may be selection-biased because only alerted transactions are reviewed.", "warning")

    add_title(doc, "6. Production Model Strategy")
    add_heading(doc, "Why production differs from research", 2)
    add_para(doc, "A model can be academically useful yet operationally unsafe at bank scale. The production strategy therefore prioritizes bounded memory, chronological incremental learning, auditable artifacts, predictable latency, and safe fallback behavior.")
    add_heading(doc, "Current production decision hierarchy", 2)
    add_steps(doc, [
        "Evaluate deterministic velocity and AML rules, including hard-rule overrides.",
        "Calculate normalized customer and peer-group behavioural scores with confidence metadata.",
        "Obtain normalized Half-Space Trees and Online One-Class SVM scores from exact registered versions.",
        "Combine component evidence using the versioned weighted risk policy.",
        "Accept layered output only for an account selected by the segment's deterministic canary and only when policy/model versions match the active pointer.",
        "Use the compatibility result for unpromoted accounts and the Isolation Forest fallback for an explicit rollback pointer or unsafe layered result.",
    ])
    add_heading(doc, "Incremental production-model rationale", 2)
    add_bullets(doc, [
        "Incremental learning processes rows sequentially and supports end-of-day updates.",
        "Model state is bounded by tree count and height rather than lifetime row count.",
        "The trainer reads Parquet in configurable record batches.",
        "Chronological learning matches the temporal nature of transaction behavior.",
        "HST provides bounded-state tree evidence while Online OCSVM provides complementary incremental boundary evidence.",
        "Candidate and complete layered scoring can run silently before affecting any production decision.",
    ])
    add_heading(doc, "Offline-only boundary", 2)
    add_para(doc, "LOF, kernel One-Class SVM, Elliptic Envelope, and PCA Reconstruction are retained for partition comparison and research reporting but are explicitly blocked from production-v2 decisions. This preserves the thesis feature without exposing algorithms whose full-data behavior is unsuitable for millions of transactions.")

    add_title(doc, "7. Reproducible Training Lifecycle")
    add_flow(doc, [("Close date", "Freeze business-day eligibility"), ("Create run", "Window, cutoff, feature/model version"), ("Export", "Chunked checksummed Parquet"), ("Train", "HST / Online OCSVM candidates"), ("Shadow", "Layered comparison evidence"), ("Canary", "Exact-version segment pointer")])
    add_heading(doc, "Immutable dataset export", 2)
    add_para(doc, "Spring uses keyset pagination to export only eligible persisted feature rows. Files are written in bounded Parquet parts. A manifest records columns, part row counts, SHA-256 checksums, and an aggregate dataset checksum. The full dataset never has to reside in Java or Python memory.")
    add_heading(doc, "Immutable artifact registration", 2)
    add_para(doc, "Python writes the model state, feature schema, parameters, metrics, and manifest into a versioned candidate bundle. Spring independently verifies the deterministic bundle checksum before registering the candidate. Candidate registration and training-run completion occur transactionally.")
    add_heading(doc, "Champion-challenger validation", 2)
    add_table(doc, ["Metric/gate", "Purpose"], [
        ["Minimum silent predictions", "Prevents decisions on a tiny observation sample."],
        ["Candidate anomaly rate", "Rejects candidates that flag implausibly few or many transactions."],
        ["Agreement and Jaccard", "Measures overlap and disagreement with the current production decision."],
        ["P50/P95/P99 score separation", "Checks whether the upper anomaly tail is distinguishable from typical scores."],
        ["Daily anomaly-rate deviation", "Measures day-to-day operational stability."],
        ["Reviewed precision", "Uses reviewed overlaps when enough outcomes exist, while acknowledging selection bias."],
    ], [3100, 6260])
    add_heading(doc, "Promotion and rollback", 2)
    add_para(doc, "Only ADMIN or AML_ADMIN can promote or roll back. Individual model promotion verifies candidate evidence and artifacts. Layered architecture promotion additionally requires a recent passing segment report, locks exact policy/HST/Online OCSVM versions, and activates deterministic account-level canary traffic. A unique action ID makes requests idempotent. Architecture rollback atomically sets the segment to Isolation Forest fallback with zero layered traffic.")

    add_title(doc, "8. Scalability for Millions of Transactions")
    add_heading(doc, "Problems and implemented responses", 2)
    add_table(doc, ["Scalability problem", "Implemented response", "Benefit"], [
        ["Millions of rows cannot be sent as one JSON payload", "Spring exports immutable Parquet files; Python reads them directly in batches.", "Lower serialization cost and bounded memory."],
        ["Full lifetime data cannot fit in one process", "Keyset-paginated export, configurable part size, record-batch decoding, and incremental learning.", "Memory use depends on chunk/model settings, not total rows."],
        ["Expensive algorithms become impractical", "Production uses HST; costly classical models are capped/offline; Isolation Forest uses controlled batch fitting.", "Predictable resource use without losing research comparison."],
        ["Training must not block transaction APIs", "Dataset export and HST training are asynchronous run states.", "Ingestion remains available while background jobs progress."],
        ["Continuous retraining may learn anomalies", "Observed/trusted profiles and learning-eligibility decisions quarantine unsafe examples.", "Reduces self-contamination."],
        ["Model changes can create operational incidents", "Silent challengers, validation gates, explicit promotion, active pointers, and rollback.", "Safer deployment and fast recovery."],
        ["Different customer populations behave differently", "Segment-specific champions with GLOBAL fallback.", "Targeted behavior without breaking unmatched segments."],
        ["Browser cannot plot millions of points", "Reports use metrics, learning curves, representative samples, and planned density aggregation.", "Responsive visualization with statistically controlled sampling."],
    ], [2700, 4300, 2360])
    add_heading(doc, "Horizontal scale path", 2)
    add_para(doc, "Spring remains stateless at the HTTP layer apart from SQL/filesystem dependencies and can be replicated behind a load balancer after shared storage and job coordination are externalized. Prediction workers can load the same approved versioned artifact and scale horizontally. Segment pointers in SQL provide deterministic routing.")
    add_heading(doc, "Operational metrics required at scale", 2)
    add_bullets(doc, [
        "Transaction throughput and ingestion lag.",
        "Prediction p50, p95, and p99 latency and ML-service error rate.",
        "Alert and case creation rate by segment, model version, and channel.",
        "False-positive and STR-generated outcome rates.",
        "Feature drift, score-distribution drift, model freshness, and data completeness.",
        "Training duration, exported/learned rows, memory, artifact size, retries, and failed runs.",
    ])

    add_title(doc, "9. Runtime Prediction and Case Flow")
    add_flow(doc, [("Ingest", "Persist transaction"), ("Feature", "Compute/persist versioned vector"), ("Components", "Rules + behaviour + HST + OCSVM"), ("Aggregate", "Versioned weighted risk"), ("Route", "Canary or IF fallback"), ("Case", "Human review and outcome")])
    add_heading(doc, "Explainability evidence", 2)
    add_para(doc, "Prediction logs preserve final risk level and score, risk-policy version, exact model versions, component evidence, anomaly votes only for compatibility records, reason codes, learning decision, and timestamp. Case responses expose the latest decision evidence so investigators can see why the case exists and which architecture produced it.")
    add_heading(doc, "Reviewer outcomes", 2)
    add_para(doc, "A reviewer can generate a draft STR XML or mark the case false positive. Notes and actions form an audit trail. Outcomes also support future supervised evaluation and learning-eligibility policy, but the system distinguishes reviewed-alert precision from unbiased population precision.")

    add_title(doc, "10. Security, Reliability, and Auditability")
    add_table(doc, ["Control", "Implementation"], [
        ["Authentication", "JWT issued by the Spring authentication module; browser requests use bearer tokens."],
        ["Role boundary", "Administrative Angular routes and governance actions require ADMIN/AML_ADMIN; self-registration creates REVIEWER only."],
        ["CORS", "Browser origins are explicitly configured for the Angular host."],
        ["Database ownership", "Only Spring accesses SQL Server, reducing duplicated persistence logic and credentials."],
        ["Checksums", "Training datasets and model artifacts use SHA-256 verification."],
        ["Idempotency", "Deployment action IDs prevent repeated promotion/rollback from being applied twice."],
        ["Atomicity", "Pointer, registry-status, and deployment-event updates occur in a single database transaction."],
        ["Fail-safe scoring", "Unavailable ML returns auditable allow-and-log fallback; it does not invent a fraud classification."],
        ["Traceability", "Run IDs, timestamps, actors, reasons, model/feature versions, scores, and reviewer notes are persisted."],
    ], [2600, 6760])
    add_callout(doc, "Production hardening still required", "The demonstration defaults, including local database credentials and JWT secret fallback, must be replaced with a secrets manager, TLS, centralized identity, stricter authorization enforcement, and production monitoring before institutional deployment.", "risk")

    add_title(doc, "11. Kafka Architecture Roadmap")
    add_callout(doc, "Current status", "Kafka is a planned bank-scale extension, not part of the current runtime. The current system uses synchronous API/file ingestion, SQL persistence, asynchronous in-process training jobs, and Parquet artifacts.", "warning")
    add_heading(doc, "What Kafka would solve", 2)
    add_bullets(doc, [
        "Durable transaction-event transport with replay.",
        "Ordering within an account partition when partitioned by account ID.",
        "Back-pressure isolation between ingestion, feature generation, prediction, audit, case, and training consumers.",
        "Independent horizontal scaling and deployment of consumers.",
        "A reliable event source for rebuilding state after failure or model/feature changes.",
    ])
    add_heading(doc, "What Kafka would not do", 2)
    add_para(doc, "Kafka would not train models, calculate clusters, replace model governance, or automatically make distributed consumers correct. Feature-state management, idempotency, schema evolution, dead-letter handling, and exactly-once business effects still require explicit design.")
    add_heading(doc, "Suggested topics", 2)
    add_table(doc, ["Topic", "Example consumers"], [
        ["transactions.raw", "Validation, persistence, feature stream processor"],
        ["transactions.features", "Prediction workers, feature store, monitoring"],
        ["anomaly.predictions", "Audit persistence, alert/case policy, monitoring"],
        ["model.training.commands", "Python training workers"],
        ["model.training.results", "Spring registry/orchestration"],
        ["model.lifecycle.events", "Prediction-worker cache refresh, audit, monitoring"],
        ["case.events", "Case workflow, notifications, downstream reporting"],
    ], [3000, 6360])
    add_heading(doc, "Recommended adoption sequence", 2)
    add_steps(doc, [
        "Introduce an outbox pattern so committed SQL transactions publish durable events without dual-write loss.",
        "Move feature generation to a stateful stream processor partitioned by account ID.",
        "Persist date-partitioned feature events to durable Parquet/object storage.",
        "Run horizontally scalable prediction consumers using the active-model pointer and versioned artifact cache.",
        "Publish prediction and lifecycle events; make case creation idempotent by transaction/action identity.",
        "Add consumer-lag, replay, schema-registry, dead-letter, and disaster-recovery procedures.",
    ])

    add_title(doc, "12. Clustering Evaluation and Visualization Roadmap")
    add_callout(doc, "Current status", "The current implemented UI focuses on anomaly-model comparison. Full clustering models, clustering metrics, and cluster visualizations discussed for the next research stage are not yet complete and should be presented as planned work.", "info")
    add_heading(doc, "Recommended scalable clustering models", 2)
    add_table(doc, ["Model", "Use"], [
        ["MiniBatchKMeans", "Primary scalable partitioning baseline with bounded mini-batches and fast assignment."],
        ["BIRCH", "Incremental clustering through a compact clustering-feature tree."],
        ["Bisecting KMeans", "Optional batch challenger for hierarchical splits on controlled datasets."],
        ["HDBSCAN", "Capped-sample exploratory challenger for variable-density clusters and noise; not a full-million-row browser workflow."],
    ], [2800, 6560])
    add_heading(doc, "Clustering measures", 2)
    add_table(doc, ["Measure", "Interpretation", "Direction"], [
        ["Silhouette score", "Cohesion versus separation for each sampled point.", "Higher"],
        ["Davies-Bouldin index", "Average similarity between each cluster and its most similar neighbor.", "Lower"],
        ["Calinski-Harabasz index", "Between-cluster dispersion relative to within-cluster dispersion.", "Higher"],
        ["ARI / NMI stability", "Agreement between cluster assignments across resamples or time windows.", "Higher"],
        ["Cluster balance", "Whether the solution is dominated by one cluster or excessive tiny clusters.", "Context-dependent"],
        ["Noise rate", "Proportion left unassigned by density-based methods.", "Context-dependent"],
        ["Runtime metrics", "Training time, memory, assignment latency, and throughput.", "Lower cost at acceptable quality"],
    ], [2600, 4900, 1860])
    add_heading(doc, "Visualization at scale", 2)
    add_para(doc, "PCA can reduce feature dimensions for a two-dimensional explanatory plot, but PCA is not itself a clustering algorithm. UMAP may be used for exploratory display on a fixed sample. The UI should never send millions of points to the browser; it should use a reproducible representative sample, hexbin/density aggregation, cluster centroids, and summarized distributions.")

    add_title(doc, "13. Validation and Quality Assurance")
    add_heading(doc, "Current automated validation", 2)
    add_bullets(doc, [
        "Angular production build completes successfully.",
        "Spring Boot automated suite contains 101 passing tests across feature engineering, behavioural scorers, rules, weighted aggregation, profiles, learning eligibility, training export, HST/Online OCSVM lifecycle, layered validation, canary routing, rollback, comparison, and fraud-review behavior.",
        "Python suite contains 10 passing tests covering artifact reproducibility, incremental HST and Online OCSVM training, persisted-feature alignment, production/shadow score contracts, deprecated-route metadata, finite-value validation, and offline-model rejection from v2.",
    ])
    add_heading(doc, "Manual acceptance scenarios", 2)
    add_bullets(doc, [
        "Upload valid and invalid batches and verify row counts/error handling.",
        "Run 10/25/50/100 model comparison and inspect growth trends.",
        "Create a closed-date HST run, export Parquet, train a candidate, and verify checksums/statuses.",
        "Accumulate layered shadow predictions, validate exact versions, activate and expand a deterministic canary, and then roll back to Isolation Forest using an audit reason.",
        "Create anomalous test transactions for account AC00455 and verify evidence/case behavior.",
        "Mark one case false positive and generate draft STR XML for another; verify incompatible actions are blocked.",
    ])

    add_title(doc, "14. Limitations and Future Work")
    add_table(doc, ["Area", "Current limitation", "Recommended next step"], [
        ["Labels", "No representative complete fraud ground truth.", "Use reviewer outcomes, delayed confirmations, and controlled backtesting while measuring selection bias."],
        ["Streaming", "Kafka/stateful stream processing not yet implemented.", "Adopt outbox + Kafka + account-partitioned feature consumers."],
        ["Storage", "Local paths are acceptable for development but not resilient shared storage.", "Move Parquet and artifacts to versioned object storage with retention and access controls."],
        ["Job control", "Asynchronous jobs expose status but lack full distributed queue, cancellation, and retry orchestration.", "Introduce durable job commands, worker leases, retry policy, and resource quotas."],
        ["Clustering", "Metrics and visuals are planned but not complete.", "Implement MiniBatchKMeans/BIRCH track with fixed-sample metrics and density visualization."],
        ["Security", "Development secrets and local transport remain.", "Use TLS, secrets manager, enterprise identity, server-side RBAC, and security audit."],
        ["Observability", "Basic health and stored audit exist; production telemetry is incomplete.", "Add metrics, distributed tracing, structured logs, dashboards, and alerting."],
        ["STR", "Generated XML is a draft and not jurisdiction-certified.", "Map to the required regulator schema and add approval/submission workflow."],
    ], [1800, 3740, 3820])

    add_title(doc, "Conclusion")
    add_para(doc, "The project evolved from a small unsupervised fraud prototype into a controlled anomaly-intelligence platform. Its main academic value is retained: stakeholders can observe how multiple offline models change as chronological data grows. Production instead uses persisted features, deterministic controls, customer and peer behaviour, incremental HST and Online OCSVM evidence, weighted policy aggregation, shadow validation, deterministic canaries, and Isolation Forest rollback to create a credible path toward bank-scale operation.")
    add_para(doc, "The result is not a claim of autonomous fraud detection. It is a governed decision-support system that records evidence, limits unsafe learning, supports human investigation, and provides a clear roadmap for Kafka streaming, scalable clustering analysis, distributed workers, stronger security, and regulatory integration.")

    add_title(doc, "Appendix A - Key API Groups")
    add_table(doc, ["API group", "Representative purpose"], [
        ["/api/auth", "Register reviewer, login, current user, token validation."],
        ["/api/v1/uploads", "Bulk transaction upload and batch review."],
        ["/api/v1/transactions", "Create/search/retrieve transactions."],
        ["/api/v1/anomaly-model-comparisons", "Database snapshots, partitions, offline training, comparison configuration."],
        ["/api/v1/aml/training-runs", "Create/list/status/export/train AML runs."],
        ["/api/v1/aml/models", "Registry, validation, active pointers, promotion, rollback, deployment history."],
        ["/api/v1/aml/layered-shadow", "Layered prediction evidence, scenario labels, and architecture validation."],
        ["/api/v1/aml/layered-deployments", "Segment canary activation, expansion, active pointers, history, and rollback."],
        ["/api/cases", "Case list/detail, notes, false positive, STR XML."],
        ["FastAPI /api/v2/fraud/predict", "Production prediction from persisted versioned features."],
        ["FastAPI /api/v1/aml/training/incremental", "Incremental HST training from verified Parquet."],
    ], [3600, 5760])

    add_title(doc, "Appendix B - Model Lifecycle States")
    add_flow(doc, [("CANDIDATE", "Trained and registered"), ("VALIDATED", "Passed silent gates"), ("CHAMPION", "Active approved version"), ("CHALLENGER", "Previous or comparative version"), ("RETIRED", "No longer eligible")])
    add_para(doc, "A failed validation does not automatically delete a candidate. It remains available for investigation, retuning, or explicit later lifecycle action. Validation alone never changes the active production pointer.")

    path = OUTPUT_DIR / "Fraud_Transaction_Detector_System_Design_Implementation_and_Scalability_v2.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(build_user_guide())
    print(build_system_report())
