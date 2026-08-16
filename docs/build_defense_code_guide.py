from __future__ import annotations

from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "deliverables" / "Fraud_Transaction_Detector_Defense_Code_and_Business_Guide.docx"
SCHEMA_INVENTORY = ROOT / "tmp" / "schema-inventory.txt"
TODAY = date.today().strftime("%B %d, %Y")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "102A43"
INK = "243B53"
MUTED = "627D98"
WHITE = "FFFFFF"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E9F7EF"
PALE_AMBER = "FFF7E6"
PALE_RED = "FDECEC"
GREEN = "1B7F5A"
AMBER = "9A6700"
RED = "B42318"
GRID = "CBD5E1"
WIDTH_DXA = 9360


TABLE_PURPOSES = {
    "transactions": "Authoritative transaction events, source metadata, processing state, and supervised label lifecycle.",
    "bulk_upload_batches": "One audit record per uploaded file, including success and failure counts.",
    "uploaded_datasets": "Logical upload or database snapshot metadata used by comparison workflows.",
    "dataset_partitions": "Chronological percentage partitions and row counts for model-growth comparison.",
    "training_runs": "Compatibility training-run records used by earlier comparison/training flows.",
    "aml_training_runs": "Governed immutable training windows, export state, selected model, artifact, and failure metadata.",
    "aml_business_days": "Open/closed business dates used to prevent mutable historical exports.",
    "aml_transaction_features": "Persisted point-in-time feature vectors used for reproducible scoring and training.",
    "aml_feature_learning_status": "Eligibility, quarantine, review, and learning-state decisions per transaction.",
    "aml_customer_observed_profile": "Fast operational profile updated from observed customer activity.",
    "aml_customer_trusted_profile": "Profile derived only from transactions approved for trusted learning.",
    "aml_customer_recent_transactions": "Bounded recent-history support for velocity and behavior calculations.",
    "fraud_prediction_logs": "Immutable request/response and explanation audit for each scored transaction.",
    "fraud_alerts": "Suspicious-event alerts created from qualifying production decisions.",
    "case_records": "Investigation case header, workflow state, assignment, priority, and linked transaction/alert.",
    "case_notes": "Attributed analyst and system notes attached to a case.",
    "str_reports": "Generated suspicious transaction report metadata and XML content/state.",
    "app_users": "Application identities, BCrypt password hashes, role, and activation state.",
    "app_config": "Typed runtime configuration used by settings, tuning, research, and risk policy services.",
    "anomaly_configs": "Compatibility anomaly-policy records retained for earlier configuration flows.",
    "aml_model_registry": "Immutable model versions, artifact paths/checksums, metrics, lifecycle state, and feature contract.",
    "aml_active_models": "Atomic active/previous model pointers by model type and segment.",
    "aml_model_validations": "Candidate validation evidence and pass/fail governance result.",
    "aml_model_deployments": "Promotion and rollback event history for model versions.",
    "model_versions": "Compatibility registry for earlier partition-model lifecycle workflows.",
    "aml_shadow_predictions": "Silent challenger/layered scores that cannot directly create production outcomes.",
    "aml_shadow_scenario_labels": "Expected labels for synthetic or controlled shadow scenarios.",
    "aml_layered_validations": "Validation reports for the complete layered risk architecture.",
    "aml_layered_deployment_pointers": "Active layered architecture pointer, locked policy/models, and rollout percentage.",
    "aml_layered_deployment_events": "Audit history for layered promotion, expansion, and rollback.",
    "aml_growth_studies": "Unsupervised data-growth study status and serialized result.",
    "aml_growth_metrics": "Per-model/per-partition unsupervised quality, stability, rate, latency, and capacity metrics.",
    "aml_supervised_growth_studies": "Supervised chronological comparison status and cached result.",
    "aml_agreement_studies": "Unsupervised model agreement and overlap study cache.",
    "comparison_runs": "Earlier scenario-comparison run header retained in the backend.",
    "comparison_results": "Earlier model-comparison metrics retained for compatibility.",
    "comparison_scenarios": "Earlier reusable comparison scenario definitions.",
    "scenario_sets": "Earlier scenario-set metadata retained although its UI page is inactive.",
}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, size=None, bold=None, color=INK, italic=None, font="Calibri"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), font)
    rpr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code_style = doc.styles.add_style("CodePath", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = rgb(DARK_BLUE)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("FRAUD TRANSACTION DETECTOR  |  DEFENSE CODE & BUSINESS GUIDE")
    set_run(r, size=8, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item))


def add_steps(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "271")
    ppr.append(indent)
    level.append(ppr)
    abstract.append(level)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_ref = OxmlElement("w:numId")
        num_ref.set(qn("w:val"), str(num_id))
        num_pr.append(num_ref)
        set_run(p.add_run(item))


def add_explicit_steps(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run(p.add_run(f"{index}.  {item}"))


def add_code_paths(doc, paths):
    for path in paths:
        p = doc.add_paragraph(style="CodePath")
        set_run(p.add_run(path), size=8.5, color=DARK_BLUE, font="Consolas")


def add_callout(doc, label, text, kind="info"):
    palette = {
        "info": (PALE_BLUE, BLUE),
        "success": (PALE_GREEN, GREEN),
        "warning": (PALE_AMBER, AMBER),
        "risk": (PALE_RED, RED),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), bold=True, color=accent)
    set_run(p.add_run(text))
    table_geometry(table, [WIDTH_DXA])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=8.9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        shade(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            if row_index % 2:
                shade(cell, "F8FAFC")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run(run, size=font_size)
    table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def keep_table_block_together(table):
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def add_chunked_table(doc, label, headers, rows, widths, font_size=8.9, chunk_size=18):
    chunks = [rows[index:index + chunk_size] for index in range(0, len(rows), chunk_size)] or [[]]
    for chunk_index, chunk in enumerate(chunks):
        if chunk_index:
            add_heading(doc, f"{label} (continued)", 3)
        table = add_table(doc, headers, chunk, widths, font_size)
        keep_table_block_together(table)


def add_flow(doc, steps):
    widths = [WIDTH_DXA // len(steps)] * len(steps)
    widths[-1] += WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=len(steps))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, (label, detail) in enumerate(steps):
        cell = table.cell(0, index)
        shade(cell, PALE_BLUE if index % 2 == 0 else "F8FAFC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(label), size=8.8, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_run(p2.add_run(detail), size=8.1)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(74)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("PROJECT DEFENSE HANDBOOK"), size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run(p.add_run("Fraud Transaction Detector"), size=29, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run(p.add_run("Code Map, Business Logic, Database, ML, Security, Operations, and Defense Questions"), size=13, color=MUTED)

    add_table(doc, ["Guide", "Scope"], [
        ["Audience", "Student developer, supervisor, external examiner, technical reviewer"],
        ["Implementation snapshot", "Current repository and live SQL schema as reviewed on August 15, 2026"],
        ["Architecture", "Angular UI + Spring Boot modular monolith + Python FastAPI ML service + Microsoft SQL Server"],
        ["Purpose", "Explain where each feature is implemented, why it exists, and how to defend the design honestly"],
    ], [2100, 7260], 9.2)
    add_callout(doc, "Core claim", "The system identifies suspicious transactions and prioritizes investigation. It does not claim that an alert proves fraud; reviewed outcomes provide the labels needed for supervised learning.", "warning")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    set_run(p.add_run(f"Prepared {TODAY}"), size=9.5, color=MUTED)
    doc.add_page_break()


def parse_inventory():
    tables = defaultdict(list)
    configs = []
    config_mode = False
    for raw in SCHEMA_INVENTORY.read_text(encoding="utf-8", errors="replace").splitlines():
        if raw == "===CONFIG===":
            config_mode = True
            continue
        if not raw.strip():
            continue
        parts = raw.split("|")
        if config_mode:
            key, value_type, description = parts[0], parts[2], "|".join(parts[3:])
            configs.append((key, value_type, description.replace("�", "-")))
        else:
            table, ordinal, name, data_type, max_length, precision, scale, nullable, identity = parts
            if data_type in {"varchar", "nvarchar", "char", "nchar", "varbinary"}:
                length = "MAX" if max_length == "-1" else max_length
                type_text = f"{data_type}({length})"
            elif data_type in {"decimal", "numeric"}:
                type_text = f"{data_type}({precision},{scale})"
            else:
                type_text = data_type
            tables[table].append((int(ordinal), name, type_text, "Yes" if nullable == "1" else "No", "Yes" if identity == "1" else "No"))
    return dict(sorted(tables.items())), configs


def build():
    tables, configs = parse_inventory()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup(doc)
    cover(doc)

    add_heading(doc, "How to Use This Handbook")
    add_para(doc, "Use the first sections for your presentation and viva answers. Use the middle sections when an examiner asks where a rule is implemented. Use the appendices as a complete lookup for routes, APIs, tables, columns, and configuration keys.")
    add_callout(doc, "Accuracy of this guide", "Current source code and the live SQL schema are the authority. Older phase documents remain useful history, but some mention models and pages that were later replaced.", "info")
    add_heading(doc, "Contents", 2)
    add_bullets(doc, [
        "1. Five-minute defense narrative",
        "2. Architecture and repository map",
        "3. Authentication, authorization, and security",
        "4. Angular pages and user journeys",
        "5. Transaction ingestion, feature engineering, and prediction",
        "6. Layered risk scoring and decisions",
        "7. Unsupervised and supervised learning modes",
        "8. Training, comparison, artifacts, and governance",
        "9. Labels, cases, false positives, and STR",
        "10. Scalability, performance, reliability, and auditability",
        "11. Manual demonstration sequence",
        "12. Defense questions and honest limitations",
        "Appendix A. API index",
        "Appendix B. Complete database table and column reference",
        "Appendix C. Runtime configuration catalog",
        "Appendix D. Code navigation cheat sheet",
    ])

    add_heading(doc, "1. Five-Minute Defense Narrative")
    add_heading(doc, "Problem", 2)
    add_para(doc, "Banks process many transactions, confirmed fraud labels are rare and delayed, customer behavior differs, and a single model can create too many false positives. Investigators also need reasons, evidence, reproducibility, and a controlled path from model training to production.")
    add_heading(doc, "Solution", 2)
    add_para(doc, "The project combines point-in-time feature engineering, configurable customer and peer behavior, deterministic AML rules, and a selectable machine-learning ensemble. It supports unsupervised learning when labels are unavailable and supervised learning after reliable labels accumulate. Every suspicious decision can create a case for human review.")
    add_flow(doc, [
        ("Ingest", "File or API transaction"),
        ("Build", "Point-in-time features"),
        ("Score", "Behavior + peer + ML + rules"),
        ("Decide", "Weighted risk policy"),
        ("Investigate", "Case, notes, FP, STR"),
    ])
    add_heading(doc, "Why the design is defensible", 2)
    add_bullets(doc, [
        "Spring Boot owns security, SQL data, workflows, audit, and final business decisions; Python owns numerical training and inference.",
        "Training data is chronological and point-in-time, reducing future leakage.",
        "Unsupervised comparison uses label-free diagnostics; supervised comparison uses PR-AUC, precision, recall, F1, balanced accuracy, Brier score, and confusion matrices.",
        "Model output is only one risk layer. Customer behavior, peer behavior, and AML rules remain explicit and explainable.",
        "Training creates candidates; governance and version pointers separate a trained artifact from a production decision.",
        "MEDIUM creates a review case; HIGH can generate an STR workflow. Human review remains the final authority.",
    ])
    add_callout(doc, "One-sentence defense", "This is an auditable AML decision-support platform that can begin without labels, learn from reviewed outcomes, compare models as data grows, and preserve human control over suspicious-transaction decisions.", "success")

    add_heading(doc, "2. Architecture and Repository Map")
    add_table(doc, ["Component", "Responsibility", "Main location"], [
        ["Angular UI", "Forms, reports, comparison views, training operations, risk policy, transaction checks, and cases.", "fraud-transaction-ui/src/app"],
        ["Spring Boot", "Authentication, authorization, business workflows, SQL access, features, risk aggregation, alerts, cases, model registry, and audit.", "fraud-transaction-detector/src/main/java"],
        ["FastAPI ML service", "Model training, artifact loading, prediction, growth analysis, agreement, supervised evaluation, and policy backtesting.", "fraud-ml-service/app"],
        ["SQL Server", "Authoritative events, labels, features, profiles, configs, model governance, predictions, alerts, and cases.", "fraud-transaction-detector/src/main/resources/db"],
    ], [1500, 4550, 3310])
    add_heading(doc, "Runtime communication", 2)
    add_flow(doc, [
        ("Browser :4200", "JWT + JSON/HTTP"),
        ("Spring :8080", "Business authority"),
        ("FastAPI :8000", "ML computation"),
        ("SQL Server", "System of record"),
    ])
    add_para(doc, "The browser does not connect directly to SQL Server or FastAPI. This keeps credentials, authorization, audit, and final decision rules in Spring Boot.")
    add_heading(doc, "Modular monolith decision", 2)
    add_para(doc, "The Java services were merged into one Spring Boot application. The code remains separated by packages such as auth, uploads, transactions, aml, comparison, cases, and config. This reduces deployment complexity for the project while preserving module boundaries that can later be separated if scale or organizational ownership requires it.")

    add_heading(doc, "3. Authentication, Authorization, and Security")
    add_heading(doc, "Authentication flow", 2)
    add_explicit_steps(doc, [
        "A user registers or logs in through Angular.",
        "Spring validates the username and BCrypt password hash stored in app_users.",
        "Spring signs an HS256 JWT containing subject, role, full name, issued time, and expiry.",
        "Angular sends the token as Authorization: Bearer <token>.",
        "JwtAuthenticationFilter verifies the signature and expiry, reloads the active user from SQL, and creates the Spring Security principal.",
        "SecurityConfig authorizes the route by HTTP method and role. Unauthorized and forbidden responses are structured JSON.",
    ])
    add_table(doc, ["Control", "Implementation", "Defense point"], [
        ["Password storage", "BCrypt with work factor 12", "Passwords are not stored in clear text or reversibly encrypted."],
        ["Session model", "Stateless JWT", "No server HTTP session is required; every protected request is authenticated."],
        ["Self-registration", "Always creates REVIEWER", "A public user cannot grant themselves ADMIN privileges."],
        ["API authorization", "Spring Security route and method rules", "Angular guards improve UX; backend rules provide real enforcement."],
        ["Account activation", "Active state checked from SQL", "A disabled user cannot continue using a previously issued identity unchecked."],
        ["CORS", "Configured allowed origins and OPTIONS support", "Browser cross-origin access is explicit rather than accidental."],
    ], [1900, 3200, 4260])
    add_heading(doc, "Roles", 2)
    add_table(doc, ["Role", "Typical access"], [
        ["REVIEWER", "Transaction checks, comparisons, case review, notes, false-positive decisions, STR generation, and manual cases."],
        ["ADMIN / AML_ADMIN", "Reviewer access plus uploads, training, ML configuration, settings, risk policy, validation, promotion, and rollback."],
    ], [2200, 7160])
    add_heading(doc, "Key code", 2)
    add_code_paths(doc, [
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/auth/security/SecurityConfig.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/auth/security/JwtAuthenticationFilter.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/auth/service/JwtService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/auth/service/AuthService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/auth/service/PasswordHashService.java",
        "fraud-transaction-ui/src/app/core/auth.guard.ts",
    ])
    add_callout(doc, "Production hardening", "Local database credentials and any development JWT fallback must be replaced by environment variables or a secret manager. Never present a development default as production security.", "risk")

    add_heading(doc, "4. Angular Pages and User Journeys")
    add_table(doc, ["Route", "Purpose", "Access", "Main component"], [
        ["/login", "Simple sign-in page.", "Public", "login-page.component"],
        ["/register", "Create a REVIEWER account.", "Public", "register-page.component"],
        ["/uploads", "Upload CSV/Excel transactions; required label guidance changes with system mode.", "Authenticated; upload API admin", "upload-page.component"],
        ["/datasets", "Unsupervised comparison, agreement, and 10/25/50/100 growth analysis.", "Authenticated", "datasets-page.component"],
        ["/supervised-comparison", "Supervised model, fusion, confusion-matrix, and full risk-policy performance.", "Authenticated", "supervised-comparison-page.component"],
        ["/training-operations", "Close dates, create immutable runs, export data, choose configured models, and train candidates.", "Admin", "training-operations-page.component"],
        ["/model-tuning", "Edit controls for the active supervised or unsupervised model family.", "Admin", "model-tuning-page.component"],
        ["/config", "Edit four-layer risk weights, thresholds, and production ML composition.", "Admin", "anomaly-config-page.component"],
        ["/settings", "Select global learning mode and configure cold-start behavior.", "Admin", "cold-start-config-page.component"],
        ["/transaction-check", "Create/check a single API transaction and inspect explanation, models, and features.", "Authenticated", "transaction-check-page.component"],
        ["/cases", "Paginated investigation queue, evidence, notes, false positive, and STR actions.", "Authenticated", "cases-page.component"],
        ["/cases/manual", "Search stored transactions and create a manual investigation case.", "Authenticated", "manual-case-page.component"],
    ], [1450, 3560, 1650, 2700], 8.3)
    add_callout(doc, "Inactive UI scope", "Scenario Library, Comparison Runs, and Simulator pages were removed from active Angular routes. Their older backend tables/controllers may remain for compatibility and should not be presented as current user journeys.", "info")

    add_heading(doc, "5. Transaction Ingestion, Features, and Prediction")
    add_heading(doc, "Bulk upload", 2)
    add_steps(doc, [
        "Angular sends a CSV or Excel file to Spring.",
        "ExcelBulkUploadService normalizes headers, parses rows, validates required fields, records the batch, and writes transactions.",
        "Uploaded historical rows preserve FraudLabel when supplied. In supervised mode, null labels are excluded from supervised training.",
        "Upload and training remain separate so data can arrive continuously and an administrator can train later.",
    ])
    add_heading(doc, "Single transaction", 2)
    add_steps(doc, [
        "TransactionCreateService rejects a duplicate transaction ID and stores the API event.",
        "FeatureEngineeringService loads only information available before the transaction timestamp.",
        "The versioned feature vector is persisted in aml_transaction_features.",
        "Observed/trusted profile services and learning eligibility are updated according to policy.",
        "Cold-start settings may skip ML when trustworthy history is insufficient.",
        "AmlPredictionOrchestrator obtains ML evidence and invokes layered scoring.",
        "Spring persists fraud_prediction_logs, creates an alert/case when required, and returns risk, action, reasons, model diagnostics, and feature summary.",
    ])
    add_heading(doc, "Feature families", 2)
    add_table(doc, ["Family", "Examples", "Reason"], [
        ["Current event", "amount, balance, channel, location, type, hour, weekday", "Describes the transaction itself."],
        ["Customer history", "last-5/30 averages, median, std, max/min, debit/credit/cash ratios", "Measures deviation from personal behavior."],
        ["Velocity", "counts and amount sums over 10m, 1h, 24h, 7d, 30d", "Detects bursts, rapid movement, and layering."],
        ["Novelty", "new beneficiary, location, channel, device", "Detects first-seen contextual changes."],
        ["Peer", "peer average, median, std, z-score, frequency percentile", "Compares similar customers, including occupation/age-derived groups when available."],
        ["Customer profile", "risk rating, expected turnover, profile confidence", "Controls how strongly behavior evidence should be trusted."],
        ["Terminal/context", "bounded terminal/location risk representation", "Adds operational risk without unbounded one-hot memory growth."],
    ], [1700, 4050, 3610], 8.5)
    add_callout(doc, "Leakage protection", "Historical aggregates must use transaction_date < current transaction date. Training and evaluation use chronological boundaries rather than a random future-aware feature calculation.", "success")
    add_heading(doc, "Key code", 2)
    add_code_paths(doc, [
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/uploads/service/ExcelBulkUploadService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/transactions/service/TransactionCreateService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/feature/application/FeatureEngineeringService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/feature/infrastructure/FeaturePersistenceService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/profile/application/CustomerProfileService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/learning/application/LearningEligibilityService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/prediction/AmlPredictionOrchestrator.java",
    ])

    add_heading(doc, "6. Layered Risk Scoring and Decisions")
    add_table(doc, ["Layer", "Input", "Output meaning"], [
        ["Customer behavior", "Personal amount, frequency, novelty, time gap, unusual hour, and profile confidence.", "How abnormal this event is for the customer."],
        ["Peer behavior", "Peer amount deviation, frequency percentile, and expected-turnover context.", "How abnormal it is for comparable customers."],
        ["ML ensemble", "Enabled production model scores/decisions and normalized internal weights.", "Combined nonlinear evidence from the selected learning mode."],
        ["AML rules", "Structuring, velocity, amount, turnover, sanctions/location, and other deterministic controls.", "Explicit policy evidence; severe rules can override weighted scoring."],
    ], [1900, 4550, 2910])
    add_heading(doc, "Calculation", 2)
    add_para(doc, "Each layer produces a normalized score from 0 to 1. The active policy assigns four weights that must total 1.0. The aggregation engine calculates:")
    add_callout(doc, "Formula", "Final risk = customer score × customer weight + peer score × peer weight + ML score × ML weight + rule score × rule weight. A configured hard rule can raise the result to at least the HIGH threshold.", "info")
    add_para(doc, "Inside the ML layer, enabled model allocations are normalized to 100%. An anomalous model decision can consume its allocated share, which makes unanimous selected models produce an ML component score of 1.0. The top-level ML weight still limits its contribution to the final score.")
    add_table(doc, ["Band", "Meaning", "Typical action"], [
        ["NORMAL", "Below LOW threshold", "ALLOW"],
        ["LOW", "At or above LOW but below MEDIUM", "ALLOW_AND_LOG"],
        ["MEDIUM", "At or above MEDIUM but below HIGH", "ALLOW_AND_ALERT; create review case"],
        ["HIGH", "At or above HIGH or hard-rule override", "HOLD_FOR_REVIEW; STR workflow may be generated"],
    ], [1500, 3700, 4160])
    add_heading(doc, "Why three anomalous ML models may not automatically mean HIGH", 2)
    add_para(doc, "Unanimous ML evidence makes the ML component high, not necessarily the final score high. The final score also depends on the ML top-level weight, the other three layers, and the MEDIUM/HIGH thresholds. This is intentional: model agreement is strong evidence, but the business policy controls its operational effect.")
    add_heading(doc, "Key code", 2)
    add_code_paths(doc, [
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/prediction/LayeredShadowScoringService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/risk/application/WeightedRiskAggregationEngine.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/behaviour/customer/CustomerBehaviourScorer.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/behaviour/peer/PeerBehaviourScorer.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/rules/engine/DeterministicAmlRuleEngine.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/risk/infrastructure/AppConfigRiskPolicyRepository.java",
    ])

    add_heading(doc, "7. Unsupervised and Supervised Learning Modes")
    add_para(doc, "The global system.learning_mode setting controls upload guidance, model-tuning controls, training selections, production ML catalog, and the appropriate comparison page.")
    add_table(doc, ["Mode", "When to use", "Current models", "Valid evaluation claim"], [
        ["UNSUPERVISED", "No reliable FraudLabel exists.", "Isolation Forest, Autoencoder, Behavioral Cluster Outlier", "Label-free separation, anomaly-rate control, agreement, stability, growth, latency; not fraud accuracy."],
        ["SUPERVISED", "Reliable 0/1 labels exist from source data or reviewed outcomes.", "XGBoost, Class-Balanced Random Forest, Extra Trees; optional Temporal Stacked Ensemble", "PR-AUC, precision, recall, F1, balanced accuracy, Brier score, confusion matrices, and risk-policy backtest."],
    ], [1450, 2200, 3000, 2710], 8.2)
    add_heading(doc, "Why two modes are necessary", 2)
    add_bullets(doc, [
        "A new client may have transactions but no confirmed fraud labels, so unsupervised learning provides an honest starting point.",
        "After investigators mark false positives or suspicious/STR outcomes, reviewed evidence can create supervised labels.",
        "Rows with FraudLabel = NULL are excluded from supervised training; they are not silently assumed legitimate.",
        "A new API transaction that creates no case may become label 0 under the accepted business rule; uploaded unlabeled history remains NULL unless the file supplies a label.",
        "The modes share feature governance and risk layers, but their ML training and evaluation logic differ.",
    ])
    add_callout(doc, "Avoid confirmation bias", "Do not convert every model alert into label 1. The label must come from reviewed evidence or an accepted business outcome, otherwise the supervised model only learns the previous model's mistakes.", "risk")

    doc.add_page_break()
    add_heading(doc, "8. Training, Comparison, Artifacts, and Governance")
    add_heading(doc, "Training Operations", 2)
    add_steps(doc, [
        "Close business dates so the historical window cannot change during export.",
        "Create a training run with feature version, date range, cutoff, and selected models from the active learning mode.",
        "Export eligible rows in keyset-paginated chunks to immutable Parquet parts with checksums.",
        "Train each selected model independently on the same snapshot.",
        "Write immutable artifact bundles, metrics, thresholds, feature columns/schema, and checksums.",
        "Register candidate versions. A candidate is not production simply because training completed.",
    ])
    add_heading(doc, "Supervised split", 2)
    add_para(doc, "The supervised pipeline uses chronological data: approximately 60% fit, 10% tuning, 10% threshold calibration, and 20% untouched evaluation. This keeps threshold selection and final reporting separate from model fitting.")
    add_heading(doc, "Unsupervised growth comparison", 2)
    add_para(doc, "The oldest 10%, 25%, 50%, and 100% of history are compared. Partitions under the minimum row requirement are skipped. The report shows how detector behavior changes as data grows, including separation/EM-AUC-style diagnostics, anomaly rate, score distribution, training time, throughput, agreement, and stability.")
    add_heading(doc, "Supervised growth comparison", 2)
    add_para(doc, "The same chronological growth idea is applied to labeled rows. Individual models, fusion strategies, confusion matrices, and the complete four-layer risk policy are evaluated on held-out rows. Results are cached in aml_supervised_growth_studies so page reloads do not recompute a completed study.")
    add_heading(doc, "Model tuning versus risk policy", 2)
    add_table(doc, ["Control", "What it changes", "What it does not change"], [
        ["Model Tuning", "Training hyperparameters, evaluation protocol, enabled candidate generation.", "It does not automatically select a model for production."],
        ["Risk Policy", "Four layer weights, internal ML allocations, LOW/MEDIUM/HIGH thresholds, rule thresholds.", "It does not retrain model parameters."],
        ["Training Operations", "Which configured models are trained against a selected immutable window.", "It does not activate a candidate."],
        ["Model registry/deployment", "Which validated immutable model version is active and rollback history.", "It does not rewrite historical metrics or artifacts."],
    ], [1800, 3900, 3660])
    add_heading(doc, "Large-data protections", 2)
    add_bullets(doc, [
        "SQL keyset pagination and Parquet record batches avoid loading all eligible rows in Java memory.",
        "High-cardinality terminal/location data is compacted into bounded deterministic buckets instead of producing more than ten thousand object columns.",
        "Tuning uses a bounded representative sample and bounded tree counts; the final selected fit can use the approved full snapshot.",
        "Comparison result caching prevents repeated expensive recomputation after page refresh.",
        "Training and comparisons are asynchronous; status is persisted and polled by the UI.",
    ])
    add_heading(doc, "Key code", 2)
    add_code_paths(doc, [
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/training/application/TrainingPipelineJobLauncher.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/training/application/TrainingDatasetExportService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/training/application/ProductionCandidateTrainingService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/training/application/AmlModelRegistryService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/research/application/GrowthAnalysisService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/aml/research/application/SupervisedGrowthStudyService.java",
        "fraud-ml-service/app/training.py",
        "fraud-ml-service/app/supervised_training.py",
        "fraud-ml-service/app/research/growth_analysis.py",
        "fraud-ml-service/app/research/supervised_growth_analysis.py",
        "fraud-ml-service/app/research/risk_policy_backtest.py",
        "fraud-ml-service/app/feature_compaction.py",
    ])

    add_heading(doc, "9. Labels, Cases, False Positives, and STR")
    add_heading(doc, "Label lifecycle", 2)
    add_table(doc, ["Situation", "FraudLabel", "Training effect"], [
        ["Historical upload without label", "NULL", "Excluded from supervised training; still usable for unsupervised learning."],
        ["Historical upload with source label", "0 or 1", "Eligible if it passes all other learning controls."],
        ["New API transaction with no case", "0 under the accepted business rule", "Can become legitimate supervised evidence after persistence and eligibility checks."],
        ["Case open/unreviewed", "Pending or unchanged", "Do not treat a generated alert as confirmed fraud."],
        ["Marked false positive", "0", "Adds reviewed legitimate evidence and updates linked alert/case state."],
        ["Suspicious outcome / STR generated", "1 under the accepted project rule", "Adds positive evidence; production institutions should align this with formal investigation policy."],
    ], [2450, 2550, 4360])
    add_heading(doc, "Automatic case flow", 2)
    add_steps(doc, [
        "A MEDIUM or HIGH final production risk creates or reuses a case linked to the transaction and alert.",
        "The case stores status, priority, assignment, creation source, and the latest decision evidence.",
        "The analyst reads user-friendly reasons, component scores, model diagnostics, feature summary, and policy/model versions.",
        "The analyst adds notes and either marks false positive or generates STR XML; mutually invalid transitions are blocked.",
    ])
    add_heading(doc, "Case performance controls", 2)
    add_bullets(doc, [
        "Paginated queries bound the number of cases returned.",
        "Search and status filters are applied by the API instead of requiring a manual refresh-only client workaround.",
        "Summary/evidence loading avoids per-row N+1 queries in the queue.",
        "Duplicate manual or automatic cases for the same event are handled idempotently where applicable.",
    ])
    add_heading(doc, "Key code", 2)
    add_code_paths(doc, [
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/cases/service/CaseManagementService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/cases/service/StrXmlService.java",
        "fraud-transaction-detector/src/main/java/com/ftd/fraud_transaction_detector/transactions/service/TransactionFraudLabelService.java",
        "fraud-transaction-ui/src/app/pages/cases-page.component.ts",
        "fraud-transaction-ui/src/app/pages/manual-case-page.component.ts",
    ])

    add_heading(doc, "10. Scalability, Performance, Reliability, and Auditability")
    add_table(doc, ["Problem", "Implemented response", "Remaining production work"], [
        ["Millions of rows", "Chunked SQL export, Parquet, bounded feature compaction, asynchronous jobs, cached comparisons.", "Distributed object storage/compute, workload scheduling, and production capacity testing."],
        ["Slow single transaction", "Persisted profiles/recent history, bounded lookups, in-memory Python artifacts, and no retraining on request.", "Profile cache, connection-pool tuning, p95/p99 tracing, model serving replicas."],
        ["ML unavailable", "Orchestrator/logging surfaces failure and preserves business control; current path has a fail-safe response where configured.", "Circuit breaker, retries with budgets, queueing, and explicit operational SLOs."],
        ["Schema drift", "Feature version, model feature schema, stored feature JSON, artifact metadata, and registry validation.", "Automated migration pipeline and compatibility contracts in CI/CD."],
        ["Model risk", "Candidate registry, validation, exact-version pointers, promotion/rollback, shadow studies, immutable metrics.", "Independent model-risk approval and periodic monitoring governance."],
        ["Explainability", "Reason codes, component scores, model payloads, feature summary, policy and model version in logs/cases.", "SHAP or model-specific local explanations for supervised decisions."],
        ["Audit", "Transactions, features, prediction logs, labels, cases, notes, STR, validations, deployments, and configs are persisted.", "Retention, encryption, tamper-evidence, SIEM integration, and access reviews."],
    ], [1800, 4200, 3360], 8.2)
    add_callout(doc, "Kafka position", "Kafka is not required to prove the current MVP. At bank scale it can decouple ingestion, feature updates, scoring, alerts, and retraining triggers; support replay; and absorb bursts. SQL remains the system of record and consumers must be idempotent.", "info")

    add_heading(doc, "11. Manual Demonstration Sequence")
    add_steps(doc, [
        "Open /login and authenticate. Explain JWT, BCrypt, roles, and backend enforcement.",
        "Open /settings. Choose UNSUPERVISED or SUPERVISED and explain why the whole application follows this mode.",
        "Open /uploads. Upload a matching dataset and explain label rules.",
        "Open /training-operations. Close required dates, create a run, export the immutable snapshot, choose models, and train candidates.",
        "Open the matching comparison page. Show 10%, 25%, 50%, and 100% growth; explain only metrics valid for that mode.",
        "Open /model-tuning. Explain that these controls affect the next training/comparison, not current production artifacts.",
        "Open /config. Explain four-layer weights, internal ML allocation, thresholds, and why totals must be exactly 100%. Save creates a versioned policy.",
        "Open /transaction-check. Submit a realistic event and explain the final score, reasons, component scores, model diagnostics, feature contract, and latency.",
        "Open /cases. Show pagination, evidence, notes, false-positive action, and STR XML workflow.",
        "Finish with model registry, audit tables, limitations, and how reviewed outcomes enable supervised learning.",
    ])
    add_callout(doc, "Demo advice", "Do not start a long comparison during the defense. Run it beforehand and demonstrate the cached completed result. Keep one small dataset available for a short live workflow.", "warning")

    add_heading(doc, "12. Defense Questions and Honest Answers")
    add_heading(doc, "Common examiner questions", 2)
    qa = [
        ("Why not call every anomaly fraud?", "An anomaly is only unusual. Fraud requires reliable labels or human investigation. The platform creates suspicious cases rather than claiming proof."),
        ("Why both supervised and unsupervised learning?", "Some clients begin without labels. Unsupervised models provide a safe starting point; reviewed outcomes later support supervised classifiers and valid precision/recall metrics."),
        ("Why PR-AUC instead of only accuracy?", "Fraud is imbalanced. A model can achieve high accuracy by predicting legitimate almost always. PR-AUC focuses on positive-class ranking quality."),
        ("Why chronological splits?", "Random splits can leak future behavioral patterns. Chronology better approximates training on the past and evaluating on later transactions."),
        ("Why can the ensemble be worse than one model?", "Ensembling only helps when models contribute complementary information and weights/thresholds are validated. Correlated errors or poor calibration can reduce performance."),
        ("Does a 90% target manipulate results?", "A target is acceptable only as a business operating point chosen on calibration data and honestly reported with its recall and case volume. It must not be tuned on the final test set."),
        ("Why keep rules if ML exists?", "Rules encode known AML obligations and provide deterministic, explainable overrides. ML discovers nonlinear patterns but should not replace explicit policy."),
        ("How does data growth analysis help?", "It shows whether quality, rate, agreement, latency, and stability improve or deteriorate as the chronological history increases from 10% to 100%."),
        ("How is reproducibility achieved?", "Training windows, cutoff, feature version, Parquet checksums, artifact checksums, random seed, metrics, policies, and model versions are persisted."),
        ("Why is Java the business authority?", "Security, SQL consistency, labels, cases, thresholds, and regulatory workflow belong in a transactional application layer. Python remains focused on ML computation."),
        ("Can this handle a bank's full volume today?", "The design contains bank-scale patterns, but production capacity is not proven until load tests, distributed deployment, observability, resilience, security hardening, and regulatory validation are completed."),
        ("What is the biggest academic contribution?", "A dual-learning, growth-aware, layered risk architecture that compares models honestly, preserves explainable AML logic, and connects investigation outcomes back to future supervised learning."),
    ]
    for question, answer in qa:
        add_heading(doc, question, 3)
        add_para(doc, answer)
    add_heading(doc, "Known limitations", 2)
    add_bullets(doc, [
        "Synthetic or limited project datasets cannot prove performance on a real bank population.",
        "STR XML is a draft artifact and not a certified regulator submission integration.",
        "Production secrets, encryption, retention, privacy, and infrastructure controls require institutional deployment work.",
        "A high score is decision support; model drift, reviewer consistency, and label quality must be monitored continuously.",
        "Some compatibility controllers/tables remain although their Angular pages are inactive; they should be cleaned after migration assurance.",
    ])

    doc.add_page_break()
    add_heading(doc, "Appendix A. API Index")
    api_rows = [
        ["Auth", "POST /api/auth/register", "Create REVIEWER account", "Public"],
        ["Auth", "POST /api/auth/login", "Issue JWT", "Public"],
        ["Auth", "GET /api/auth/me; GET /validate", "Current identity/token check", "Authenticated"],
        ["Upload", "POST /api/v1/uploads/transactions", "Import CSV/Excel", "Admin"],
        ["Upload", "GET /api/v1/uploads/batches[/latest/{batchNo}]", "Upload history", "Authenticated"],
        ["Transactions", "POST /api/v1/transactions", "Persist, feature, score, alert/case", "Authenticated"],
        ["Transactions", "GET /api/v1/transactions[/{transactionId}]", "Search/read transactions", "Authenticated"],
        ["Cases", "GET /api/cases", "Paginated/filter case queue", "Authenticated"],
        ["Cases", "POST /api/cases; GET /api/cases/{id}", "Manual create/read detail", "Authenticated"],
        ["Cases", "PUT /api/cases/{id}/status; POST /notes", "Workflow and notes", "Authenticated"],
        ["Cases", "POST /api/cases/{id}/false-positive", "Resolve as false positive", "Authenticated"],
        ["Cases", "POST /api/cases/{id}/str-xml", "Generate draft STR XML", "Authenticated"],
        ["Training", "POST /api/v1/aml/training-runs", "Create governed run", "Admin"],
        ["Training", "POST /api/v1/aml/training-runs/pipeline/start", "Create/export/train pipeline", "Admin"],
        ["Training", "POST .../{runId}/dataset; POST .../training", "Export/train selected run", "Admin"],
        ["Training", "GET /api/v1/aml/training-runs", "Run status/history", "Authenticated"],
        ["Business day", "POST /api/v1/aml/business-days/close[/range]", "Seal dates", "Admin"],
        ["Registry", "GET /api/v1/aml/models; GET .../active", "Registry and active pointers", "Authenticated"],
        ["Registry", "POST .../{version}/validate|promote|rollback", "Governed lifecycle", "Admin"],
        ["Research", "POST /api/v1/aml/growth-analysis/training-runs/{id}", "Unsupervised growth", "Admin"],
        ["Research", "POST /api/v1/aml/growth-analysis/layer-ablation", "Layer effect replay", "Admin"],
        ["Research", "GET/POST /api/v1/aml/growth-studies", "Cached unsupervised studies", "Auth/Admin"],
        ["Research", "GET/POST /api/v1/aml/supervised-growth-studies", "Cached supervised studies", "Auth/Admin"],
        ["Research", "GET/POST /api/v1/aml/agreement-studies", "Agreement studies", "Auth/Admin"],
        ["Layered", "POST /api/v1/aml/layered-shadow/validate", "Validate full architecture", "Admin"],
        ["Layered", "POST /api/v1/aml/layered-deployments/promote|rollback", "Layered routing control", "Admin"],
        ["Config", "/api/v1/anomaly-model-comparisons/risk-policy", "Risk weights/thresholds", "Auth/Admin writes"],
        ["Config", "/api/v1/anomaly-model-comparisons/model-tuning", "Tuning controls", "Auth/Admin writes"],
        ["Config", "/api/v1/anomaly-model-comparisons/model-catalog", "Mode-specific model catalog", "Authenticated"],
        ["Python", "POST /api/v2/fraud/predict", "Production ML inference", "Internal Spring call"],
        ["Python", "POST /api/v1/aml/research/growth-analysis", "Unsupervised research compute", "Internal Spring call"],
        ["Python", "POST /api/v1/aml/model-agreement", "Agreement compute", "Internal Spring call"],
    ]
    add_table(doc, ["Area", "Endpoint", "Purpose", "Access"], api_rows, [1350, 3800, 2910, 1300], 7.8)

    doc.add_page_break()
    add_heading(doc, "Appendix B. Complete Database Table and Column Reference")
    add_para(doc, f"Live inventory: {len(tables)} tables. Data types, nullability, and identity flags were read from the active fraud-transaction-detector SQL Server schema. Table descriptions summarize current application use.")
    grouped = {
        "Identity, configuration, and ingestion": ["app_users", "app_config", "bulk_upload_batches", "uploaded_datasets", "transactions"],
        "Features, profiles, and learning": ["aml_transaction_features", "aml_feature_learning_status", "aml_customer_observed_profile", "aml_customer_trusted_profile", "aml_customer_recent_transactions", "aml_business_days"],
        "Predictions, alerts, cases, and reports": ["fraud_prediction_logs", "fraud_alerts", "case_records", "case_notes", "str_reports", "aml_shadow_predictions", "aml_shadow_scenario_labels"],
        "Training, registry, validation, and deployment": ["aml_training_runs", "training_runs", "aml_model_registry", "aml_active_models", "aml_model_validations", "aml_model_deployments", "aml_layered_validations", "aml_layered_deployment_pointers", "aml_layered_deployment_events", "model_versions"],
        "Research and comparison": ["aml_growth_studies", "aml_growth_metrics", "aml_supervised_growth_studies", "aml_agreement_studies", "dataset_partitions", "comparison_runs", "comparison_results", "comparison_scenarios", "scenario_sets", "anomaly_configs"],
    }
    covered = set()
    for group, names in grouped.items():
        add_heading(doc, group, 2)
        for name in names:
            if name not in tables:
                continue
            covered.add(name)
            add_heading(doc, name, 3)
            add_para(doc, TABLE_PURPOSES.get(name, "Supporting application table."))
            rows = [[str(ordinal), column, data_type, nullable, identity] for ordinal, column, data_type, nullable, identity in tables[name]]
            add_chunked_table(doc, name, ["#", "Column", "SQL type", "Nullable", "Identity"], rows, [600, 3450, 2500, 1400, 1410], 7.8)
    remaining = [name for name in tables if name not in covered]
    if remaining:
        add_heading(doc, "Other compatibility/support tables", 2)
        for name in remaining:
            add_heading(doc, name, 3)
            add_para(doc, TABLE_PURPOSES.get(name, "Supporting or compatibility table in the current live schema."))
            rows = [[str(ordinal), column, data_type, nullable, identity] for ordinal, column, data_type, nullable, identity in tables[name]]
            add_chunked_table(doc, name, ["#", "Column", "SQL type", "Nullable", "Identity"], rows, [600, 3450, 2500, 1400, 1410], 7.8)

    doc.add_page_break()
    add_heading(doc, "Appendix C. Runtime Configuration Catalog")
    add_para(doc, "The app_config table is the runtime configuration authority. Values are intentionally omitted here because they are mutable; the key, type, and purpose are the stable defense reference.")
    config_groups = defaultdict(list)
    for key, value_type, description in configs:
        if key.startswith("aml.risk."):
            group = "Risk policy and rules"
        elif key.startswith("ml.supervised."):
            group = "Supervised learning"
        elif key.startswith("aml.research."):
            group = "Research comparison"
        elif key.startswith("aml."):
            group = "AML training, export, registry, and validation"
        elif key.startswith("ml."):
            group = "ML and cold-start compatibility"
        elif key.startswith("system."):
            group = "System mode"
        else:
            group = "Other"
        config_groups[group].append((key, value_type, description))
    for group in sorted(config_groups):
        add_heading(doc, group, 2)
        add_chunked_table(doc, group, ["Key", "Type", "Purpose"], config_groups[group], [3500, 1300, 4560], 7.7, 16)

    doc.add_page_break()
    add_heading(doc, "Appendix D. Code Navigation Cheat Sheet")
    add_table(doc, ["Question", "Start here"], [
        ["Where is login/register logic?", "auth/web/AuthController.java; auth/service/AuthService.java"],
        ["Where is JWT verified?", "auth/security/JwtAuthenticationFilter.java; auth/service/JwtService.java"],
        ["Where are roles enforced?", "auth/security/SecurityConfig.java and Angular core/auth.guard.ts"],
        ["Where is file upload parsed?", "uploads/service/ExcelBulkUploadService.java"],
        ["Where is a transaction scored?", "transactions/service/TransactionCreateService.java; aml/prediction/AmlPredictionOrchestrator.java"],
        ["Where are features calculated?", "aml/feature/application/FeatureEngineeringService.java"],
        ["Where are features persisted?", "aml/feature/infrastructure/FeaturePersistenceService.java"],
        ["Where is customer behavior scored?", "aml/behaviour/customer/CustomerBehaviourScorer.java"],
        ["Where is peer behavior scored?", "aml/behaviour/peer/PeerBehaviourScorer.java"],
        ["Where are AML rules evaluated?", "aml/rules/engine/DeterministicAmlRuleEngine.java"],
        ["Where is final score calculated?", "aml/risk/application/WeightedRiskAggregationEngine.java"],
        ["Where is risk policy loaded?", "aml/risk/infrastructure/AppConfigRiskPolicyRepository.java"],
        ["Where are labels changed?", "transactions/service/TransactionFraudLabelService.java"],
        ["Where are cases and STR handled?", "cases/service/CaseManagementService.java; cases/service/StrXmlService.java"],
        ["Where is the training pipeline?", "aml/training/application/TrainingPipelineJobLauncher.java"],
        ["Where is Parquet exported?", "aml/training/application/TrainingDatasetExportService.java"],
        ["Where are candidate models mapped?", "aml/training/application/ProductionCandidateTrainingService.java"],
        ["Where is model lifecycle handled?", "aml/training/application/AmlModelRegistryService.java; aml/deployment/application/ModelDeploymentService.java"],
        ["Where is unsupervised comparison?", "aml/research/application/GrowthAnalysisService.java; fraud-ml-service/app/research/growth_analysis.py"],
        ["Where is supervised comparison?", "aml/research/application/SupervisedGrowthStudyService.java; fraud-ml-service/app/research/supervised_growth_analysis.py"],
        ["Where is full policy backtested?", "fraud-ml-service/app/research/risk_policy_backtest.py"],
        ["Where are models trained in Python?", "fraud-ml-service/app/training.py; app/supervised_training.py"],
        ["Where are Python endpoints?", "fraud-ml-service/app/main.py"],
        ["Where are Angular routes?", "fraud-transaction-ui/src/app/app.routes.ts"],
        ["Where are Angular API calls/types?", "fraud-transaction-ui/src/app/core"],
    ], [3350, 6010], 8.1)
    add_callout(doc, "Final defense rule", "If you are unsure, state what the current code proves, distinguish prototype behavior from production assurance, and avoid promising accuracy or scale that has not been measured on real bank data.", "success")

    doc.core_properties.title = "Fraud Transaction Detector - Defense Code and Business Guide"
    doc.core_properties.subject = "Authentication, authorization, business logic, ML, database, APIs, operations, and defense preparation"
    doc.core_properties.author = "Fraud Transaction Detector Project"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
