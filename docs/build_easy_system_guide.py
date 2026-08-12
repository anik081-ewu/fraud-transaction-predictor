from pathlib import Path

from docx import Document

from build_project_documents import (
    add_bullets,
    add_callout,
    add_code,
    add_cover,
    add_flow,
    add_heading,
    add_para,
    add_steps,
    add_table,
    add_title,
    setup_document,
)


OUTPUT = Path(__file__).resolve().parent / "deliverables" / "AML_System_Explained_Simply.docx"


def build_guide() -> Path:
    document = Document()
    setup_document(
        document,
        guide=True,
        running_title="AML System Explained Simply",
    )
    add_cover(
        document,
        "Plain-Language Learning Guide",
        "How the AML Anomaly Detection System Works",
        "A simple explanation of the transaction flow, models, training, scalability, safety controls, and case management",
        "Students, professors, project evaluators, AML users, and new technical team members",
        "1.0",
    )

    add_title(document, "Start Here: The Whole System in One Minute")
    add_para(
        document,
        "Think of the system as a careful bank-security assistant. It watches each transaction, compares it with the customer's normal behaviour and similar customers, checks known AML rules, asks two scalable machine-learning models for additional evidence, and combines everything into one risk score. If the evidence is strong enough, it creates a case for a human reviewer.",
    )
    add_callout(
        document,
        "Most important idea",
        "The system does not say that fraud is proven. It says that a transaction is suspicious enough to investigate.",
        "warning",
    )
    add_flow(
        document,
        [
            ("Transaction", "Save the bank event"),
            ("Understand", "Build behaviour features"),
            ("Score", "Rules + behaviour + ML"),
            ("Decide", "Calculate final risk"),
            ("Review", "Create a human case"),
        ],
    )
    add_heading(document, "What the project is trying to solve", 2)
    add_bullets(
        document,
        [
            "Banks may receive millions of transactions, so the solution cannot load everything into memory for every decision.",
            "Confirmed fraud labels are limited and delayed, so the system must begin with anomaly evidence instead of pretending to know perfect accuracy.",
            "Every customer behaves differently, so one fixed amount threshold is not enough.",
            "A suspicious transaction must not immediately teach the system that suspicious behaviour is normal.",
            "New models must be tested safely before they affect customers or investigators.",
            "Investigators need reasons, versions, history, and rollback rather than an unexplained black-box answer.",
        ],
    )

    add_title(document, "1. The Three Applications")
    add_table(
        document,
        ["Application", "Easy description", "Main responsibility"],
        [
            ["Angular UI", "The screen users see", "Upload data, run research comparison, operate training, govern models, and review cases."],
            ["Spring Boot", "The system manager", "Own business rules, security, database work, feature orchestration, risk aggregation, alerts, cases, and audits."],
            ["Python FastAPI", "The ML worker", "Train model artifacts and calculate HST, Online OCSVM, and offline research scores."],
            ["SQL Server", "The official memory", "Store transactions, features, configuration, models, validations, deployments, predictions, alerts, and cases."],
        ],
        [1800, 3000, 4560],
    )
    add_flow(
        document,
        [
            ("Angular :4200", "User actions"),
            ("Spring :8080", "Business control"),
            ("FastAPI :8000", "ML work"),
            ("SQL Server", "Authoritative records"),
        ],
    )
    add_callout(
        document,
        "Why this separation helps",
        "The browser never talks directly to the database or ML service. Spring stays in control of security, audit, and business decisions while Python focuses on ML.",
        "success",
    )

    add_title(document, "2. What Happens When a Transaction Arrives")
    add_steps(
        document,
        [
            "Spring receives the transaction from an API or uploaded file and stores the original event in SQL Server.",
            "It loads only history that happened before the new transaction time. This prevents future information from leaking into the decision.",
            "It calculates a versioned feature vector describing amount, time, velocity, location, channel, login risk, customer behaviour, peer behaviour, and profile confidence.",
            "It permanently saves that feature vector so the exact decision can be reproduced later.",
            "It asks the scoring components for normalized evidence between 0 and 1.",
            "The versioned risk policy combines those component scores and applies any hard AML-rule override.",
            "The system stores the result, reasons, component evidence, policy version, and exact model versions.",
            "If the final decision is suspicious, Spring creates an alert and automatically opens a review case.",
            "Only after the decision does the learning policy decide whether the transaction may update trusted customer behaviour.",
        ],
    )
    add_heading(document, "Why save the feature vector?", 2)
    add_para(
        document,
        "Without a saved feature vector, the same historical transaction could produce different inputs later because the customer's profile has changed. Saving the original vector means the team can replay the exact input with the exact model and policy versions.",
    )
    add_heading(document, "Example feature ideas", 2)
    add_table(
        document,
        ["Feature", "Question it answers"],
        [
            ["Amount versus customer average", "Is this amount unusually high for this account?"],
            ["Amount z-score", "How many normal deviations away is the amount?"],
            ["Location changed", "Did the customer suddenly transact somewhere different?"],
            ["Transactions in recent window", "Is money moving unusually quickly?"],
            ["Night/weekend indicator", "Did the event happen at an unusual time?"],
            ["Peer-group deviation", "Is this unusual for comparable customers?"],
            ["Profile confidence", "Do we have enough trustworthy history to rely on personal behaviour?"],
        ],
        [3200, 6160],
    )

    add_title(document, "3. The Production Scoring Layers")
    add_para(
        document,
        "Production does not ask five classical models to vote equally. It combines different kinds of evidence because each component sees a different part of the problem.",
    )
    add_table(
        document,
        ["Component", "What it notices", "Why it is useful"],
        [
            ["Velocity and AML rules", "Fast transfers, threshold patterns, unusual sequences, and configured red flags.", "Deterministic, explainable, and able to force a high-risk result."],
            ["Customer behaviour", "Deviation from the customer's trusted amount, time, place, and channel patterns.", "Personalizes the decision."],
            ["Peer-group behaviour", "Deviation from people with similar occupation, age, account, or business characteristics.", "Helps when personal history is limited."],
            ["Half-Space Trees", "Incremental tree-based isolation patterns.", "Uses bounded memory and learns sequentially."],
            ["Online One-Class SVM", "Whether the transaction falls outside an incrementally learned normal boundary.", "Provides complementary ML evidence."],
        ],
        [2200, 4100, 3060],
    )
    add_heading(document, "How weighted risk works", 2)
    add_para(
        document,
        "Every component returns a normalized score. A versioned policy assigns each component a weight. The system multiplies each score by its weight and adds the results. Confidence can reduce weak behavioural evidence, while a severe rule can override the normal weighted calculation.",
    )
    add_table(
        document,
        ["Example component", "Score", "Weight", "Contribution"],
        [
            ["Rules", "0.80", "25%", "0.20"],
            ["Customer behaviour", "0.90", "20%", "0.18"],
            ["Peer behaviour", "0.60", "15%", "0.09"],
            ["Half-Space Trees", "0.75", "25%", "0.19"],
            ["Online OCSVM", "0.70", "15%", "0.11"],
            ["Final score", "", "100%", "0.77"],
        ],
        [3000, 1500, 1500, 3360],
    )
    add_callout(
        document,
        "This is only an example",
        "The real values come from the active versioned risk policy. A final score of 0.77 may become suspicious depending on the configured thresholds and rule overrides.",
        "info",
    )

    add_title(document, "4. A Simple Example: Account AC00455")
    add_para(
        document,
        "Suppose AC00455 normally transfers about 5,000 in familiar locations during daytime hours. A new transaction arrives for 20,000 from a new location at night after several login attempts.",
    )
    add_steps(
        document,
        [
            "The amount is four times the normal average, so customer amount deviation becomes high.",
            "The location and time differ from trusted behaviour, increasing customer anomaly evidence.",
            "Recent login attempts and transaction speed may trigger deterministic risk rules.",
            "The peer scorer checks whether this amount and channel are also unusual for comparable customers.",
            "HST and Online OCSVM score the saved feature vector against their learned normal structures.",
            "The weighted policy combines all evidence and stores reason codes such as amount deviation, location change, unusual time, or velocity risk.",
            "If the final score crosses the suspicious threshold, a case is opened for an analyst. The transaction is not automatically called fraud.",
        ],
    )
    add_callout(
        document,
        "Learning protection",
        "Because the transaction is suspicious, it can be quarantined from trusted learning. This stops a possible anomaly from immediately changing AC00455's definition of normal.",
        "warning",
    )

    add_title(document, "5. Cold Start: What If the Customer Is New?")
    add_para(
        document,
        "A new customer may have too little history for a reliable personal average. The system marks this as cold start instead of pretending that weak history is trustworthy.",
    )
    add_bullets(
        document,
        [
            "Profile confidence tells the policy how much personal evidence can be trusted.",
            "Peer-group behaviour provides broader context while personal history grows.",
            "Deterministic AML rules still work immediately.",
            "Configurable minimum-history controls prevent unstable personal comparisons.",
            "The learning policy gradually builds trusted behaviour from eligible transactions.",
        ],
    )

    add_title(document, "6. Research Comparison Versus Production")
    add_para(
        document,
        "The project has two model tracks for different purposes. Mixing them together would make the system confusing and unsafe.",
    )
    add_table(
        document,
        ["Question", "Offline research comparison", "Layered production"],
        [
            ["Main goal", "Learn how model behaviour changes as data grows.", "Score live transactions safely and explainably."],
            ["Models/components", "Isolation Forest, LOF, kernel One-Class SVM, Elliptic Envelope, PCA Reconstruction.", "Rules, customer behaviour, peer behaviour, HST, Online OCSVM, weighted policy."],
            ["Data sizes", "Oldest 10%, 25%, 50%, and 100% of one frozen snapshot.", "One saved feature vector per incoming transaction."],
            ["Result", "Fit score, proxy quality, stability, anomaly rate, latency, and growth trend.", "Final risk score, risk level, reasons, alert/case decision."],
            ["Can it deploy automatically?", "No.", "Only through authorized validation and controlled promotion."],
        ],
        [1800, 3780, 3780],
    )
    add_callout(
        document,
        "Why the five-model experiment still matters",
        "It answers the academic question: does a model become better, more stable, slower, or noisier as the available chronological data grows? It is research evidence, not confirmed-fraud accuracy.",
        "success",
    )

    add_title(document, "7. How Training Works Without Loading Millions of Rows")
    add_flow(
        document,
        [
            ("Close dates", "Freeze eligible history"),
            ("Create run", "Record window and cutoff"),
            ("Export", "Checksummed Parquet parts"),
            ("Train", "Read small batches"),
            ("Register", "Immutable candidate"),
        ],
    )
    add_steps(
        document,
        [
            "An administrator closes completed business dates so the training window cannot silently change.",
            "Spring creates a run containing the date range, cutoff, feature version, model type, segment, and actor.",
            "Spring reads eligible feature rows using keyset pagination and writes several Parquet files instead of one enormous JSON request.",
            "A manifest records columns, row counts, and SHA-256 checksums.",
            "Python verifies the dataset and reads manageable record batches.",
            "HST or Online OCSVM updates incrementally, so memory depends mainly on model settings and batch size rather than lifetime row count.",
            "Python writes a versioned artifact bundle. Spring verifies it and registers it as a candidate.",
            "The candidate remains inactive until validation and authorized promotion succeed.",
        ],
    )
    add_heading(document, "What makes training reproducible?", 2)
    add_bullets(
        document,
        [
            "Closed business dates and an exact cutoff time.",
            "Feature and model versions.",
            "Dataset and artifact checksums.",
            "Recorded parameters, metrics, row counts, actor, and timestamps.",
            "Immutable candidate bundles rather than overwriting the current model.",
        ],
    )

    add_title(document, "8. How a New Architecture Reaches Production Safely")
    add_para(
        document,
        "A model that works in training can still create too many alerts, respond slowly, or behave differently on real traffic. The system therefore uses shadow testing and gradual release.",
    )
    add_steps(
        document,
        [
            "Shadow mode calculates the layered result beside the compatibility result but does not let it create production cases.",
            "Validation checks sample size, observation days, alert volume, overlap, daily stability, synthetic scenarios, reviewed outcomes, latency, availability, and exact model versions.",
            "Only a recent PASSED report can authorize promotion for that exact customer segment.",
            "The deployment pointer locks the risk policy, HST, Online OCSVM, and validation versions together.",
            "A deterministic SHA-256 account bucket sends a stable percentage of accounts to layered production.",
            "The team can expand from a small canary toward 100% after operational review.",
            "If versions mismatch or layered scoring fails, the system returns the compatibility/fallback result.",
            "Rollback sets the entire segment to Isolation Forest fallback with zero layered traffic.",
        ],
    )
    add_heading(document, "Why deterministic canaries are better than random routing", 2)
    add_para(
        document,
        "Random routing could send the same customer through different architectures on consecutive transactions. Deterministic hashing keeps that customer in the same bucket. Expanding from 10% to 25% keeps the original 10% and adds a stable new group.",
    )

    add_title(document, "9. How Case Management Works")
    add_flow(
        document,
        [
            ("Suspicious result", "Threshold reached"),
            ("Alert", "Prediction evidence stored"),
            ("Case", "Investigator queue"),
            ("Review", "Notes and evidence"),
            ("Outcome", "False positive or draft STR"),
        ],
    )
    add_bullets(
        document,
        [
            "The case shows transaction details, final risk, versions, component/model evidence, and reasons.",
            "An analyst can add attributed and timestamped notes.",
            "A false-positive decision updates the case and related alert outcome.",
            "A suspicious case can generate draft STR XML for further institutional review.",
            "A transaction can also be searched and turned into a manual case.",
            "Duplicate manual creation returns the existing case rather than creating repeated work.",
        ],
    )
    add_callout(
        document,
        "Regulatory boundary",
        "The STR XML is a draft technical output. A real bank must map it to the required regulator schema and approval/submission process.",
        "warning",
    )

    add_title(document, "10. How the Main Problems Are Solved")
    add_table(
        document,
        ["Problem", "Solution in this system"],
        [
            ["Millions of rows do not fit in memory", "Keyset pagination, chunked Parquet, record-batch reading, and bounded incremental models."],
            ["Research models are too expensive for production", "Keep LOF, kernel SVM, Elliptic Envelope, and PCA offline; use HST and Online OCSVM for production ML."],
            ["Every customer behaves differently", "Combine customer history, peer groups, profile confidence, and global controls."],
            ["Suspicious events can poison learning", "Separate observed behaviour from trusted learning eligibility and quarantine unsafe events."],
            ["Models may change unpredictably", "Immutable candidates, silent scoring, validation gates, exact versions, canaries, and rollback."],
            ["A black-box answer is hard to investigate", "Persist component scores, reason codes, policy/model versions, feature vectors, and audit events."],
            ["ML service can fail", "Store the transaction and use an auditable compatibility/fallback result instead of inventing a fraud answer."],
            ["Labels are incomplete", "Use honest unsupervised proxy/stability evidence and reviewer outcomes without claiming population accuracy."],
            ["Too many alerts overload analysts", "Use weighted thresholds, hard-rule control, validation of alert volume, canary rollout, and reviewer feedback."],
            ["Concurrent customer transactions may corrupt state", "Keep SQL authoritative and design account-keyed ordering/locking; Kafka is the planned scale extension."],
        ],
        [3300, 6060],
    )

    add_title(document, "11. Where Kafka Fits Later")
    add_para(
        document,
        "Kafka is not the machine-learning model. It is a durable event highway that helps different workers process millions of events without tightly blocking each other.",
    )
    add_bullets(
        document,
        [
            "Partition transactions by account ID so one customer's events stay ordered.",
            "Replay events when rebuilding features or recovering from failure.",
            "Separate ingestion, feature calculation, prediction, monitoring, case creation, and training consumers.",
            "Absorb bursts through back-pressure instead of overloading one synchronous service.",
            "Publish model lifecycle events so prediction workers refresh approved versions.",
        ],
    )
    add_callout(
        document,
        "What Kafka does not solve",
        "It does not train models, guarantee correct business logic, or replace SQL. The system still needs idempotency, schema evolution, state management, dead-letter handling, monitoring, and an outbox to avoid lost dual writes.",
        "info",
    )

    add_title(document, "12. What Users Do on Each Page")
    add_table(
        document,
        ["Page", "Use it for"],
        [
            ["Upload Data", "Load historical CSV or Excel transactions."],
            ["Research Comparison", "Compare five offline models at 10%, 25%, 50%, and 100% chronological data."],
            ["Training Operations", "Close dates, create runs, export Parquet, and train HST/Online OCSVM candidates."],
            ["Model Governance", "Review evidence, validate, activate/expand layered canaries, and roll back."],
            ["Model Tuning", "Configure offline search and model parameters."],
            ["Anomaly Config", "Maintain the compatibility decision configuration for unpromoted segments."],
            ["Cold Start", "Control behaviour when trusted customer history is insufficient."],
            ["Case Management", "Investigate generated cases, record notes, resolve false positives, and generate draft STR XML."],
            ["Manual Case", "Search a transaction and open a case on demand."],
        ],
        [2600, 6760],
    )

    add_title(document, "13. A Normal Daily Operating Story")
    add_heading(document, "During business hours", 2)
    add_steps(
        document,
        [
            "Transactions arrive and are stored, featured, scored, logged, and routed to cases when suspicious.",
            "AML reviewers work the case queue, add notes, and record outcomes.",
            "Operators monitor availability, latency, fallback frequency, and unusual alert volume.",
        ],
    )
    add_heading(document, "After business hours", 2)
    add_steps(
        document,
        [
            "Confirm ingestion is complete and close the business date.",
            "Create the approved incremental training run and export the frozen eligible dataset.",
            "Train candidate HST or Online OCSVM versions.",
            "Keep candidates in shadow until enough evidence exists.",
            "Validate and promote only through an authorized governance decision.",
        ],
    )
    add_heading(document, "Periodically", 2)
    add_bullets(
        document,
        [
            "Run offline research comparison after meaningful data growth.",
            "Review false-positive outcomes and alert volume by segment.",
            "Run fresh rebuilds and historical replay when implemented.",
            "Review feature drift, score drift, model freshness, latency, and canary results.",
            "Practice rollback rather than assuming it works.",
        ],
    )

    add_title(document, "14. What the System Does Not Claim")
    add_bullets(
        document,
        [
            "An anomaly is not automatically fraud.",
            "Proxy quality is not real-world fraud accuracy.",
            "A passed validation does not prove perfect detection.",
            "A research leaderboard does not authorize production deployment.",
            "Kafka is planned, not part of the current runtime.",
            "The draft STR XML is not yet a complete regulator submission workflow.",
            "Development credentials, local files, and local transport require production hardening before bank deployment.",
        ],
    )

    add_title(document, "15. Quick Glossary")
    add_table(
        document,
        ["Term", "Simple meaning"],
        [
            ["Anomaly", "Something different from learned or configured normal behaviour."],
            ["Feature", "A number or category that describes one useful aspect of a transaction."],
            ["Behaviour scorer", "Transparent statistical comparison against customer or peer patterns."],
            ["HST", "Half-Space Trees, an incremental production anomaly model with bounded state."],
            ["Online OCSVM", "Incremental production model that learns a boundary around normal feature patterns."],
            ["Risk policy", "Versioned weights, thresholds, confidence handling, and override rules."],
            ["Shadow mode", "Calculate and record a new result without allowing it to control production."],
            ["Canary", "A small stable group that receives the new architecture first."],
            ["Fallback", "The safe known decision path used when layered scoring is disabled or unsafe."],
            ["Candidate", "A trained model version that is not yet active."],
            ["Validation", "Evidence checks required before promotion."],
            ["STR", "Suspicious Transaction Report; this system currently creates draft XML."],
        ],
        [2600, 6760],
    )

    add_title(document, "Final Picture to Remember")
    add_flow(
        document,
        [
            ("Remember", "Store exact data and features"),
            ("Understand", "Compare customer and peers"),
            ("Detect", "Rules + scalable ML"),
            ("Control", "Validate, canary, rollback"),
            ("Investigate", "Human case decision"),
        ],
    )
    add_para(
        document,
        "The system solves the problem by separating concerns. Research models answer how behaviour changes as data grows. Scalable production components score transactions. Governance controls model change. SQL preserves evidence. Human reviewers make the final investigation decision. This separation is what makes the project understandable, safer, and capable of growing toward bank-scale operation.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_guide())
