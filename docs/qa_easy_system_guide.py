from pathlib import Path
import re
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


PATH = Path("docs/deliverables/AML_System_Explained_Simply.docx")


def width_value(element, child_name: str) -> int | None:
    child = element.find(qn(child_name))
    if child is None:
        return None
    value = child.get(qn("w:w"))
    return int(value) if value else None


with ZipFile(PATH) as package:
    assert package.testzip() is None, "DOCX package is corrupt"
    required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
    assert required.issubset(set(package.namelist())), "Required DOCX parts are missing"

document = Document(PATH)
headings = [paragraph for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")]
paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
table_text = [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
all_text = paragraphs + table_text
assert len(paragraphs) >= 110, "Guide is unexpectedly short"
assert len(headings) >= 25, "Guide needs stronger semantic navigation"
assert any("whole system in one minute" in text.lower() for text in all_text)
assert any("AC00455" in text for text in all_text)
assert any("does not say that fraud is proven" in text.lower() for text in all_text)
assert not any(re.search(r"turn\d+(search|fetch)\d+", text) for text in all_text), "Tool citation leaked"

for table_index, table in enumerate(document.tables, start=1):
    table_width = width_value(table._tbl.tblPr, "w:tblW")
    grid_widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
    assert table_width == sum(grid_widths), f"Table {table_index} width mismatch"
    for row_index, row in enumerate(table.rows, start=1):
        cell_widths = []
        for cell in row.cells:
            width = width_value(cell._tc.get_or_add_tcPr(), "w:tcW")
            assert width is not None, f"Missing width in table {table_index}, row {row_index}"
            assert len(cell.text) < 700, f"Over-dense table {table_index}"
            cell_widths.append(width)
        assert sum(cell_widths) == table_width, f"Cell width mismatch in table {table_index}, row {row_index}"

print(
    f"PASS {PATH.name}: paragraphs={len(paragraphs)}, headings={len(headings)}, "
    f"tables={len(document.tables)}, size={PATH.stat().st_size} bytes"
)
