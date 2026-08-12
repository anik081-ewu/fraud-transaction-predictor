from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


FILES = [
    Path("docs/deliverables/Fraud_Transaction_Detector_Full_User_Guide_v2.docx"),
    Path("docs/deliverables/Fraud_Transaction_Detector_System_Design_Implementation_and_Scalability_v2.docx"),
]


def width_value(element, child_name: str) -> int | None:
    child = element.find(qn(child_name))
    if child is None:
        return None
    value = child.get(qn("w:w"))
    return int(value) if value else None


for path in FILES:
    with ZipFile(path) as package:
        bad = package.testzip()
        assert bad is None, f"Corrupt package member in {path}: {bad}"
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        assert required.issubset(set(package.namelist())), f"Missing DOCX parts in {path}"

    document = Document(path)
    headings = [p for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert headings, f"No semantic headings in {path}"
    assert all(p.text.strip() for p in headings), f"Empty heading in {path}"

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    assert len(paragraphs) >= 100, f"Document is unexpectedly short: {path}"
    assert not any("turn" in text and "search" in text for text in paragraphs), f"Tool citation leaked: {path}"

    for table_index, table in enumerate(document.tables, start=1):
        tbl_pr = table._tbl.tblPr
        table_width = width_value(tbl_pr, "w:tblW")
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert table_width == sum(grid_widths), f"Table {table_index} width mismatch in {path}"
        for row_index, row in enumerate(table.rows, start=1):
            cell_widths = []
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                value = width_value(tc_pr, "w:tcW")
                assert value is not None, f"Missing cell width in table {table_index}, row {row_index}: {path}"
                cell_widths.append(value)
                assert len(cell.text) < 900, f"Over-dense cell in table {table_index}: {path}"
            assert sum(cell_widths) == table_width, f"Cell widths mismatch in table {table_index}, row {row_index}: {path}"

    print(
        f"PASS {path.name}: paragraphs={len(paragraphs)}, headings={len(headings)}, "
        f"tables={len(document.tables)}, size={path.stat().st_size} bytes"
    )
