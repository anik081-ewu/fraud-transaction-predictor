from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = OUTPUT_DIR / "Anomaly_Model_Measurement_Guide.docx"

NAVY = "17324D"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
PALE_BLUE = "F5F8FC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "087A55"
PALE_GREEN = "E9F8F1"
AMBER = "8A5A00"
PALE_AMBER = "FFF5D6"
RED = "A12424"
WHITE = "FFFFFF"
BLACK = "111827"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    properties = table._tbl.tblPr
    width_element = properties.first_child_found_in("w:tblW")
    if width_element is None:
        width_element = OxmlElement("w:tblW")
        properties.append(width_element)
    width_element.set(qn("w:w"), str(sum(widths)))
    width_element.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, header_fill=NAVY):
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.2)
                    run.font.color.rgb = RGBColor.from_string(BLACK)


def add_table(document, headers, rows, widths, header_fill=NAVY):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
    set_table_widths(table, widths)
    style_table(table, header_fill)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(document, label, text, fill=LIGHT_BLUE, label_color=BLUE):
    table = document.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(label_color)
    paragraph.add_run(text)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_definition(document, term, definition):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    label = paragraph.add_run(f"{term}. ")
    label.bold = True
    label.font.color.rgb = RGBColor.from_string(NAVY)
    paragraph.add_run(definition)


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.7)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)


def add_page_number(paragraph):
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def build_document():
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    header = section.header.paragraphs[0]
    header.text = "FRAUD TRANSACTION DETECTOR | TECHNICAL GUIDE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.runs[0]
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Anomaly model measurement methodology | ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_page_number(footer)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(22)
    kicker.paragraph_format.space_after = Pt(6)
    kicker_run = kicker.add_run("ANOMALY MODEL COMPARISON")
    kicker_run.bold = True
    kicker_run.font.size = Pt(10)
    kicker_run.font.color.rgb = RGBColor.from_string(BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("How the Models and Report Metrics Are Measured")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(25)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        "A professor-ready explanation of native anomaly scores, proxy quality, resampling stability, "
        "data-growth behavior, fit score, and assessment labels"
    )
    subtitle_run.font.size = Pt(12.5)
    subtitle_run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    add_callout(
        document,
        "Central claim",
        "Because the current dataset has no confirmed fraud labels, this system does not report real fraud "
        "accuracy, precision, recall, or F1. It reports label-free model-selection indicators that measure "
        "separation on controlled synthetic challenges, prediction consistency, and behavior as data grows.",
        fill=PALE_AMBER,
        label_color=AMBER,
    )

    document.add_heading("1. The Three Measurement Layers", level=1)
    add_definition(
        document,
        "Layer 1 - Native model score",
        "Each algorithm produces its own score with its own meaning. These raw values are useful within one "
        "trained model, but their magnitudes must not be directly compared across different algorithms.",
    )
    add_definition(
        document,
        "Layer 2 - Common proxy evaluation",
        "Every candidate model is tested on the same chronological validation rows and the same generated "
        "anomaly challenges. ROC AUC, average precision, anomaly-rate control, and bootstrap agreement convert "
        "the different models into common 0-100 indicators.",
    )
    add_definition(
        document,
        "Layer 3 - Data-growth comparison",
        "The same process is repeated on the oldest 10%, 25%, 50%, and 100% of transactions. The UI averages "
        "the common indicators and measures how quality changes from 10% to 100%.",
    )

    document.add_heading("2. Native Measurement Unit of Each Model", level=1)
    paragraph = document.add_paragraph(
        "All input features are standardized before fitting. Numeric variables are therefore expressed as "
        "z-scores, while one-hot categorical indicators are transformed by the same scaler. Except for PCA's "
        "reconstruction error, the native decision values are dimensionless, model-specific scores."
    )
    paragraph.paragraph_format.space_after = Pt(8)

    native_table = add_table(
        document,
        ["Model", "What it measures", "Native output used", "How to read it"],
        [
            [
                "Isolation Forest",
                "How quickly random trees isolate a transaction. Shorter paths imply unusualness.",
                "decision_function; dimensionless path-based score",
                "Below 0 = anomaly; above 0 = inlier. Larger positive values are more normal.",
            ],
            [
                "Local Outlier Factor (LOF)",
                "Local density relative to nearby transactions.",
                "decision_function; shifted local-density score",
                "Below 0 = anomaly. The underlying LOF factor is near 1 for typical local density and larger for outliers.",
            ],
            [
                "One-Class SVM",
                "Position relative to a learned boundary around normal data in kernel space.",
                "decision_function; signed boundary value",
                "Below 0 = anomaly; above 0 = inlier. Magnitude is not a probability.",
            ],
            [
                "Elliptic Envelope",
                "Robust covariance distance from an assumed elliptical normal-data region.",
                "decision_function; shifted robust distance score",
                "Below 0 = anomaly. Best suited when standardized data is approximately elliptical.",
            ],
            [
                "PCA Reconstruction",
                "Information lost when a transaction is compressed and reconstructed.",
                "Mean squared reconstruction error; squared standardized-feature units",
                "Error above the learned percentile threshold = anomaly. UI decision value = threshold minus error.",
            ],
        ],
        [1600, 2900, 2200, 2660],
    )
    set_repeat_table_header(native_table.rows[0])
    add_callout(
        document,
        "Do not compare raw scores",
        "For example, an SVM decision value of -2 is not automatically more anomalous than an Isolation Forest "
        "value of -0.2. Each value is calibrated only against that model's own zero boundary and score distribution.",
    )

    document.add_heading("3. How Training and Evaluation Work", level=1)
    add_table(
        document,
        ["Stage", "Current implementation"],
        [
            ["Chronological partitions", "Oldest 10%, 25%, 50%, and 100% of transactions by transaction_date."],
            [
                "Chronological holdout",
                "The newest part inside each partition is validation data. Default size is 20%, bounded between 50 and 2,000 rows.",
            ],
            [
                "Optimization training limit",
                "Candidate tuning uses at most the latest 5,000 pre-validation rows to control training cost.",
            ],
            [
                "Synthetic anomaly challenge",
                "A copy of every validation row is perturbed in 20% of features by 3 to 6 standardized units in random directions.",
            ],
            [
                "Candidate selection",
                "Several hyperparameter combinations are fitted; the candidate with the highest proxy score is selected.",
            ],
            [
                "Stability test",
                "The selected candidate is refitted three times on bootstrap samples containing 85% of training-row count.",
            ],
            [
                "Final partition model",
                "The selected hyperparameters are used to train the model artifact for that partition.",
            ],
        ],
        [2100, 7260],
    )

    document.add_heading("4. Proxy Quality Metrics", level=1)
    add_definition(
        document,
        "Proxy Average Precision (AP), range 0 to 1",
        "Ranks the generated anomaly challenges ahead of original validation rows, emphasizing precision-recall "
        "behavior. Higher is better. It is a proxy because the positive class is generated, not reviewer-confirmed fraud.",
    )
    add_definition(
        document,
        "Proxy ROC AUC, range 0 to 1",
        "The probability that a randomly selected generated anomaly receives a more anomalous score than a randomly "
        "selected original validation row. A value near 0.5 indicates chance ranking; 1.0 indicates perfect proxy separation.",
    )
    add_definition(
        document,
        "Validation anomaly rate, fraction 0 to 1",
        "The share of original validation rows predicted as anomalies. This behaves like a proxy false-positive rate, "
        "but it cannot be called a true false-positive rate because original rows may contain unknown real anomalies.",
    )
    add_definition(
        document,
        "Rate-control score, range 0 to 100",
        "Rewards a validation anomaly rate near the configured target, currently 5%. It is calculated as "
        "100 x max(0, 1 - |observed rate - target rate| / max(target rate, 0.01)).",
    )

    add_callout(
        document,
        "Proxy score equation",
        "ProxyScore = 100 x (0.50 x AveragePrecision + 0.35 x ROC_AUC + 0.15 x RateControlFraction). "
        "The result is on a 0-100 scale.",
        fill=PALE_GREEN,
        label_color=GREEN,
    )

    document.add_heading("5. Resampling Stability", level=1)
    paragraph = document.add_paragraph(
        "Stability asks: if the available training rows change slightly, does the model still classify the same "
        "validation transactions in the same way?"
    )
    paragraph.paragraph_format.space_after = Pt(5)
    add_definition(
        document,
        "Procedure",
        "Create three bootstrap training samples. Each sample draws with replacement and has a size equal to 85% "
        "of the original training-row count, with at least 20 rows. Fit three copies of the selected model and "
        "predict the same chronological validation set.",
    )
    add_definition(
        document,
        "Agreement calculation",
        "For each validation row, compare the binary anomaly/inlier decisions for model pairs (1,2), (1,3), and "
        "(2,3). Compute the fraction of equal decisions for each pair and take the mean.",
    )
    add_callout(
        document,
        "Stability equation",
        "StabilityScore = 100 x mean(pairwise prediction agreement across the three bootstrap fits). "
        "100 means identical decisions; 80 means 80% average agreement.",
    )
    add_callout(
        document,
        "Interpretation limit",
        "High stability means repeatability, not correctness. A consistently wrong model can still be highly stable.",
        fill=PALE_AMBER,
        label_color=AMBER,
    )

    document.add_heading("6. What the Training Report Columns Mean", level=1)
    report_table = add_table(
        document,
        ["UI column", "Unit and formula", "Interpretation"],
        [
            [
                "Fit score",
                "0-100 points. 0.65 x AvgQuality + 0.35 x AvgStability.",
                "Overall ranking score used to sort models. Higher is preferred under the current label-free policy.",
            ],
            [
                "Avg. quality",
                "0-100 points. Arithmetic mean of qualityScore across 10%, 25%, 50%, and 100%.",
                "Average candidate quality over increasing data volumes.",
            ],
            [
                "Avg. stability",
                "Percent. Arithmetic mean of bootstrap agreement scores across all four partitions.",
                "How repeatable the model's binary decisions are when training data is resampled.",
            ],
            [
                "Avg. anomaly rate",
                "Percent. Mean of anomaly_count / partition_rows across all four trained partitions.",
                "Operational alert volume. It is not accuracy and is not automatically better when lower.",
            ],
            [
                "10% to 100%",
                "Percentage-point change: Quality_100% - Quality_10%.",
                "Positive means measured quality improved with more data; negative means it declined.",
            ],
            [
                "Assessment",
                "Rule-based label from stability and 10%-to-100% quality change.",
                "Strong and stable, Moderately stable, Needs review, or Insufficient metrics.",
            ],
        ],
        [1550, 3550, 4260],
    )
    set_repeat_table_header(report_table.rows[0])

    document.add_heading("7. The Exact Meaning of 'Avg. Quality'", level=1)
    paragraph = document.add_paragraph(
        "In the current code, the metric stored as qualityScore is not the raw ProxyScore alone. It is already a "
        "composite that includes stability and rate control:"
    )
    paragraph.paragraph_format.space_after = Pt(5)
    add_callout(
        document,
        "Current partition quality equation",
        "QualityScore = 0.65 x ProxyScore + 0.25 x StabilityScore + 0.10 x RateControlScore.",
        fill=PALE_GREEN,
        label_color=GREEN,
    )
    paragraph = document.add_paragraph(
        "The UI then computes FitScore = 0.65 x AvgQuality + 0.35 x AvgStability. Therefore, stability contributes "
        "inside AvgQuality and again directly in FitScore. The page wording 'proxy quality (65%) and resampling "
        "stability (35%)' is a simplified description, but technically the first component is composite quality."
    )
    paragraph.paragraph_format.space_after = Pt(6)
    add_callout(
        document,
        "Expanded fit equation",
        "FitScore = 0.4225 x AvgProxyScore + 0.5125 x AvgStability + 0.065 x AvgRateControl. "
        "This follows by substituting the current QualityScore equation.",
        fill=PALE_AMBER,
        label_color=AMBER,
    )
    add_callout(
        document,
        "Recommended professor wording",
        "Say: 'The current fit score is a label-free composite dominated by resampling stability, with additional "
        "synthetic-challenge separation and anomaly-rate calibration.' Do not describe it as fraud accuracy.",
    )

    document.add_heading("8. Assessment Rules", level=1)
    assessment_table = add_table(
        document,
        ["Assessment", "Rule used by the UI", "Plain-language meaning"],
        [
            [
                "Strong and stable",
                "AvgStability >= 80 and quality change >= -5 points.",
                "Highly repeatable and does not materially deteriorate as data grows.",
            ],
            [
                "Moderately stable",
                "AvgStability >= 65 and quality change >= -10 points.",
                "Reasonably repeatable with tolerable data-growth degradation.",
            ],
            [
                "Needs review",
                "Any result below the moderate thresholds.",
                "Predictions vary too much or quality falls too sharply as data grows.",
            ],
            [
                "Insufficient metrics",
                "Stability is unavailable.",
                "The system cannot produce a defensible assessment.",
            ],
        ],
        [1900, 3350, 4110],
    )
    set_repeat_table_header(assessment_table.rows[0])

    document.add_heading("9. Worked Example", level=1)
    paragraph = document.add_paragraph(
        "Suppose one model has partition quality scores of 76, 79, 82, and 84; stability scores of 80%, 82%, "
        "84%, and 86%; and anomaly rates of 5%, 5.5%, 4.8%, and 5.2%."
    )
    add_table(
        document,
        ["Calculation", "Result"],
        [
            ["AvgQuality = (76 + 79 + 82 + 84) / 4", "80.25 points"],
            ["AvgStability = (80 + 82 + 84 + 86) / 4", "83.00%"],
            ["AvgAnomalyRate = (5 + 5.5 + 4.8 + 5.2) / 4", "5.125%"],
            ["10% to 100% = 84 - 76", "+8.00 points"],
            ["FitScore = 0.65 x 80.25 + 0.35 x 83.00", "81.21 / 100"],
            ["Assessment", "Strong and stable"],
        ],
        [6200, 3160],
        header_fill=BLUE,
    )
    paragraph = document.add_paragraph(
        "Conclusion: the model is ranked strongly because its quality rises as data grows, its decisions remain "
        "consistent across resampled training sets, and its anomaly volume remains close to the target. This still "
        "does not prove that the flagged transactions are real fraud."
    )

    document.add_heading("10. How to Explain the Five Models to a Professor", level=1)
    speaking_table = add_table(
        document,
        ["Model", "One-sentence explanation", "Main strength", "Main risk"],
        [
            [
                "Isolation Forest",
                "Anomalies are transactions that random trees isolate using fewer splits.",
                "Scales well and detects global anomalies.",
                "Contamination and random sampling affect the boundary.",
            ],
            [
                "LOF",
                "Anomalies have much lower local density than their nearest neighbors.",
                "Finds anomalies hidden inside different local groups.",
                "Sensitive to neighborhood size and high-dimensional distance.",
            ],
            [
                "One-Class SVM",
                "Anomalies lie outside a nonlinear boundary learned around normal behavior.",
                "Flexible nonlinear boundary.",
                "Sensitive to scaling, gamma, nu, and dataset size.",
            ],
            [
                "Elliptic Envelope",
                "Anomalies are far from a robust multivariate center under an elliptical assumption.",
                "Interpretable covariance-based distance.",
                "Weak fit when the data is strongly non-elliptical or multimodal.",
            ],
            [
                "PCA Reconstruction",
                "Anomalies cannot be reconstructed well from the dominant normal-data components.",
                "Captures unusual combinations of correlated features.",
                "Linear PCA can miss nonlinear patterns.",
            ],
        ],
        [1500, 3150, 2200, 2510],
    )
    set_repeat_table_header(speaking_table.rows[0])

    document.add_heading("11. What Can and Cannot Be Claimed", level=1)
    claims_table = add_table(
        document,
        ["Safe claims now", "Claims that require reviewer-confirmed labels"],
        [
            ["Model A is more stable under bootstrap resampling.", "Model A has higher real fraud precision."],
            ["Model A separates generated anomaly challenges more strongly.", "Model A catches more actual fraud."],
            ["Model A's anomaly rate is near the configured operating target.", "The reported anomaly rate is a true false-positive rate."],
            ["Model A improves or deteriorates as training data grows.", "Model A has validated production accuracy or F1."],
            ["The ensemble creates cases when configured votes reach the threshold.", "A generated case confirms that fraud occurred."],
        ],
        [4680, 4680],
        header_fill=BLUE,
    )
    set_repeat_table_header(claims_table.rows[0])
    add_callout(
        document,
        "Future validation",
        "After reviewers label alerts as suspicious or false positive, report precision, recall, F1, false-positive "
        "rate, confusion matrix, precision-recall curves, calibration, subgroup checks, and time-based backtesting.",
        fill=PALE_GREEN,
        label_color=GREEN,
    )

    document.add_heading("12. Short Viva Script", level=1)
    add_definition(
        document,
        "Why unsupervised models?",
        "The dataset does not contain confirmed fraud labels, so supervised fraud accuracy cannot yet be estimated.",
    )
    add_definition(
        document,
        "How do you compare different algorithms?",
        "Native scores are not compared directly. Every model is evaluated on a shared chronological holdout and "
        "shared synthetic anomaly challenges, then converted to common 0-100 quality and stability indicators.",
    )
    add_definition(
        document,
        "Why use four dataset sizes?",
        "The 10%, 25%, 50%, and 100% partitions reveal whether a model remains stable and improves when more historical data becomes available.",
    )
    add_definition(
        document,
        "What does the best-fit model mean?",
        "It is the model with the highest current label-free composite score, not necessarily the model with the "
        "highest real fraud-detection accuracy.",
    )
    add_definition(
        document,
        "Why use voting?",
        "Different anomaly models capture different structures. Voting reduces dependence on one algorithm and "
        "lets the business choose how much agreement is required before creating a case.",
    )

    document.add_heading("13. References and Implementation Basis", level=1)
    references = [
        "Project implementation: fraud-ml-service/app/training.py - candidate evaluation, synthetic anomalies, "
        "stability, quality score, anomaly rate, and decision percentiles.",
        "Project implementation: fraud-transaction-ui/src/app/pages/datasets-page.component.ts - averages, fit "
        "score, 10%-to-100% change, sorting, and assessment thresholds.",
        "Scikit-learn User Guide: Novelty and Outlier Detection, including Isolation Forest, LOF, One-Class SVM, "
        "and Elliptic Envelope. https://scikit-learn.org/stable/modules/outlier_detection.html",
        "Scikit-learn API documentation: IsolationForest. "
        "https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.IsolationForest.html",
        "Scikit-learn API documentation: LocalOutlierFactor. "
        "https://scikit-learn.org/stable/modules/generated/sklearn.neighbors.LocalOutlierFactor.html",
        "Scikit-learn API documentation: OneClassSVM. "
        "https://scikit-learn.org/stable/modules/generated/sklearn.svm.OneClassSVM.html",
        "Mihelich et al. (2024), Interplay of ROC and Precision-Recall AUCs, Proceedings of Machine Learning Research 235.",
    ]
    for index, reference in enumerate(references, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        number = paragraph.add_run(f"{index}. ")
        number.bold = True
        paragraph.add_run(reference)

    add_callout(
        document,
        "Final takeaway",
        "The report is a defensible engineering comparison for model selection under missing labels. Its strongest "
        "scientific language is 'proxy quality and robustness assessment.' Real fraud effectiveness must be added "
        "later using reviewer-confirmed outcomes.",
        fill=PALE_BLUE,
        label_color=BLUE,
    )

    for table in document.tables:
        if table.rows:
            set_repeat_table_header(table.rows[0])

    document.core_properties.title = "Anomaly Model Measurement Guide"
    document.core_properties.subject = "Fraud transaction detector model comparison methodology"
    document.core_properties.author = "Fraud Transaction Detector Project"
    document.core_properties.keywords = "anomaly detection, model comparison, stability, proxy quality"
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
