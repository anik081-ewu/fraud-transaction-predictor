from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "deliverables" / "Intelligent_Fraud_Transaction_Detection_and_AML_Case_Management_Report_with_UI_Walkthrough.docx"
OUTPUT = ROOT / "docs" / "deliverables" / "Intelligent_Fraud_Transaction_Detection_and_AML_Case_Management_Report_Updated_Current_Results.docx"
ASSET_DIR = ROOT / "docs" / "assets" / "current-results"

UNSUPERVISED_IMAGE = ROOT / "docs" / "assets" / "ui-report" / "08-unsupervised-models.png"
SUPERVISED_SOURCE_IMAGE = Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-b0494cd7-7cc1-410b-b877-261203d88c41.png")
POLICY_SOURCE_IMAGE = Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-01788572-3b3f-4283-8338-c96588ba3e52.png")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeatable_font(run, size=9.5, color="22324A", bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def find_paragraph(document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text)


def replace_exact(document, old: str, new: str):
    paragraph = find_paragraph(document, old)
    paragraph.text = new


def remove_range(start, end):
    parent = start._p.getparent()
    elements = list(parent)
    start_index = elements.index(start._p)
    end_index = elements.index(end._p)
    for element in elements[start_index:end_index]:
        parent.remove(element)


def add_paragraph_before(target, text="", style=None, keep=False, page_break=False):
    paragraph = target.insert_paragraph_before(text, style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = keep
    paragraph.paragraph_format.page_break_before = page_break
    return paragraph


def add_bullet_before(target, text: str):
    paragraph = add_paragraph_before(target, text, style="List Bullet1")
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_table_before(document, target, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "DCE8FF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_repeatable_font(run, 9, "173B7A", True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_repeatable_font(run, 8.7)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    target._p.addprevious(table._tbl)
    return table


def add_picture_before(target, path: Path, width: float, caption: str):
    paragraph = target.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    caption_paragraph = target.insert_paragraph_before(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    set_repeatable_font(caption_run, 9, "53657D", italic=True)


def copy_result_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    supervised = ASSET_DIR / "supervised-current-measured.png"
    policy = ASSET_DIR / "risk-policy-current-measured.png"
    if SUPERVISED_SOURCE_IMAGE.exists():
        shutil.copy2(SUPERVISED_SOURCE_IMAGE, supervised)
    if POLICY_SOURCE_IMAGE.exists():
        shutil.copy2(POLICY_SOURCE_IMAGE, policy)
    return supervised, policy


def update_existing_content(document):
    replacements = {
        "When labels are unavailable, the system uses three complementary detectors. Isolation Forest isolates rare observations through randomized trees; the Autoencoder learns to reconstruct common feature patterns and treats high reconstruction error as unusual; Local Outlier Factor compares the local density around one transaction with the density around its neighbors. Their outputs indicate abnormality rather than confirmed fraud.":
            "When labels are unavailable, the system uses three complementary detectors. Isolation Forest isolates rare observations through randomized trees; the Autoencoder learns common feature patterns and treats high reconstruction error as unusual; Behavioral Cluster Outlier measures cluster-conditional deviation from comparable transaction behaviour. Their outputs indicate abnormality rather than confirmed fraud.",
        "When an auditable binary label is available, supervised classifiers can learn the relationship between historical features and reviewed outcomes. XGBoost provides a strong nonlinear boosted-tree model, Random Forest provides a robust class-balanced ensemble, and Logistic Regression provides an interpretable probability baseline. The comparison protocol emphasizes precision-recall measures because fraud datasets are normally imbalanced.":
            "When an auditable binary label is available, supervised classifiers learn relationships between point-in-time features and reviewed outcomes. XGBoost provides scalable nonlinear boosting, Class-Balanced Random Forest provides a robust bagged-tree benchmark, and Extra Trees adds stronger split randomization and computational efficiency. Fusion strategies remain research comparisons; the production selector exposes only the three base classifiers. The comparison protocol emphasizes precision-recall measures because fraud is highly imbalanced.",
        "Supervised mode requires FraudLabel values and activates XGBoost, Random Forest, and Logistic Regression workflows.":
            "Supervised mode requires FraudLabel values and activates XGBoost, Class-Balanced Random Forest, and Extra Trees workflows.",
        "Current Isolation Forest, Autoencoder, and Local Outlier Factor diagnostics.":
            "Current Isolation Forest, Autoencoder, and Behavioral Cluster Outlier diagnostics.",
        "The page exposes only controls consumed by training or comparison, including chronological holdout settings, reproducibility seed, and model-specific hyperparameters.":
            "The page exposes only controls consumed by training or comparison, including chronological holdout settings, reproducibility seed, class-imbalance handling, and model-specific hyperparameters for XGBoost, Class-Balanced Random Forest, and Extra Trees.",
        "The current evidence proves that the revised standalone model exceeds 70% precision and 70% recall on an untouched chronological period. It does not yet prove that the complete four-layer policy beats Extra Trees. That comparison must be rerun after migration and retraining; the Supervised Comparison page reports both results on the same held-out rows and should be used to make that claim only when its measured deltas are positive.":
            "The current evidence proves that Extra Trees exceeds 70% precision and 70% recall on the untouched newest period of the 264,500-row benchmark. The latest measured four-layer policy replay on the earlier 25,000-row integration snapshot does not outperform the standalone classifiers. The report therefore treats model quality and complete policy quality as separate claims and requires a fresh measured replay after every policy or feature change.",
    }
    for old, new in replacements.items():
        try:
            replace_exact(document, old, new)
        except StopIteration:
            pass

    catalog = document.tables[3]
    catalog.cell(1, 1).text = "Isolation Forest; Autoencoder; Behavioral Cluster Outlier"
    catalog.cell(2, 1).text = "XGBoost; Class-Balanced Random Forest; Extra Trees"

    problems = document.tables[8]
    problems.cell(1, 1).text = "Global unsupervised mode with Isolation Forest, Autoencoder, and Behavioral Cluster Outlier; label-free metrics"
    problems.cell(2, 1).text = "Separate supervised mode with XGBoost, Class-Balanced Random Forest, Extra Trees, and supervised metrics"

    references = {
        "[5] D. P. Kingma and M. Welling, 'Auto-Encoding Variational Bayes,' International Conference on Learning Representations, 2014.":
            "[5] M. Sakurada and T. Yairi, 'Anomaly Detection Using Autoencoders with Nonlinear Dimensionality Reduction,' MLSDA, 2014, doi:10.1145/2689746.2689747.",
    }
    for old, new in references.items():
        try:
            replace_exact(document, old, new)
        except StopIteration:
            pass


def add_research_matrix(document):
    target = find_paragraph(document, "CHAPTER 3")
    add_paragraph_before(target, "2.6 Research-Guided Design Decisions", style="Heading 2", keep=True)
    add_paragraph_before(
        target,
        "The project follows published methods as design guidance rather than claiming to reproduce every paper exactly. Each source contributes a concrete decision that is visible in code, evaluation, or governance.",
    )
    add_table_before(document, target,
        ["Research source", "Decision adopted in this system"],
        [
            ["Fraud Detection Handbook (Le Borgne et al., 2022)", "Reproducible simulator data, chronological evaluation, severe class imbalance, PR-oriented metrics, and investigator-capacity awareness."],
            ["Liu, Ting and Zhou (Isolation Forest)", "A scalable label-free detector that isolates rare observations and supports bounded subsampling."],
            ["Sakurada and Yairi (Autoencoder anomaly detection)", "Reconstruction error represents deviation from learned normal behaviour."],
            ["Goix, Sabourin and Clemencon (Excess-Mass curves)", "EM-AUC is used as a label-free quality diagnostic instead of inventing accuracy without labels."],
            ["Breiman; Geurts et al.; Chen and Guestrin", "Class-Balanced Random Forest, Extra Trees, and XGBoost form the supervised base-model comparison."],
            ["Bach Nguyen et al. (future information)", "Real-time scoring uses past-only features; later transactions may be used only for posterior investigation or later retraining, preventing online leakage."],
        ],
        widths=[2.2, 4.2],
    )


def rebuild_results(document, supervised_image: Path, policy_image: Path):
    start = find_paragraph(document, "6.2 Current Unsupervised Training Result")
    chapter_seven = find_paragraph(document, "CHAPTER 7")
    remove_range(start, chapter_seven)

    add_paragraph_before(chapter_seven, "6.2 Dataset Sources and Evidence Boundaries", style="Heading 2", keep=True, page_break=True)
    add_paragraph_before(
        chapter_seven,
        "Two evidence tracks are intentionally separated. The unsupervised track uses an unlabeled 50,000-row development snapshot and therefore reports label-free diagnostics only. The main supervised benchmark uses the open Fraud Detection Handbook simulated-data-transformed source created by Le Borgne, Siblini, Lebichot, and Bontempi. The project selected complete customer histories and preserved source chronology, amounts, terminal identity, and TX_FRAUD labels; TERMINAL_ID is mapped to Location so the application can derive novelty without using the label as a feature.",
    )
    add_table_before(document, chapter_seven,
        ["Evidence set", "Rows / period", "Purpose and boundary"],
        [
            ["Unlabeled development snapshot", "50,000 rows", "Unsupervised EM-AUC, score shape, anomaly rate, runtime. No fraud-accuracy claim."],
            ["Fraud Detection Handbook adapted subset", "264,500 rows; 749 customers; 183 days; 1 Apr-30 Sep 2018", "Main reproducible supervised benchmark; 2,340 fraud rows (0.8847%)."],
            ["Application integration snapshot", "25,000 labelled rows", "End-to-end UI, fusion, growth, and four-layer policy replay; synthetic demonstration evidence."],
        ],
        widths=[1.7, 1.7, 3.0],
    )
    note = add_paragraph_before(chapter_seven)
    run = note.add_run("Leakage boundary. ")
    set_repeatable_font(run, 9.5, "174EA6", True)
    run = note.add_run("Fraud labels and fraud-scenario identifiers are never model inputs. Training uses the oldest 70%, threshold calibration uses the next 10%, and final reporting uses the newest untouched 20% for the reproducible benchmark.")
    set_repeatable_font(run, 9.5)

    add_paragraph_before(chapter_seven, "6.3 Current Unsupervised Diagnostic Result", style="Heading 2", keep=True, page_break=True)
    add_picture_before(chapter_seven, UNSUPERVISED_IMAGE, 6.35, "Figure 8. Current label-free diagnostics on the 50,000-row unsupervised snapshot")
    add_table_before(document, chapter_seven,
        ["Detector", "EM-AUC", "Score skewness", "Anomaly rate", "Training time"],
        [
            ["Isolation Forest", "67.2 / 100", "77.5 / 100", "1.00%", "1,309 ms"],
            ["Autoencoder", "85.6 / 100", "100.0 / 100", "2.50%", "3,194 ms"],
            ["Behavioral Cluster Outlier", "89.2 / 100", "100.0 / 100", "0.82%", "5,595 ms"],
        ],
        widths=[2.0, 1.0, 1.2, 1.0, 1.2],
    )
    add_paragraph_before(
        chapter_seven,
        "Behavioral Cluster Outlier currently has the strongest EM-AUC and the lowest anomaly rate, while Isolation Forest is fastest. Autoencoder also provides strong score separation. These measurements compare anomaly-score structure; they do not establish precision, recall, or confirmed fraud detection.",
    )

    add_paragraph_before(chapter_seven, "6.4 Current Reproducible Supervised Benchmark", style="Heading 2", keep=True, page_break=True)
    add_paragraph_before(
        chapter_seven,
        "On the 264,500-row Fraud Detection Handbook subset, Extra Trees produced the strongest balanced operating point on the newest untouched 20%. The threshold was chosen only on the calibration period to maximize recall while maintaining at least 70% calibration precision.",
    )
    add_table_before(document, chapter_seven,
        ["Metric", "Extra Trees result", "Interpretation"],
        [
            ["PR-AUC", "73.3%", "Strong ranking quality under 0.8847% fraud prevalence"],
            ["Precision", "74.8%", "323 of 432 fraud alerts were true fraud"],
            ["Recall", "71.8%", "323 of 450 fraud transactions were captured"],
            ["F1", "73.2%", "Balanced precision-recall operating point"],
            ["Confusion matrix", "TP 323; FP 109; FN 127; TN 52,341", "Untouched newest 52,900 transactions"],
        ],
        widths=[1.25, 1.85, 3.1],
    )
    add_paragraph_before(
        chapter_seven,
        "Random Forest reached 74.9% precision and 66.4% recall at its precision-constrained threshold. XGBoost reached 81.9% precision but only 48.2% recall. The result demonstrates why the project compares multiple algorithms and selects an operating point instead of manipulating a fixed 90% precision target.",
    )

    add_paragraph_before(chapter_seven, "6.5 Measured Fusion and Full-Policy Replay", style="Heading 2", keep=True, page_break=True)
    if supervised_image.exists():
        add_picture_before(chapter_seven, supervised_image, 6.4, "Figure 9. Measured standalone and fusion results on the 25,000-row integration snapshot")
    add_paragraph_before(
        chapter_seven,
        "On the same 25,000 held-out integration rows, fusion did not materially outperform the strongest standalone classifier. The weighted probability ensemble produced 58.7% precision, 65.9% recall, and 62.0% F1. Temporal stacking remains a research comparison but has been removed from the production-model selection UI; production uses weighted base classifiers only.",
    )
    if policy_image.exists():
        add_picture_before(chapter_seven, policy_image, 6.4, "Figure 10. Measured four-layer risk-policy replay on the same held-out integration rows")
    add_table_before(document, chapter_seven,
        ["Measured system decision", "Precision", "Recall", "F1 / balanced", "Operational volume"],
        [
            ["Create case (MEDIUM/HIGH)", "54.6%", "50.2%", "52.3% F1; 74.1% balanced accuracy", "43.6 cases per 1,000"],
            ["Generate STR (HIGH)", "66.3%", "25.7%", "62.5% balanced accuracy", "92 HIGH decisions"],
        ],
        widths=[1.85, 0.85, 0.85, 1.55, 1.4],
    )
    add_paragraph_before(
        chapter_seven,
        "The frozen policy used 25% customer behaviour, 10% peer behaviour, 60% ML ensemble, and 5% AML rules. Its lower case-level result shows that adding layers does not automatically improve classification. Customer, peer, and rule layers serve explainability and AML controls, but their weights and thresholds still require calibration on a separate validation period. This limitation is reported rather than hidden.",
    )

    add_paragraph_before(chapter_seven, "6.6 Explainable Transaction Scenario", style="Heading 2", keep=True)
    add_paragraph_before(
        chapter_seven,
        "In a tested scenario, the selected ML detectors flagged an anomaly while customer, peer, and rules supplied separate reasons such as amount deviation, peer-frequency deviation, and turnover thresholds. The UI shows model diagnostics, normalized component scores, reasons, the final 0-100 score, and the policy version. MEDIUM creates a review case; HIGH also prepares a draft STR. This keeps a suspicious decision explainable and reviewable rather than claiming autonomous confirmation of fraud.",
    )

    add_paragraph_before(chapter_seven, "6.7 Automated Validation and Scalability", style="Heading 2", keep=True)
    add_bullet_before(chapter_seven, "Angular and Spring Boot production builds pass; the current Angular bundle reports only existing size-budget warnings.")
    add_bullet_before(chapter_seven, "The 264,500-row model pipeline avoids the earlier 19.8 GiB one-hot allocation by using compact bounded feature representations.")
    add_bullet_before(chapter_seven, "Training snapshots are immutable and checksummed; model artifacts are versioned and do not become active automatically.")
    add_bullet_before(chapter_seven, "Online scoring reuses loaded artifacts and bounded point-in-time features; case and STR work occurs after the transaction decision path.")

    add_paragraph_before(chapter_seven, "6.8 Scientific Interpretation Limits", style="Heading 2", keep=True)
    add_paragraph_before(
        chapter_seven,
        "The presentation-only cache is never treated as measured evidence. Unsupervised EM-AUC is not fraud accuracy. The 25,000-row integration snapshot is synthetic demonstration data. The 264,500-row Fraud Detection Handbook benchmark is reproducible simulator evidence, not a claim about a bank's live population. A production claim requires institution-specific labels, temporal drift monitoring, calibration, investigator-capacity metrics such as precision at k, fairness checks, and independent validation.",
    )


def add_references(document):
    target = find_paragraph(document, "APPENDIX A: MAIN USER WORKFLOW")
    references = [
        "[12] Y.-A. Le Borgne, W. Siblini, B. Lebichot, and G. Bontempi, Reproducible Machine Learning for Credit Card Fraud Detection - Practical Handbook, Universite Libre de Bruxelles, 2022. https://github.com/Fraud-Detection-Handbook/fraud-detection-handbook",
        "[13] P. Geurts, D. Ernst, and L. Wehenkel, 'Extremely Randomized Trees,' Machine Learning, vol. 63, no. 1, pp. 3-42, 2006. doi:10.1007/s10994-006-6226-1.",
        "[14] N. Goix, A. Sabourin, and S. Clemencon, 'On Anomaly Ranking and Excess-Mass Curves,' PMLR, vol. 38, pp. 287-295, 2015.",
        "[15] N. Goix, 'How to Evaluate the Quality of Unsupervised Anomaly Detection Algorithms?' ICML Workshop on Anomaly Detection, 2016.",
        "[16] V. Bach Nguyen, K. Ghosh Dastidar, M. Granitzer, and W. Siblini, 'The Importance of Future Information in Credit Card Fraud Detection,' PMLR, vol. 151, pp. 10067-10077, 2022.",
    ]
    for reference in references:
        paragraph = add_paragraph_before(target, reference)
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        paragraph.paragraph_format.space_after = Pt(4)


def main():
    supervised_image, policy_image = copy_result_assets()
    document = Document(SOURCE)
    update_existing_content(document)
    add_research_matrix(document)
    rebuild_results(document, supervised_image, policy_image)
    add_references(document)

    properties = document.core_properties
    properties.title = "Intelligent Fraud Transaction Detection and AML Case Management - Updated Results"
    properties.subject = "Dual-mode fraud detection, journal-guided design, dataset provenance, and current measured results"
    properties.keywords = "fraud detection, AML, supervised learning, unsupervised learning, Fraud Detection Handbook"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
