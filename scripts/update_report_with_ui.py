from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "deliverables" / "Intelligent_Fraud_Transaction_Detection_and_AML_Case_Management_Report.docx"
OUTPUT = ROOT / "docs" / "deliverables" / "Intelligent_Fraud_Transaction_Detection_and_AML_Case_Management_Report_with_UI_Walkthrough.docx"
ASSETS = ROOT / "docs" / "assets" / "ui-report"


PAGES = [
    (
        "5.9.1 Login and Access Control",
        "01-login.png",
        "Login page used to establish an authenticated application session.",
        "The page provides a focused entry point without exposing operational navigation before authentication.",
        [
            "Credentials are submitted to the Spring Boot authentication endpoint and a signed JWT is returned after successful verification.",
            "The Angular interceptor attaches the token to protected API requests, while backend role checks restrict administrative operations.",
        ],
    ),
    (
        "5.9.2 System Settings and Cold Start",
        "02-settings.png",
        "System-wide learning mode and cold-start configuration.",
        "Settings controls whether the application operates in supervised or unsupervised mode and defines how accounts with insufficient history are treated.",
        [
            "Supervised mode requires FraudLabel values and activates XGBoost, Random Forest, and Logistic Regression workflows.",
            "Cold-start handling prevents unstable customer-history features from being treated as reliable before the configured minimum history is reached.",
        ],
    ),
    (
        "5.9.3 Transaction Dataset Upload",
        "11-upload.png",
        "Mode-aware transaction upload page.",
        "The upload page communicates the active system type and validates whether labelled or unlabelled input is expected before ingestion begins.",
        [
            "CSV and Excel files are accepted, normalized, validated, and inserted through the backend upload pipeline.",
            "In supervised mode FraudLabel 1 means confirmed suspicious/fraud, FraudLabel 0 means legitimate, and blank labels are excluded from supervised learning.",
        ],
    ),
    (
        "5.9.4 Training Operations",
        "07-training-operations.png",
        "Controlled model-training workflow and recent pipeline activity.",
        "An operator chooses the historical period and enabled models. The backend creates one immutable feature snapshot so all selected models receive equivalent data.",
        [
            "Uploaded business dates are sealed automatically, while live API dates can be closed explicitly before training.",
            "Candidate rows show the number of labelled records actually learned rather than the larger mixed snapshot row count.",
        ],
    ),
    (
        "5.9.5 Model Tuning",
        "04-model-tuning.png",
        "Supervised model parameters and shared evaluation protocol.",
        "The page exposes only controls consumed by training or comparison, including chronological holdout settings, reproducibility seed, and model-specific hyperparameters.",
        [
            "Changes apply to subsequent runs and do not silently rewrite already registered model artifacts.",
            "Training enable switches determine which models are available in Training Operations.",
        ],
    ),
    (
        "5.9.6 Risk Decision Policy",
        "03-risk-policy.png",
        "Layered risk weighting and operational decision thresholds.",
        "The policy combines customer behaviour, peer behaviour, the selected ML ensemble, and deterministic AML rules into a normalized score from 0 to 100.",
        [
            "Each component can consume its remaining percentage without unexpectedly changing unrelated controls during manual adjustment.",
            "The medium threshold creates a review case, while the high threshold can trigger automatic STR preparation according to policy.",
        ],
    ),
    (
        "5.9.7 Unsupervised Current-Model Comparison",
        "08-unsupervised-models.png",
        "Current Isolation Forest, Autoencoder, and Local Outlier Factor diagnostics.",
        "This comparison is used when confirmed fraud labels are unavailable. It reports label-free distribution diagnostics rather than claiming fraud accuracy.",
        [
            "EM-AUC and score-skewness values help compare score separation across detector types.",
            "Anomaly rate, alert count, training duration, and learned rows expose operational cost and alert-volume implications.",
        ],
    ),
    (
        "5.9.8 Unsupervised Model Agreement",
        "09-model-agreement.png",
        "Pairwise overlap and consensus among unsupervised detectors.",
        "The agreement study determines whether models flag the same records or contribute complementary anomaly evidence.",
        [
            "Jaccard overlap measures shared flagged sets without requiring class labels.",
            "Exactly-one, exactly-two, and unanimous counts help configure ensemble voting and understand alert concentration.",
        ],
    ),
    (
        "5.9.9 Unsupervised Growth Analysis",
        "10-unsupervised-growth.png",
        "Detector behaviour as chronological history grows.",
        "The oldest 10%, 25%, 50%, and 100% partitions are evaluated to show whether detector behaviour stabilizes as more history becomes available.",
        [
            "Stored growth studies prevent expensive multi-model analysis from rerunning whenever the page is opened.",
            "Trend lines make changes in separation, anomaly rate, and training cost visible to non-technical reviewers.",
        ],
    ),
    (
        "5.9.10 Supervised Model Ranking",
        "12-supervised-ranking.png",
        "Labelled fraud-classification ranking across growth stages.",
        "The ranking compares PR-AUC, precision, recall, F1, calibration, and throughput using chronological evaluation rather than random future-data leakage.",
        [
            "The latest implementation also reports accuracy, balanced accuracy, PR-AUC lift, calibrated decision thresholds, and a confusion matrix.",
            "Results are cached per immutable snapshot, so reopening the page reads stored evidence instead of retraining twelve model-partition combinations.",
        ],
    ),
    (
        "5.9.11 Supervised Growth Matrix",
        "13-supervised-growth.png",
        "Supervised performance at every labelled-data partition.",
        "The matrix reveals whether model quality improves, declines, or remains stable as labelled history expands.",
        [
            "PR-AUC is emphasized because fraud is rare; plain accuracy alone can appear high even when minority-class detection is weak.",
            "On the 25,000-row demonstration snapshot, Random Forest achieved approximately 95.3% accuracy, 70.5% balanced accuracy, and 8.9 times PR-AUC lift over the fraud prevalence baseline.",
        ],
    ),
    (
        "5.9.12 Manual Case Creation",
        "05-manual-case.png",
        "On-demand case creation from an existing transaction.",
        "Investigators can search transaction or account identifiers, inspect the selected record, set priority and assignment, and create a case without waiting for automatic scoring.",
        [
            "The workflow supports intelligence-led review and external referrals that are not initiated by the ML ensemble.",
            "The selected transaction remains the auditable source record for the manually opened case.",
        ],
    ),
    (
        "5.9.13 Case Management",
        "06-case-management.png",
        "Investigation queue, case evidence, and review actions.",
        "Automatically generated and manual cases are listed in a paginated review queue with status, priority, transaction, account, assignment, and decision evidence.",
        [
            "Investigators can mark a case as a false positive or generate an STR when evidence supports escalation.",
            "Human-readable reason descriptions explain why the transaction was considered unusual instead of exposing only technical reason codes.",
        ],
    ),
]


def set_repeatable_font(run, name: str, size: float, color: str | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade_paragraph(paragraph, fill: str):
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    spacing = properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        properties.append(spacing)
    spacing.set(qn("w:before"), "100")
    spacing.set(qn("w:after"), "100")


def add_picture_before(target, image_path: Path, width: float, alt_text: str):
    paragraph = target.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    return paragraph


def main():
    document = Document(SOURCE)
    testing_heading = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "TESTING AND RESULTS"
    )

    heading = testing_heading.insert_paragraph_before("5.9 Application Page Walkthrough", style="Heading 2")
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True

    intro = testing_heading.insert_paragraph_before(
        "This section presents the implemented Angular pages in operational order. Each page is connected to backend controls, persisted evidence, or investigator actions described in the preceding sections."
    )
    intro.paragraph_format.space_after = Pt(8)

    note = testing_heading.insert_paragraph_before()
    shade_paragraph(note, "EAF2FF")
    run = note.add_run("Interface evidence. ")
    set_repeatable_font(run, "Aptos", 10, "174EA6", True)
    run = note.add_run(
        "Screenshots document the working prototype as tested on 12 August 2026. Numerical results are demonstration evidence from synthetic or label-free datasets and are not claims of bank-production fraud accuracy."
    )
    set_repeatable_font(run, "Aptos", 10, "22324A")

    for figure_number, (title, filename, caption, purpose, bullets) in enumerate(PAGES, start=1):
        page_heading = testing_heading.insert_paragraph_before(title, style="Heading 3")
        page_heading.paragraph_format.page_break_before = True
        page_heading.paragraph_format.keep_with_next = True

        purpose_paragraph = testing_heading.insert_paragraph_before(purpose)
        purpose_paragraph.paragraph_format.space_after = Pt(7)

        image = ASSETS / filename
        width = 4.65 if filename == "01-login.png" else 5.0 if filename == "04-model-tuning.png" else 6.25
        add_picture_before(testing_heading, image, width, f"{title}: {caption}")

        caption_paragraph = testing_heading.insert_paragraph_before()
        caption_paragraph.style = "Caption"
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.keep_with_next = True
        caption_paragraph.paragraph_format.space_after = Pt(7)
        caption_run = caption_paragraph.add_run(f"Figure 5.{figure_number}. {caption}")
        set_repeatable_font(caption_run, "Aptos", 9, "53657D", False)
        caption_run.italic = True

        for bullet in bullets:
            paragraph = testing_heading.insert_paragraph_before(bullet, style="List Bullet1")
            paragraph.paragraph_format.space_after = Pt(3)

    testing_heading.paragraph_format.page_break_before = True
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
