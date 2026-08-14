from __future__ import annotations

import math
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\personal\Final Year Project\fraud-transaction-predictor")
TMP = ROOT / "tmp" / "project-report"
OUT = ROOT / "docs" / "deliverables"
REFERENCE = Path(r"C:\Users\Administrator\Downloads\Evening Program Project Report Template.docx")
FINAL = OUT / "Intelligent_Fraud_Transaction_Detection_and_AML_Case_Management_Report.docx"
FIG_DIR = TMP / "figures"

SCREENSHOTS = {
    "comparison": Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-cbcc5c2e-70a4-4a2b-b3a0-5ea7704b4b1f.png"),
    "risk_policy": Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-9fd0571f-e173-4fd9-ab97-295f128e46a1.png"),
    "transaction": Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-3c067d1e-e4e3-4872-9cc6-3f149cb71686.png"),
    "decision": Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-a928816c-3f75-43c8-b529-e3a3093aee9b.png"),
    "components": Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-d260ff06-20d5-408c-a192-4a25054c2e27.png"),
    "payload": Path(r"C:\Users\ADMINI~1\AppData\Local\Temp\codex-clipboard-7ddb8515-0b2b-4e17-863c-8b2549d2a88d.png"),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=110, bottom=110, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def unlink_headers_footers(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False


def clear_header_footer(section, page_number=True) -> None:
    section.header.paragraphs[0].clear()
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_number:
        add_page_field(footer, "PAGE")
        for run in footer.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def configure_section(section, cover=False) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(2.0 if cover else 1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.5)
    unlink_headers_footers(section)
    clear_header_footer(section, page_number=not cover)


def font_run(run, size=12, bold=False, italic=False, color="000000") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, bold=False, italic=False, size=12) -> None:
    run = paragraph.add_run(text)
    font_run(run, size=size, bold=bold, italic=italic)


def add_body(doc, text: str, first_line=True, keep=False) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.3) if first_line else None
    p.paragraph_format.keep_together = keep
    add_text(p, text)


def ensure_numbering(doc: Document, kind: str, fresh=False) -> int:
    cache = getattr(doc, "_report_numbering", {})
    if kind in cache and not fresh:
        return cache[kind]
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(2):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
        num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal"); lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else f"%{level + 1}.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab"); lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), str(720 + level * 360)); tabs.append(tab); p_pr.append(tabs)
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), str(720 + level * 360)); ind.set(qn("w:hanging"), "360"); p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id)); num.append(abstract_ref)
    numbering.append(num)
    cache[kind] = num_id
    setattr(doc, "_report_numbering", cache)
    return num_id


def apply_numbering(doc: Document, paragraph, kind: str, level: int = 0, num_id: int | None = None) -> None:
    num_id = num_id if num_id is not None else ensure_numbering(doc, kind)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(level)); num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId"); num_id_node.set(qn("w:val"), str(num_id)); num_pr.append(num_id_node)


def add_bullets(doc, items: list[str], level=0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        apply_numbering(doc, p, "bullet", level)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(3)
        add_text(p, item)


def add_numbered(doc, items: list[str]) -> None:
    num_id = ensure_numbering(doc, "decimal", fresh=True)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        apply_numbering(doc, p, "decimal", 0, num_id)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(3)
        add_text(p, item)


def add_chapter(doc, number: int, title: str, page_break=True) -> None:
    if page_break and len(doc.paragraphs) > 0:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_text(p, f"CHAPTER {number}", bold=True, size=16)
    p2 = doc.add_paragraph(style="Heading 1")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    p2.paragraph_format.keep_with_next = True
    add_text(p2, title.upper(), bold=True, size=16)


def add_heading(doc, number: str, title: str, level=2) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, f"{number} {title}", bold=True, size=13 if level == 2 else 12)


def add_caption(doc, label: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    add_text(p, label, italic=True, size=10)


def add_figure(doc, path: Path, caption: str, width=5.7) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], caption: str) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, "D9E2F3")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, bold=True, size=9)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, str(value), size=8.5)
    doc.add_paragraph()


def add_toc_line(doc, text: str, page: str, level=0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(5.45))
    add_text(p, text, bold=level == 0, size=10.5)
    add_text(p, "\t" + page, size=10.5)


def create_diagram(filename: str, title: str, boxes: list[tuple[float, float, float, float, str, str]], arrows: list[tuple[int, int, str]]) -> Path:
    from PIL import Image, ImageDraw, ImageFont

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path(r"C:\Windows\Fonts\times.ttf")
    bold_path = Path(r"C:\Windows\Fonts\timesbd.ttf")
    font = ImageFont.truetype(str(font_path), 28)
    small = ImageFont.truetype(str(font_path), 23)
    bold = ImageFont.truetype(str(bold_path), 36)
    draw.text((width / 2, 42), title, fill="#111111", font=bold, anchor="mm")
    positions = []
    for x, y, w, h, text, color in boxes:
        px = int(x * width)
        py = int(y * height)
        pw = int(w * width)
        ph = int(h * height)
        positions.append((px, py, pw, ph))
        draw.rounded_rectangle((px, py, px + pw, py + ph), radius=20, fill=color, outline="#244062", width=3)
        words = text.split()
        lines, current = [], ""
        for word in words:
            candidate = (current + " " + word).strip()
            if draw.textlength(candidate, font=small) > pw - 34 and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        total = len(lines) * 29
        for li, line in enumerate(lines):
            draw.text((px + pw / 2, py + ph / 2 - total / 2 + li * 29), line, fill="#111111", font=small, anchor="ma")
    for source, target, label in arrows:
        sx, sy, sw, sh = positions[source]
        tx, ty, tw, th = positions[target]
        start = (sx + sw, sy + sh / 2) if tx > sx else (sx + sw / 2, sy + sh)
        end = (tx, ty + th / 2) if tx > sx else (tx + tw / 2, ty)
        draw.line((start, end), fill="#244062", width=4)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        size = 16
        p1 = (end[0] - size * math.cos(angle - 0.55), end[1] - size * math.sin(angle - 0.55))
        p2 = (end[0] - size * math.cos(angle + 0.55), end[1] - size * math.sin(angle + 0.55))
        draw.polygon((end, p1, p2), fill="#244062")
        if label:
            mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
            draw.rectangle((mx - 70, my - 17, mx + 70, my + 17), fill="white")
            draw.text((mx, my), label, fill="#244062", font=font, anchor="mm")
    path = FIG_DIR / filename
    image.save(path)
    return path


def build_figures() -> dict[str, Path]:
    architecture = create_diagram(
        "architecture.png",
        "Three-Application Architecture",
        [
            (0.04, 0.30, 0.20, 0.25, "Angular User Interface", "#EAF2F8"),
            (0.31, 0.20, 0.25, 0.45, "Spring Boot Business API Security, AML Workflow, Cases and Governance", "#D9EAD3"),
            (0.64, 0.14, 0.25, 0.24, "Python FastAPI ML Training and Prediction", "#FCE5CD"),
            (0.64, 0.53, 0.25, 0.24, "Microsoft SQL Server Auditable System of Record", "#FFF2CC"),
        ],
        [(0, 1, "REST"), (1, 2, "ML API"), (1, 3, "JPA/SQL")],
    )
    scoring = create_diagram(
        "layered_scoring.png",
        "Layered Risk Scoring (All Weights Normalize to 100%)",
        [
            (0.03, 0.24, 0.19, 0.22, "Customer Behaviour", "#D9EAF7"),
            (0.27, 0.24, 0.19, 0.22, "Peer Behaviour", "#EADCF8"),
            (0.51, 0.24, 0.19, 0.22, "ML Ensemble", "#D9EAD3"),
            (0.75, 0.24, 0.19, 0.22, "AML Rules", "#FCE5CD"),
            (0.37, 0.62, 0.27, 0.18, "Final Risk Score and Risk Level", "#FFF2CC"),
        ],
        [(0, 4, "weighted"), (1, 4, "weighted"), (2, 4, "weighted"), (3, 4, "weighted")],
    )
    training = create_diagram(
        "training_pipeline.png",
        "Governed Training and Comparison Pipeline",
        [
            (0.02, 0.27, 0.15, 0.22, "Close Business Dates", "#EAF2F8"),
            (0.21, 0.27, 0.15, 0.22, "Create Training Run", "#D9EAD3"),
            (0.40, 0.27, 0.15, 0.22, "Export Checksummed Parquet Snapshot", "#FFF2CC"),
            (0.59, 0.27, 0.15, 0.22, "Train Selected Models", "#FCE5CD"),
            (0.78, 0.27, 0.18, 0.22, "Register Candidate and Compare", "#EADCF8"),
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, "")],
    )
    transaction = create_diagram(
        "transaction_flow.png",
        "Point-in-Time Transaction Decision Flow",
        [
            (0.02, 0.27, 0.14, 0.22, "Receive and Persist Transaction", "#EAF2F8"),
            (0.20, 0.27, 0.14, 0.22, "Load Prior History Only", "#D9EAD3"),
            (0.38, 0.27, 0.14, 0.22, "Persist Feature Vector", "#FFF2CC"),
            (0.56, 0.27, 0.14, 0.22, "Layered Risk Scoring", "#FCE5CD"),
            (0.74, 0.13, 0.22, 0.20, "MEDIUM: Create Case", "#EADCF8"),
            (0.74, 0.51, 0.22, 0.20, "HIGH: Create Case and Draft STR XML", "#F4CCCC"),
        ],
        [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (3, 5, "")],
    )
    return {"architecture": architecture, "scoring": scoring, "training": training, "transaction": transaction}


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    for name in ("Normal", "Default Paragraph Font"):
        if name in styles:
            styles[name].font.name = "Times New Roman"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            styles[name].font.size = Pt(12)
    normal = styles["Normal"]
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for idx, size in ((1, 16), (2, 13), (3, 12)):
        style = styles[f"Heading {idx}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        if list_name not in styles:
            style = styles.add_style(list_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
        style = styles[list_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.35)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.18)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.65)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.18)
    styles["List Number"].paragraph_format.left_indent = Inches(0.35)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.18)
    if "Caption" in styles:
        styles["Caption"].font.name = "Times New Roman"
        styles["Caption"].font.size = Pt(10)


def add_front_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    add_text(p, title.upper(), bold=True, size=14)


def add_front_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.first_line_indent = Inches(0.25)
    add_text(p, text, size=10.5)


def build_report() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    figures = build_figures()
    doc = Document(REFERENCE)
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    while len(doc.sections) > 1:
        break
    configure_styles(doc)
    first = doc.sections[0]
    configure_section(first, cover=True)
    first.different_first_page_header_footer = True
    first.first_page_footer.paragraphs[0].clear()

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "INTELLIGENT FRAUD TRANSACTION DETECTION AND AML CASE MANAGEMENT SYSTEM", bold=True, size=18)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "A Project", bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "Presented to the")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "University of Dhaka, Bangladesh", bold=True)
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "In Partial Fulfillment")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "of the Requirements for the Degree")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "Master in Information Technology (MIT)", bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "or")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "Post Graduate Diploma in Information Technology (PGDIT)", bold=True)
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "By", bold=True)
    for value in ("<Student Name>", "<Roll / Registration Number>", "<Academic Year>"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, value)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, "August 2026", bold=True)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)

    # Front matter section with Roman numbering.
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front)
    set_page_number_format(front, "lowerRoman", 1)
    add_front_heading(doc, "Signature Page")
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.75, 4.25])
    signature_rows = [
        ("PROJECT:", "Intelligent Fraud Transaction Detection and AML Case Management System"),
        ("AUTHOR:", "<Student Name>"),
        ("ROLL:", "<Roll / Registration Number>"),
        ("DATE SUBMITTED:", "August 2026"),
        ("SUPERVISED BY:", "<Supervisor Name>"),
        ("DESIGNATION:", "<Supervisor Designation>"),
        ("INSTITUTION:", "University of Dhaka, Bangladesh"),
        ("SUPERVISOR'S APPROVAL:", "Signature and Date: ______________________________"),
    ]
    for row, values in zip(table.rows, signature_rows):
        for idx, value in enumerate(values):
            set_cell_margins(row.cells[idx], 130, 130, 130, 130)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, bold=idx == 0, size=10.5)
    doc.add_page_break()
    add_front_heading(doc, "Acknowledgements")
    add_body(doc, "I express sincere gratitude to my project supervisor for continuous guidance, technical criticism, and encouragement throughout this work. I also acknowledge the faculty members and classmates whose feedback helped refine the scope from a basic anomaly detector into an auditable AML platform. Finally, I am grateful to my family for their patience and support during the design, implementation, testing, and documentation of this project.")
    add_body(doc, "This project benefited from open-source technologies and publicly documented machine-learning methods. Their communities made it possible to construct a complete prototype spanning transaction ingestion, feature engineering, model training, model comparison, risk policy configuration, case management, and suspicious transaction report generation.")
    doc.add_page_break()
    add_front_heading(doc, "Declaration")
    add_body(doc, "I declare that this project report is an original account of the work completed for the Intelligent Fraud Transaction Detection and AML Case Management System. Material derived from external publications and software documentation is acknowledged in the references. The system is an academic prototype and does not claim that an anomaly is legally proven fraud; suspicious outcomes require authorized human review.")
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); add_text(p, "Student's Name: <Student Name>", bold=True)
    p = doc.add_paragraph(); add_text(p, "Student's Signature: ______________________________", bold=True)
    doc.add_page_break()
    add_front_heading(doc, "Abstract")
    add_front_body(doc, "Financial institutions process large volumes of transactions while needing to identify unusual activity quickly, explain decisions to analysts, and preserve evidence for regulatory review. A single machine-learning model is insufficient because institutions may have no confirmed fraud labels, transaction behaviour changes as history grows, customer-specific patterns differ from peer behaviour, and model predictions alone do not represent deterministic anti-money-laundering obligations.")
    add_front_body(doc, "This project implements an Intelligent Fraud Transaction Detection and AML Case Management System using Angular, Spring Boot, FastAPI, and Microsoft SQL Server. A global setting selects unsupervised or supervised learning. Unsupervised mode uses Isolation Forest, Autoencoder, and Local Outlier Factor; supervised mode uses XGBoost, Random Forest, and Logistic Regression. Both modes feed a layered risk score combining customer behaviour, peer behaviour, the selected ML ensemble, and deterministic AML rules. Weights normalize to 100%, MEDIUM and HIGH outcomes create cases, and HIGH outcomes generate draft STR XML after transaction commit.")
    add_front_body(doc, "The main research feature evaluates how model behaviour changes when the oldest 10%, 25%, 50%, and 100% of an immutable snapshot are used. Label-free diagnostics are reported for unsupervised models, while labelled chronological evaluation is used for supervised models. Scalability is addressed through persisted point-in-time features, closed business dates, keyset export, checksummed Parquet snapshots, bounded evaluation, asynchronous case work, pagination, and versioned artifacts. The result separates anomaly evidence from confirmed fraud and supports analyst review, false-positive feedback, manual cases, and STR preparation.")
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0); add_text(p, "Keywords: ", bold=True, size=10); add_text(p, "fraud detection, anomaly detection, anti-money laundering, supervised learning, unsupervised learning, model comparison, explainable risk scoring, case management, STR XML, scalable architecture", size=10)

    # Contents
    doc.add_page_break(); add_front_heading(doc, "Table of Contents")
    toc = [
        ("Signature Page", "i", 0), ("Acknowledgements", "ii", 0), ("Declaration", "iii", 0), ("Abstract", "iv", 0),
        ("Table of Contents", "v", 0), ("List of Tables", "vi", 0), ("List of Figures", "vii", 0),
        ("Chapter 1: Introduction", "1", 0), ("1.1 Background", "1", 1), ("1.2 Motivation", "1", 1), ("1.3 Problem Statement", "1", 1), ("1.4 Objectives", "2", 1),
        ("Chapter 2: Background and Related Approaches", "4", 0), ("Chapter 3: System Analysis and Architecture", "6", 0),
        ("Chapter 4: Methodology", "9", 0), ("Chapter 5: Implementation", "13", 0),
        ("Chapter 6: Testing and Results", "18", 0), ("Chapter 7: Problems and Solutions", "21", 0),
        ("Chapter 8: Conclusion and Future Work", "23", 0), ("References", "25", 0), ("Appendices", "26", 0),
    ]
    for text, page, level in toc:
        add_toc_line(doc, text, page, level)
    doc.add_page_break(); add_front_heading(doc, "List of Tables")
    table_entries = [
        ("Table 1. Functional requirements", "6"), ("Table 2. Non-functional requirements", "6"),
        ("Table 3. Dual learning model catalog", "9"), ("Table 4. Evaluation metrics by learning mode", "12"),
        ("Table 5. Main implementation technologies", "13"), ("Table 6. Current unsupervised model results on 50,000 rows", "18"),
        ("Table 7. Automated validation summary", "19"), ("Table 8. Problem-to-solution matrix", "21"),
    ]
    for text, page in table_entries:
        add_toc_line(doc, text, page, 0)
    doc.add_page_break(); add_front_heading(doc, "List of Figures")
    figure_entries = [
        ("Figure 1. Three-application architecture", "7"), ("Figure 2. Point-in-time transaction decision flow", "8"),
        ("Figure 3. Layered risk scoring", "11"), ("Figure 4. Governed training and comparison pipeline", "12"),
        ("Figure 5. Risk policy configuration interface", "14"), ("Figure 6. Transaction entry interface", "15"),
        ("Figure 7. Explainable transaction decision", "16"), ("Figure 8. Model comparison results", "18"),
    ]
    for text, page in figure_entries:
        add_toc_line(doc, text, page, 0)

    # Body section Arabic numbering.
    main = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main)
    set_page_number_format(main, "decimal", 1)

    # Chapter 1
    add_chapter(doc, 1, "Introduction", page_break=False)
    add_heading(doc, "1.1", "Background")
    add_body(doc, "Fraud monitoring and anti-money-laundering operations must identify unusual transactions without interrupting legitimate banking activity. Conventional rule engines are transparent but can be rigid and generate large numbers of false alerts. Machine-learning methods can discover complex patterns, but their raw scores are difficult to compare, may drift as the amount of history changes, and can fail silently when feature engineering differs between training and prediction.")
    add_body(doc, "The project began as an anomaly detector for a bank transaction dataset containing transaction identity, account, amount, date, type, location, channel, customer age, occupation, login attempts, balance, and previous transaction date. It evolved into a broader AML decision-support system that persists point-in-time features, compares models on growing chronological datasets, supports labelled and unlabelled learning, explains layered risk contributions, and manages the analyst workflow after a suspicious decision.")
    add_heading(doc, "1.2", "Motivation")
    add_body(doc, "A bank may store millions of transactions and cannot repeatedly send its entire lifetime history through one synchronous request. It may also lack reliable fraud labels because alerts are not equivalent to confirmed fraud. Even when a model detects an anomaly, analysts need to know whether the amount was unusual for the customer, unusual for comparable peers, associated with velocity or structuring rules, or simply outside a model boundary. These operational and scientific limitations motivated a system that separates evidence layers and preserves an audit trail.")
    add_heading(doc, "1.3", "Problem Statement")
    add_bullets(doc, [
        "Transaction datasets often have no trustworthy class label, preventing valid accuracy, precision, recall, and F1 claims.",
        "A model that looks good on a small sample may behave differently when historical data grows.",
        "Raw scores from different algorithms use incompatible scales and cannot be added directly.",
        "Scoring one transaction can become slow if the system scans large historical tables or retrains models synchronously.",
        "Customer behaviour, peer behaviour, AML rules, and ML evidence can disagree and require a governed method of combination.",
        "An anomaly alert alone is not a complete operational outcome; investigators need cases, explanations, notes, false-positive handling, and STR preparation.",
        "Administrative functions require authentication, authorization, model versioning, and reproducible training data.",
    ])
    add_heading(doc, "1.4", "Project Objectives")
    add_numbered(doc, [
        "Build an end-to-end web platform for upload, transaction checking, training, model comparison, configuration, case review, and STR XML generation.",
        "Support unsupervised and supervised learning through one system-level mode while keeping their models and evaluation metrics scientifically separate.",
        "Measure how model behaviour changes using the oldest 10%, 25%, 50%, and 100% of an immutable dataset.",
        "Combine transparent customer behaviour, peer behaviour, ML evidence, and deterministic AML rules into a normalized 0-100 risk score.",
        "Make suspicious decisions explainable through reason codes, component scores, model diagnostics, and persisted evidence.",
        "Design training, data export, and real-time decision paths that can evolve toward bank-scale workloads.",
    ])
    add_heading(doc, "1.5", "Scope")
    add_body(doc, "The implemented scope includes CSV and Excel ingestion, real-time transaction creation, persisted feature vectors, cold-start settings, global learning mode, six supported model choices across two learning modes, mode-aware tuning, chronological model comparison, weighted risk policy configuration, JWT authentication, role authorization, automatic and manual cases, pagination, analyst notes, false-positive marking, and draft STR XML. The project is an academic prototype: it does not connect to a production core-banking system, submit reports to a regulator, or replace human adjudication.")
    add_heading(doc, "1.6", "Significance")
    add_body(doc, "The primary contribution is not a claim that one algorithm proves fraud. The contribution is an integrated framework for selecting and governing evidence: it tells users which methodology is valid for their data, how models behave as history grows, how non-model rules affect the final result, and why a case was generated. This makes the prototype useful for academic evaluation and for discussing production AML design constraints.")

    # Chapter 2
    add_chapter(doc, 2, "Background and Related Approaches")
    add_heading(doc, "2.1", "Rule-Based Transaction Monitoring")
    add_body(doc, "Traditional monitoring uses deterministic thresholds for structuring, velocity, unusual amount, sanctions, location, and channel behaviour. Rules are easy to explain and can implement mandatory controls, but static thresholds may not reflect individual customer baselines. This project therefore retains rules as one explicit layer rather than asking ML models to rediscover every legal or operational control.")
    add_heading(doc, "2.2", "Unsupervised Anomaly Detection")
    add_body(doc, "When labels are unavailable, the system uses three complementary detectors. Isolation Forest isolates rare observations through randomized trees; the Autoencoder learns to reconstruct common feature patterns and treats high reconstruction error as unusual; Local Outlier Factor compares the local density around one transaction with the density around its neighbours. Their outputs indicate abnormality rather than confirmed fraud.")
    add_heading(doc, "2.3", "Supervised Fraud Classification")
    add_body(doc, "When an auditable binary label is available, supervised classifiers can learn the relationship between historical features and reviewed outcomes. XGBoost provides a strong nonlinear boosted-tree model, Random Forest provides a robust class-balanced ensemble, and Logistic Regression provides an interpretable probability baseline. The comparison protocol emphasizes precision-recall measures because fraud datasets are normally imbalanced.")
    add_heading(doc, "2.4", "Explainable Layered Risk")
    add_body(doc, "Pure voting creates unintuitive outcomes when three models agree but the ML layer owns only part of the total risk budget. The project replaces this ambiguity with an explicit layered score. Each component produces a normalized value between zero and one; configured top-level weights sum to one; selected ML model weights also sum to one inside the ensemble. The UI shows both anomaly votes for diagnostics and the weighted score that actually determines the risk level.")
    add_heading(doc, "2.5", "Research Gap Addressed")
    add_body(doc, "Many demonstrations train one model on one fixed dataset and report either accuracy without proving label quality or anomaly counts without studying stability. This project adds a chronological growth experiment, separates label-free and label-based metrics, includes transparent non-ML layers, and connects evaluation to a governed operational workflow. It therefore studies both model behaviour and system behaviour.")

    # Chapter 3
    add_chapter(doc, 3, "System Analysis and Architecture")
    add_heading(doc, "3.1", "Stakeholders and Use Cases")
    add_body(doc, "Administrators configure the learning mode, tuning values, risk policy, training runs, and model governance. Reviewers inspect cases, search transactions, add notes, mark false positives, and generate STR XML. External transaction sources upload files or call the transaction API. The ML service trains and scores models but does not own business records or case decisions.")
    add_table(doc, ["Actor", "Major functions"], [
        ["Administrator", "Settings, uploads, training operations, model tuning, comparison, risk policy, governance"],
        ["AML Reviewer", "Transaction checks, case search, notes, false-positive review, manual case creation, STR XML"],
        ["Transaction Source", "Bulk CSV/Excel upload or real-time transaction API"],
        ["ML Service", "Feature alignment, model training, prediction, comparison metrics, artifact persistence"],
    ], [1.4, 4.6], "Table 1. Functional requirements")
    add_heading(doc, "3.2", "Non-Functional Requirements")
    add_table(doc, ["Quality", "Requirement and implementation direction"], [
        ["Security", "Stateless JWT authentication, BCrypt password hashes, server-side role authorization, deny-by-default API policy"],
        ["Explainability", "Reason codes, component scores, model diagnostics, feature summary, policy and model versions"],
        ["Scalability", "Persisted features, Parquet snapshots, bounded evaluation, asynchronous workflows, pagination"],
        ["Integrity", "Closed business dates, chronological cutoffs, checksums, immutable model bundles, duplicate protection"],
        ["Availability", "Auditable fallback when the ML service is unavailable; case creation occurs after commit"],
        ["Maintainability", "Separate Angular, Spring Boot, and FastAPI applications with stable REST contracts"],
    ], [1.35, 4.65], "Table 2. Non-functional requirements")
    add_heading(doc, "3.3", "Overall Architecture")
    add_figure(doc, figures["architecture"], "Figure 1. Three-application architecture")
    add_body(doc, "Angular provides the browser interface. Spring Boot owns transaction persistence, security, feature orchestration, risk policy, model governance, and case workflow. FastAPI owns numerical preprocessing, model fitting, artifact loading, prediction, and research analysis. SQL Server is the auditable system of record. This division prevents the Python service from directly mutating business records and prevents the UI from accessing model artifacts or database tables.")
    add_heading(doc, "3.4", "Database and Audit Design")
    add_body(doc, "Core operational records include transactions, upload batches, persisted transaction features, learning eligibility, prediction logs, fraud alerts, cases, case notes, STR reports, application configuration, training runs, immutable model registry entries, validation reports, deployment pointers, and growth studies. Migration scripts are additive and independently verifiable. Training runs reference closed business windows and immutable exported datasets rather than mutable lifetime queries.")
    add_heading(doc, "3.5", "Transaction Decision Flow")
    add_figure(doc, figures["transaction"], "Figure 2. Point-in-time transaction decision flow")
    add_body(doc, "A transaction is first persisted with processing status. Feature engineering loads only records earlier than the transaction timestamp, preventing future information leakage. The feature vector is persisted before prediction. The layered orchestrator returns the final risk score, model evidence, feature summary, reasons, and action. Suspicious results create an alert and publish an after-commit event. MEDIUM creates a case; HIGH creates a case and automatically generates draft STR XML asynchronously so the transaction database commit is not held open by case work.")
    add_heading(doc, "3.6", "Authentication and Authorization")
    add_body(doc, "Login issues a signed JWT. Passwords are stored using BCrypt with strength 12. Spring Security permits health, login, registration, and API documentation without authentication, requires authentication for all other APIs, and restricts upload, training, AML writes, and comparison writes to ADMIN or AML_ADMIN. Angular route guards improve usability, but the backend remains the authority.")

    # Chapter 4
    add_chapter(doc, 4, "Methodology")
    add_heading(doc, "4.1", "Global Learning Mode")
    add_body(doc, "The Settings page stores system.learning_mode in app_config. UNSUPERVISED is selected when confirmed labels are absent. SUPERVISED is selected only when an auditable FraudLabel is present. Upload validation, the model catalog, Model Tuning, Training Operations, the comparison page, and the Risk Policy model list read this mode so the workflow remains consistent.")
    add_table(doc, ["Mode", "Models", "Required data", "Primary interpretation"], [
        ["Unsupervised", "Isolation Forest; Autoencoder; Local Outlier Factor", "Transactions without confirmed class labels", "Anomaly evidence and stability"],
        ["Supervised", "XGBoost; Random Forest; Logistic Regression", "Reviewed binary FraudLabel", "Estimated suspicious probability and classification quality"],
    ], [1.0, 2.1, 1.55, 1.35], "Table 3. Dual learning model catalog")
    add_heading(doc, "4.2", "Feature Engineering and Leakage Prevention")
    add_body(doc, "Features describe amount and balance ratios, hour and day, night and weekend indicators, customer history count, rolling amount statistics, z-scores, debit and credit ratios, transaction velocity, unique beneficiaries, new location/channel/device indicators, peer-group statistics, expected turnover deviation, customer type, and risk rating. Each feature is computed using records available before the prediction time. The feature schema and column order are persisted with every model bundle so training and prediction cannot silently disagree.")
    add_heading(doc, "4.3", "Customer Behaviour Scoring")
    add_body(doc, "The customer layer asks whether the transaction is unusual for the same customer. A transparent weighted formula combines amount deviation, transaction frequency, time-gap deviation, categorical novelty, and unusual hour. A confidence multiplier reduces customer-specific evidence during cold start. This prevents two or three prior transactions from being treated as a mature profile while retaining limited evidence.")
    add_heading(doc, "4.4", "Peer Behaviour Scoring")
    add_body(doc, "The peer layer compares the transaction with similar customers. Groups are derived from customer attributes such as occupation and age band, not hard-coded for one profession. If a specific occupation-age baseline lacks enough history, the scorer falls back to a broader parent segment and then to a global baseline, with decreasing confidence. The peer window is anchored to the training snapshot cutoff rather than the newest just-arrived transaction, preventing one future transaction from shifting the historical reference window.")
    add_heading(doc, "4.5", "Deterministic AML Rules")
    add_body(doc, "Rules evaluate transparent conditions such as amount multiples, customer and peer z-scores, high transaction velocity, repeated near-threshold activity, expected-turnover deviation, unusual locations, and hard compliance overrides. Rule thresholds are configured separately from model tuning. A hard override remains identifiable as a rule decision rather than being misrepresented as model consensus.")
    add_heading(doc, "4.6", "Layered Weighted Risk Aggregation")
    add_figure(doc, figures["scoring"], "Figure 3. Layered risk scoring")
    add_body(doc, "The final normalized score is customerWeight x customerScore + peerWeight x peerScore + mlWeight x mlScore + rulesWeight x rulesScore. Top-level weights total 1.0. Inside the ML ensemble, enabled model weights also total 1.0. The UI provides component-level normalization so an operator can increase one value to use the remaining percentage without unpredictably changing unrelated sections. Risk thresholds convert the final 0-1 value to a 0-100 display and NORMAL, LOW, MEDIUM, or HIGH.")
    add_heading(doc, "4.7", "Chronological Growth Comparison")
    add_body(doc, "The research experiment orders the immutable snapshot by transaction_date and trains each selected model on the oldest 10%, 25%, 50%, and 100% that meet the minimum-row rule. Each partition uses a chronological holdout, so future rows are not randomly mixed into training. This directly answers how behaviour, stability, anomaly rate, training time, and throughput change as historical evidence increases.")
    add_table(doc, ["Mode", "Metrics", "What they establish"], [
        ["Unsupervised", "EM-AUC, score separation, skewness, anomaly rate, rate control, stability, throughput", "Label-free diagnostic quality; not fraud accuracy"],
        ["Supervised", "PR-AUC, ROC-AUC, precision, recall, F1, Brier score, confusion matrix", "Classification and probability quality on reviewed labels"],
        ["Both", "Training duration, learned rows, evaluation rows, prediction latency", "Operational scalability and reproducibility"],
    ], [1.0, 2.45, 2.55], "Table 4. Evaluation metrics by learning mode")
    add_heading(doc, "4.8", "Governed Training")
    add_figure(doc, figures["training"], "Figure 4. Governed training and comparison pipeline")
    add_body(doc, "An administrator chooses when to train; automatic daily and weekly schedules were removed. Business dates must first be closed. Spring Boot exports eligible persisted features through keyset pagination into checksummed Parquet parts. Python trains only the models selected from the active mode and saves versioned artifacts. Spring registers candidates with metrics and metadata. Training does not silently activate a model.")

    # Chapter 5
    add_chapter(doc, 5, "Implementation")
    add_heading(doc, "5.1", "Tools and Technologies")
    add_table(doc, ["Layer", "Technology", "Purpose"], [
        ["Frontend", "Angular 21, TypeScript, RxJS", "Responsive pages, forms, route guards, API integration, alerts"],
        ["Business backend", "Java 17, Spring Boot 4, Spring Security, JPA", "Transactions, security, policies, cases, training orchestration"],
        ["ML service", "Python, FastAPI, pandas, NumPy, scikit-learn, XGBoost", "Training, prediction, evaluation, artifact handling"],
        ["Database", "Microsoft SQL Server", "Auditable operational and governance state"],
        ["Data interchange", "REST/JSON, multipart upload, Parquet, XML", "Online calls, datasets, and STR draft output"],
    ], [1.0, 2.0, 3.0], "Table 5. Main implementation technologies")
    add_heading(doc, "5.2", "Angular User Interface")
    add_body(doc, "The UI uses a white, responsive dashboard with a toggleable side menu and consistent success/error popup alerts. Current routes include Uploads, Datasets and model comparison, Supervised Comparison, Training Operations, Risk Policy, Settings, Model Tuning, Transaction Check, Case Management, and Manual Case Creation. Administrative pages are route-protected, and every page delegates authority to backend APIs.")
    add_heading(doc, "5.3", "Risk Policy Configuration")
    add_figure(doc, SCREENSHOTS["risk_policy"], "Figure 5. Risk policy configuration interface", width=5.85)
    add_body(doc, "The configuration interface allocates 100% across customer behaviour, peer behaviour, the ML ensemble, and AML rules. A modal selects active production models based on the global learning mode. Disabled models do not participate. The page displays effective model contribution to the total score and prevents totals above 100%.")
    add_heading(doc, "5.4", "Model Tuning")
    add_body(doc, "Model Tuning is mode-aware. In unsupervised mode it exposes controls for Isolation Forest, Autoencoder, and LOF. In supervised mode it exposes XGBoost tree count, depth, learning rate, row and feature sampling; Random Forest tree count, depth, and leaf size; and Logistic Regression regularization and iteration limits. Shared evaluation controls include minimum rows, holdout fraction, maximum evaluation rows, and random seed. Values are persisted in app_config and consumed by both training and growth analysis.")
    add_heading(doc, "5.5", "Transaction Check and Explainability")
    add_figure(doc, SCREENSHOTS["transaction"], "Figure 6. Transaction entry interface", width=4.2)
    add_figure(doc, SCREENSHOTS["decision"], "Figure 7. Explainable transaction decision", width=4.2)
    add_body(doc, "The transaction page accepts a transaction ID, account, amount, balance, type, channel, date, location, age, occupation, and login attempts. The result displays final score, risk level, action, ML anomaly votes, human-readable reason tags, key features, model diagnostics, latency, layered component scores, policy version, and raw payload for technical review. This separation prevents the raw model decision from being confused with the final policy decision.")
    add_heading(doc, "5.6", "Case Management and STR")
    add_body(doc, "Suspicious transactions create fraud alerts and cases after the transaction commit. Case lists use server-side pagination, filtering, and summary counts to avoid loading the full case table. The detail view includes transaction evidence, prediction evidence, analyst notes, reason explanations, and allowed actions. A reviewer can mark a case false positive or generate draft STR XML, but the state machine prevents STR generation from a false-positive case and prevents an STR-generated case from later being marked false positive. Manual cases reuse existing transactions and avoid duplicates.")
    add_heading(doc, "5.7", "Training and Artifact Management")
    add_body(doc, "Training operations expose model checkboxes from the active Risk Policy rather than fixed daily or weekly training types. One run can train selected candidates from the same immutable snapshot. Artifacts include the fitted estimator, scaler or preprocessing state, ordered feature columns, threshold, hyperparameters, metrics, and dataset checksum. Atomic replacement avoids partially written model directories, and model loading caches versioned artifacts for prediction.")
    add_heading(doc, "5.8", "Failure Handling")
    add_body(doc, "If FastAPI is unavailable, Spring records the failure and returns the documented safe fallback instead of corrupting transaction state. Invalid roles receive JSON 401 or 403 responses. Duplicate transaction IDs are rejected. Training failures preserve messages and do not activate candidates. Dataset export rejects open business dates. Artifact checksum and feature-version mismatches block governance actions.")

    # Chapter 6
    add_chapter(doc, 6, "Testing and Results")
    add_heading(doc, "6.1", "Testing Strategy")
    add_body(doc, "Testing combined automated unit and integration tests with manual end-to-end acceptance checks. Java tests cover feature context, customer and peer scoring, rules, layered risk arithmetic, training export, model registry, validation, deployment, case management, authentication, and comparison services. Python tests cover growth analysis, persisted-feature alignment, artifact bundles, route compatibility, finite-value checks, and model artifact reuse. Angular production builds verify template and TypeScript integration.")
    add_heading(doc, "6.2", "Current Unsupervised Training Result")
    add_figure(doc, SCREENSHOTS["comparison"], "Figure 8. Model comparison results for the 50,000-row training snapshot", width=5.9)
    add_table(doc, ["Model", "Rows", "EM-AUC", "Skewness", "Anomaly rate", "Training time"], [
        ["Isolation Forest", "50,000", "67.2", "77.5", "1.00%", "1,309 ms"],
        ["Autoencoder", "50,000", "85.6", "100.0", "2.50%", "3,194 ms"],
        ["Local Outlier Factor", "50,000", "89.2", "100.0", "0.82%", "5,595 ms"],
    ], [1.45, 0.75, 0.75, 0.85, 1.0, 1.2], "Table 6. Current unsupervised model results on 50,000 rows")
    add_body(doc, "Local Outlier Factor produced the strongest current EM-AUC score, while Autoencoder also showed strong label-free separation. Isolation Forest trained fastest but had a lower EM-AUC. These figures describe diagnostic separation and anomaly behaviour on the available snapshot; they do not prove fraud accuracy because the underlying unsupervised dataset has no confirmed fraud label. Selection must also consider growth stability, throughput, alert volume, and operational latency.")
    add_heading(doc, "6.3", "Explainable Transaction Scenario")
    add_body(doc, "In a tested transaction scenario, the three selected ML detectors all flagged an anomaly. The final score nevertheless depended on the configured 65% ML allocation plus customer, peer, and rule contributions. Peer scoring identified the amount as above the peer average and the frequency above the 95th percentile; the rule layer detected turnover and amount deviations. The result reached MEDIUM and produced ALLOW_AND_ALERT with a case. This demonstrates why model votes and final risk must be displayed separately.")
    add_heading(doc, "6.4", "Automated Validation Summary")
    add_table(doc, ["Validation", "Observed result"], [
        ["Spring Boot compilation", "319 Java source files compiled successfully"],
        ["Python automated suite", "10 tests passed in the latest focused ML validation run"],
        ["Angular production build", "Bundle generated successfully; only existing size-budget warnings"],
        ["Security design", "JWT filter, BCrypt, authenticated APIs, ADMIN/AML_ADMIN write restrictions"],
        ["Case safeguards", "False-positive/STR state conflicts rejected; pagination and summary queries implemented"],
        ["Research integrity", "Oldest-first partitions, chronological holdout, immutable dataset checksum"],
    ], [2.0, 4.0], "Table 7. Automated validation summary")
    add_heading(doc, "6.5", "Performance and Scalability Observations")
    add_body(doc, "The training screenshot shows all three unsupervised models completing on 50,000 rows within approximately 1.3 to 5.6 seconds on the development environment. This is not a bank-scale benchmark, but it validates the pipeline and exposes relative cost. Real-time latency is reduced by reusing loaded artifacts, persisting feature vectors, limiting historical reads to point-in-time windows, and moving case/STR work to an asynchronous after-commit listener. Larger deployments should benchmark p50, p95, and p99 latency under concurrent load.")
    add_heading(doc, "6.6", "Scientific Interpretation Limits")
    add_body(doc, "EM-AUC, separation, score skewness, anomaly-rate control, Silhouette score, Davies-Bouldin index, and Calinski-Harabasz index are not fraud accuracy. The first group describes anomaly score structure; the clustering indices describe geometry only when a genuine clustering experiment is performed. Precision, recall, F1, and false-positive rate require reviewed outcomes. The UI and report explicitly preserve this distinction.")

    # Chapter 7
    add_chapter(doc, 7, "Problems and Solutions")
    add_heading(doc, "7.1", "Problem-to-Solution Summary")
    add_table(doc, ["Problem", "Implemented solution", "Result"], [
        ["No reliable fraud labels", "Global unsupervised mode with Isolation Forest, Autoencoder, and LOF; label-free metrics", "Useful anomaly research without false accuracy claims"],
        ["Some clients have reviewed labels", "Separate supervised mode with XGBoost, Random Forest, Logistic Regression and supervised metrics", "One platform supports both scientifically valid workflows"],
        ["Model behaviour changes with data volume", "Oldest-first 10%, 25%, 50%, 100% growth study with chronological holdout", "Stability and scalability become visible before selection"],
        ["Raw scores are incompatible", "Normalize every model and component to 0-1 before weighted aggregation", "All configured contributions total exactly 100%"],
        ["ML alone misses AML obligations", "Separate deterministic AML rules with hard override support", "Mandatory controls remain transparent and auditable"],
        ["One customer differs from population", "Customer behaviour scorer with amount, velocity, novelty, time-gap, and confidence", "Personal deviations contribute independently"],
        ["Peer comparison unavailable for sparse groups", "Occupation-age group, parent segment, then global fallback with confidence reduction", "Every transaction can receive a bounded peer result"],
        ["Cold start creates unstable profiles", "Minimum history gate and confidence multiplier; suspicious rows excluded from trusted learning", "Immature or contaminated profiles have less authority"],
        ["Training and prediction feature drift", "Persisted versioned feature vectors and exact feature-column artifact", "Training and scoring use the same schema"],
        ["Millions of rows cannot be sent as JSON", "Keyset export, Parquet parts, checksums, capped evaluation, selected-model training", "Memory and network usage remain bounded"],
        ["Transaction checks were slow", "Cached model artifacts, persisted features, bounded history, asynchronous post-commit case work", "Decision path avoids retraining and case I/O"],
        ["Model votes conflicted with final score", "Display votes as diagnostics and calculate final score from explicit policy weights", "Users can explain why 3/3 anomalies may still be below threshold"],
        ["Weight sliders changed unrelated values", "Normalize within each component and model section; enforce totals server-side", "Policy editing is predictable and cannot exceed 100%"],
        ["Cases became expensive to list", "Server-side pagination, filtering, count summaries, and reduced N+1 loading", "Case management scales beyond small demonstrations"],
        ["Alerts did not complete the AML workflow", "Automatic cases, notes, false-positive outcome, manual cases, and draft STR XML", "Suspicious decisions become reviewable operational records"],
        ["Administrative APIs were exposed", "JWT authentication, BCrypt, role authorization, deny-by-default security", "Sensitive operations require authorized identities"],
        ["Training could silently change production", "Immutable model registry, validation, version locking, controlled promotion, rollback", "Model lifecycle is auditable and reversible"],
    ], [1.65, 2.9, 1.45], "Table 8. Problem-to-solution matrix")
    add_heading(doc, "7.2", "How Scalability Is Solved")
    add_body(doc, "Scalability is addressed by separating the online decision path from offline training. Online scoring reads bounded point-in-time state and already loaded artifacts. Offline training uses closed windows and immutable Parquet snapshots. The UI sends a training-run identifier rather than millions of feature rows. Evaluation rows are capped, large algorithms can use bounded samples, and candidate creation is manual and asynchronous. SQL Server remains the governance store rather than the numerical training engine.")
    add_heading(doc, "7.3", "Kafka Integration Roadmap")
    add_body(doc, "Kafka is not required for the current prototype, but it is the next transport layer for millions of daily events. A SQL outbox should publish committed transaction events to account-keyed topics. Independent consumers can update customer, peer, and velocity state; write feature events; request predictions; persist model results; and create cases. Account-keyed partitioning preserves ordering for one customer while consumer groups provide horizontal scaling, replay, and back-pressure isolation. Kafka transports events; it does not replace SQL governance or train models by itself.")
    add_heading(doc, "7.4", "Why the Final Design Is More Flexible")
    add_body(doc, "The client decides the learning methodology once at system level, chooses participating production models based on comparison evidence, assigns model and layer weights, and manually starts training when data is ready. The same business layers continue to operate in both modes. This avoids forcing a bank with no labels into supervised metrics and avoids preventing a bank with reviewed labels from using probability-based classifiers.")

    # Chapter 8
    add_chapter(doc, 8, "Conclusion and Future Work")
    add_heading(doc, "8.1", "Conclusion")
    add_body(doc, "The project delivers a complete prototype for transaction ingestion, feature persistence, dual-mode machine learning, model growth comparison, explainable layered risk scoring, governed training, authentication, and AML case management. Its most important design decision is methodological honesty: anomalies are treated as suspicious evidence, not confirmed fraud, unless reviewed labels justify supervised evaluation.")
    add_body(doc, "The system solves the central project problems by using point-in-time features, immutable chronological datasets, mode-specific models and metrics, normalized risk weights, transparent customer and peer scorers, deterministic rules, auditable model versions, and a reviewer workflow. The Angular interface makes these decisions visible rather than hiding them behind one prediction label.")
    add_heading(doc, "8.2", "Current Limitations")
    add_bullets(doc, [
        "The 50,000-row result is a development-scale experiment, not a production bank benchmark.",
        "Unsupervised results cannot provide true precision, recall, F1, or confirmed false-positive rate.",
        "The generated supervised dataset is synthetic and must not replace reviewed institutional labels for real deployment.",
        "Kafka, distributed feature state, centralized observability, and infrastructure deployment are architectural roadmaps rather than completed production services.",
        "Draft STR XML requires validation against the exact schema and submission process of the relevant regulator.",
        "Model calibration, fairness by segment, privacy controls, penetration testing, and disaster recovery require further work.",
    ])
    add_heading(doc, "8.3", "Future Work")
    add_bullets(doc, [
        "Integrate Kafka through a transactional SQL outbox and account-keyed consumer groups.",
        "Add online feature stores or bounded materialized state for customer and peer statistics.",
        "Collect reviewer labels and add calibration curves, threshold optimization, recall at fixed false-positive rates, and drift by business period.",
        "Run reproducible load tests with millions of rows and concurrent transaction scoring, reporting p50, p95, and p99 latency.",
        "Introduce approved clustering experiments using MiniBatchKMeans or BIRCH and display PCA/UMAP projections with Silhouette, Davies-Bouldin, and Calinski-Harabasz as geometric diagnostics.",
        "Add containerization, centralized logs and metrics, secrets management, backup, and automated deployment pipelines.",
        "Validate STR XML and workflow against the applicable jurisdiction and add digital approval controls.",
    ])
    add_heading(doc, "8.4", "Final Remarks")
    add_body(doc, "The completed platform demonstrates that a strong fraud and AML project is not only a model-selection exercise. It is a data-integrity, scalability, explainability, governance, and human-review problem. By making each of those concerns explicit, the project provides a credible foundation for further academic study and future production engineering.")

    # References
    doc.add_page_break(); add_front_heading(doc, "References")
    references = [
        "[1] F. T. Liu, K. M. Ting, and Z.-H. Zhou, 'Isolation Forest,' Proceedings of the 8th IEEE International Conference on Data Mining, 2008.",
        "[2] M. M. Breunig, H.-P. Kriegel, R. T. Ng, and J. Sander, 'LOF: Identifying Density-Based Local Outliers,' Proceedings of ACM SIGMOD, 2000.",
        "[3] T. Chen and C. Guestrin, 'XGBoost: A Scalable Tree Boosting System,' Proceedings of the 22nd ACM SIGKDD Conference, 2016.",
        "[4] L. Breiman, 'Random Forests,' Machine Learning, vol. 45, pp. 5-32, 2001.",
        "[5] D. P. Kingma and M. Welling, 'Auto-Encoding Variational Bayes,' International Conference on Learning Representations, 2014.",
        "[6] Spring, 'Spring Boot Reference Documentation' and 'Spring Security Reference,' accessed August 2026.",
        "[7] Angular, 'Angular Documentation,' accessed August 2026.",
        "[8] FastAPI, 'FastAPI Documentation,' accessed August 2026.",
        "[9] scikit-learn developers, 'scikit-learn User Guide,' accessed August 2026.",
        "[10] Apache Software Foundation, 'Apache Parquet Documentation' and 'Apache Kafka Documentation,' accessed August 2026.",
        "[11] Microsoft, 'SQL Server Documentation,' accessed August 2026.",
    ]
    for ref in references:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.hanging_indent = Inches(0.25); p.paragraph_format.space_after = Pt(5); add_text(p, ref, size=10.5)

    # Appendices
    doc.add_page_break(); add_front_heading(doc, "Appendix A: Main User Workflow")
    add_numbered(doc, [
        "Log in with an authorized account.",
        "Open Settings and choose supervised or unsupervised learning.",
        "Upload CSV/Excel data and inspect the batch result.",
        "Close the relevant business dates in Training Operations.",
        "Create a training run, generate the immutable dataset, and wait for DATASET_READY.",
        "Choose the desired active-mode models and start training.",
        "Open the corresponding comparison page and review quality, stability, anomaly/classification metrics, and runtime.",
        "Open Risk Policy, select production models, and configure normalized layer/model weights and thresholds.",
        "Use Transaction Check or the transaction API to score new activity.",
        "Review automatically generated cases; mark false positives or generate/download STR XML as appropriate.",
    ])
    doc.add_page_break(); add_front_heading(doc, "Appendix B: Selected API Endpoints")
    api_rows = [
        ["POST", "/api/auth/login", "Authenticate and obtain JWT"],
        ["POST", "/api/v1/uploads/transactions", "Upload CSV or Excel transactions"],
        ["POST", "/api/v1/transactions", "Persist and score one transaction"],
        ["POST", "/api/v1/aml/training-runs/pipeline/start", "Create governed training run"],
        ["POST", "/api/v1/aml/training-runs/{id}/dataset", "Generate immutable feature dataset"],
        ["POST", "/api/v1/aml/growth-analysis/training-runs/{id}", "Run chronological growth analysis"],
        ["GET/PUT", "/api/v1/anomaly-model-comparisons/risk-policy", "Read or save production risk policy"],
        ["GET/PUT", "/api/v1/anomaly-model-comparisons/model-tuning", "Read or save active-mode tuning"],
        ["GET", "/api/cases", "Paginated case search"],
        ["POST", "/api/cases/{id}/false-positive", "Record false-positive outcome"],
        ["POST", "/api/cases/{id}/str-xml", "Generate draft STR XML"],
    ]
    add_table(doc, ["Method", "Endpoint", "Purpose"], api_rows, [0.8, 3.2, 2.0], "Table B1. Selected API endpoints")
    doc.add_page_break(); add_front_heading(doc, "Appendix C: Configuration Principles")
    add_bullets(doc, [
        "Top-level risk-layer weights must total 100%.",
        "Enabled model weights inside the ML ensemble must total 100%.",
        "Supervised mode requires reviewed labels; unresolved labels are excluded from training.",
        "Unsupervised metrics must never be described as confirmed fraud accuracy.",
        "Training uses closed business dates and a fixed cutoff.",
        "No training run automatically activates or promotes a model.",
        "MEDIUM creates a case; HIGH creates a case and draft STR XML under the current policy.",
        "The backend is the security authority even when UI controls are hidden or disabled.",
    ])

    # Document settings and metadata.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.core_properties.title = "Intelligent Fraud Transaction Detection and AML Case Management System"
    doc.core_properties.subject = "Final project report"
    doc.core_properties.author = "<Student Name>"
    doc.core_properties.keywords = "fraud detection, AML, machine learning, case management"
    doc.core_properties.comments = "Generated from the Evening Program Project Report Template and verified against project implementation evidence."
    doc.save(FINAL)
    return FINAL


if __name__ == "__main__":
    print(build_report())
